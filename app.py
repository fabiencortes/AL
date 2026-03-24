from utils_mail_universal import parse_mail_to_navette_universal
from urllib.parse import unquote
# ============================================================
# 🐞 DEBUG GLOBAL (console) — activable via env AL_DEBUG=1
# ============================================================
import os as _os
import sys as _sys
import time as _time
import base64


try:
    from utils import debug_print, debug_enabled
except Exception:
    def debug_enabled(): return True
    def debug_print(*a, **k):
        try:
            print(*a, **k, flush=True)
        except Exception:
            pass

debug_print("🚀 APP LOADED:", __file__)
debug_print("🐍 PYTHON:", _sys.executable)
debug_print("📁 CWD:", _os.getcwd())

# ============================================================
#   AIRPORTS LINES – APP.PLANNING – VERSION OPTIMISÉE 2025
#   BLOC 1/7 : IMPORTS, CONFIG, HELPERS, SESSION
# ============================================================
DEBUG_SAFE_MODE = True
AUTO_SYNC_ENABLED = False  # 🔒 Synchro uniquement manuelle
import os
import io
import sqlite3
from datetime import datetime, date, timedelta
from typing import Dict, Any, List
try:
    from database import init_time_rules_table
except Exception as _e:
    debug_print('⚠️ import init_time_rules_table failed:', _e)
    def init_time_rules_table():
        return

from database import init_actions_table
from database import mark_navette_confirmed
from database import ensure_ack_columns
from pathlib import Path

import uuid
import streamlit.components.v1 as components
import math
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import pandas as pd
import requests
from openpyxl import load_workbook
from io import BytesIO
import streamlit as st
import database as _database
try:
    debug_print('📦 DATABASE MODULE:', _database.__file__)
except Exception:
    pass
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

from database import (
    get_planning,
    get_chauffeurs,
    get_chauffeur_planning,
    search_client,
    get_row_by_id,
    insert_planning_row,
    update_planning_row,
    apply_row_update,
    delete_planning_row,
    get_planning_columns,
    get_connection,
    init_indispo_table,
    create_indispo_request,
    get_indispo_requests,
    set_indispo_status,
    ensure_planning_updated_at_column,
    ensure_km_time_columns,
    init_chauffeur_ack_table,
    get_chauffeur_last_ack,
    set_chauffeur_last_ack,
    init_flight_alerts_table,
    ensure_flight_alerts_time_columns,
    should_notify_flight_change,
    upsert_flight_alert,
    sqlite_safe,
    get_last_sync_time,
    set_last_sync_time,
    ensure_meta_table,
    get_meta,
    set_meta,
    rebuild_planning_db_from_two_excel_files,
    ensure_planning_confirmation_and_caisse_columns,
    ensure_superseded_column,
    ensure_urgence_columns,
    find_time_conflicts,
    get_urgences,
    set_urgence_status,
)
from database import (
    split_chauffeurs,
    ensure_planning_row_key_column,
    ensure_planning_row_key_index,
    get_planning_table_columns,
    get_chauffeurs_phones,
    ensure_caisse_columns,
    ensure_planning_audit_table,
    ensure_chauffeur_messages_table,
    ensure_admin_reply_read_column,
    ensure_admin_reply_columns,
    ensure_excel_sync_column,
    cleanup_orphan_planning_rows,
    ensure_ch_manual_column,
    list_pending_actions,
    mark_actions_done,
    unlock_rows_by_row_keys,
)
from utils import add_excel_color_flags_from_dropbox, log_event, render_logs_ui, format_mail_navette_v2, parse_mail_to_navette_v2, detect_dest_code, suggest_heures_from_rules, parse_mail_to_navette_v2, format_mail_navette_v2, parse_mail_to_navette_v2_cached
def rebuild_planning_views():
    """
    🔒 Version ULTIME (corrigée)
    - Planning 7j = UNIQUEMENT les lignes dans les 7 jours
      (navettes + congés + indispos)
    - Compatible DATE_ISO, DATE dd/mm/YYYY, DATE ISO texte
    - Rapide (pas de full scan)
    """

    with get_connection() as conn:
        cur = conn.cursor()

        cur.execute("DROP VIEW IF EXISTS planning_day")
        cur.execute("DROP VIEW IF EXISTS planning_7j")
        cur.execute("DROP VIEW IF EXISTS planning_full")

        # ======================================================
        # Vue FULL (brute)
        # ======================================================
        cur.execute("""
            CREATE VIEW planning_full AS
            SELECT * FROM planning
        """)

        # ======================================================
        # Expression DATE robuste
        # Priorité :
        #   1) DATE_ISO si présent
        #   2) DATE dd/mm/YYYY
        #   3) DATE déjà ISO
        # ======================================================
        date_expr = """
            CASE
                WHEN DATE_ISO IS NOT NULL AND DATE_ISO != '' THEN DATE_ISO
                WHEN LENGTH(DATE) = 10 AND substr(DATE,3,1)='/' THEN
                    substr(DATE,7,4)||'-'||substr(DATE,4,2)||'-'||substr(DATE,1,2)
                ELSE DATE
            END
        """

        # ======================================================
        # Vue planning_7j
        # 👉 TOUT ce qui tombe dans les 7 jours
        # ======================================================
        cur.execute(f"""
            CREATE VIEW planning_7j AS
            SELECT *
            FROM planning
            WHERE
                date({date_expr})
                BETWEEN date('now') AND date('now','+6 day')
        """)

        # ======================================================
        # Vue planning_day (aujourd’hui)
        # ======================================================
        cur.execute(f"""
            CREATE VIEW planning_day AS
            SELECT *
            FROM planning
            WHERE
                date({date_expr}) = date('now')
        """)

        conn.commit()
# ============================================================
#   🔒 CRÉATION DES VUES PLANNING — UNE SEULE FOIS (SAFE)
# ============================================================

def ensure_planning_views_once():
    """
    Crée les vues planning_* UNIQUEMENT si elles n'existent pas.
    Évite tout deadlock SQLite au démarrage Streamlit.
    """
    with get_connection() as conn:
        cur = conn.cursor()

        cur.execute("""
            SELECT name
            FROM sqlite_master
            WHERE type='view' AND name='planning_7j'
        """)

        exists = cur.fetchone()

        if not exists:
            print("🛠️ Création des vues planning_*", flush=True)
            rebuild_planning_views()
        else:
            print("✅ Vues planning déjà existantes", flush=True)

def init_sqlite_pragmas():
    with get_connection() as conn:
        conn.execute("PRAGMA busy_timeout=5000;")
        conn.execute("PRAGMA foreign_keys=ON;")
        conn.execute("PRAGMA synchronous=NORMAL;")

def init_db_once():
    if st.session_state.get("db_init_done"):
        return

    print("▶️ init_db_once START", flush=True)

    # PRAGMAS SQLITE
    init_sqlite_pragmas()
    print("▶️ pragmas OK", flush=True)

    # COLONNES / TABLES SÉCURITÉ
    ensure_planning_confirmation_and_caisse_columns()
    ensure_planning_audit_table()
    ensure_chauffeur_messages_table()
    ensure_admin_reply_read_column()
    ensure_planning_updated_at_column()
    ensure_admin_reply_columns()
    ensure_excel_sync_column()
    ensure_ch_manual_column()
    print("▶️ ensure columns OK", flush=True)

    # 🔒 VUES SQLITE (UNE SEULE FOIS)
    ensure_planning_views_once()

    st.session_state.db_init_done = True
    print("▶️ init_db_once DONE", flush=True)


def get_device_bound_login() -> str | None:
    return str(st.session_state.get("device_bound_login") or "").strip().lower() or None


def set_device_bound_login(login: str):
    st.session_state["device_bound_login"] = str(login or "").strip().lower()

# ============================================================
#   SESSION STATE
# ============================================================

def init_session_state():
    defaults = {
        # 🔐 Auth
        "logged_in": False,
        "username": None,
        "role": None,
        "chauffeur_code": None,
        "remember_me": True,

        # 📅 UI planning
        "planning_start": date.today(),
        "planning_end": date.today() + timedelta(days=6),
        "planning_sort_choice": "Date + heure",

        # 🔄 Sync & refresh
        "sync_running": False,
        "last_auto_sync": 0,

        # 🧭 Rafraîchissement par onglet
        "tab_refresh": {},   # ex: {"admin": 123456789}
    }

    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ===========================================================

# ============================================================
#   🔁 SOFT REFRESH / DEBOUNCE (anti-refresh brutal)
# ============================================================
def request_soft_refresh(tab_key: str, *, clear_cache: bool = True, mute_autosync_sec: int = 5):
    """Demande un rafraîchissement contrôlé (1 seul rerun) pour un onglet.
    - clear_cache : invalide cache_data (pour voir la DB tout de suite)
    - mute_autosync_sec : évite que l'auto-sync Excel se lance en même temps
    """
    try:
        if clear_cache:
            st.cache_data.clear()
    except Exception:
        pass

    st.session_state.setdefault("tab_refresh", {})
    st.session_state["tab_refresh"][tab_key] = time.time()

    # 🔇 évite collision UI update vs auto-sync
    try:
        st.session_state["_mute_autosync_until"] = time.time() + int(mute_autosync_sec or 0)
    except Exception:
        pass


def consume_soft_refresh(tab_key: str):
    """À appeler au début du rendu de l'onglet : déclenche le rerun UNE fois."""
    ts = st.session_state.get("tab_refresh", {}).get(tab_key, 0) or 0
    seen_key = f"_tab_refresh_seen_{tab_key}"
    seen = st.session_state.get(seen_key, 0) or 0
    if ts and ts > seen:
        st.session_state[seen_key] = ts
        st.rerun()

#   CONFIG UTILISATEURS
#   (admins, restreints, chauffeurs GSM)
# ============================================================

USERS = {
    "fab":  {"password": "AL2025",  "role": "admin"},
    "oli":  {"password": "AL2025",  "role": "admin"},
    "leon": {"password": "GL2025", "role": "restricted"},

    # Comptes chauffeurs pour GSM
    "gg": {"password": "gg", "role": "driver", "chauffeur_code": "GG"},
    "fa": {"password": "fa", "role": "driver", "chauffeur_code": "FA"},
    "do": {"password": "do", "role": "driver", "chauffeur_code": "DO"},
    "ma": {"password": "ma", "role": "driver", "chauffeur_code": "MA"},
    "po": {"password": "po", "role": "driver", "chauffeur_code": "PO"},
    "gd": {"password": "gd", "role": "driver", "chauffeur_code": "GD"},
    "om": {"password": "om", "role": "driver", "chauffeur_code": "OM"},
    "ad": {"password": "ad", "role": "driver", "chauffeur_code": "AD"},
    "ro": {"password": "ro", "role": "driver", "chauffeur_code": "RO"},
    "ge": {"password": "ge", "role": "driver", "chauffeur_code": "GE"},
    "lillo": {"password": "lillo", "role": "driver", "chauffeur_code": "LILLO"},
    "jf": {"password": "jf", "role": "driver", "chauffeur_code": "JF"},
}

# Fallback si Feuil2 ne contient rien
CH_CODES = [
    "AU", "FA", "GD", "GG", "LL", "MA", "O", "RK", "RO", "SW", "NP", "DO",
    "OM", "AD", "CB", "CF", "CM", "EM", "GE", "HM", "JF", "KM", "LILLO",
    "MF", "WS", "PO"
]

# ============================================================
# 🔐 SESSION PERSISTANTE ROBUSTE (GSM / Streamlit App)
#    - client_id stable par appareil
#    - session stockée côté serveur (SQLite)
#    - relance URL une seule fois pour remonter client_id/session à Python
# ============================================================
import uuid
import json
import threading
import streamlit.components.v1 as components

LOGIN_PERSIST_KEY = "al_session"
LOGIN_CLIENT_KEY = "al_cid"
LOGIN_CLIENT_STORAGE_KEY = "al_cid_local"
LOGIN_BOOTSTRAP_FLAG = "al_bootstrap_done"
LOGIN_PERSIST_HOURS = 24 * 30  # 30 jours
LOGIN_SID_KEY = "al_sid"


def _js_quote(val: str) -> str:
    import json as _json
    return _json.dumps(str(val or ""))


def ensure_persistent_sessions_table():
    try:
        with get_connection() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS persistent_sessions (
                    client_id TEXT PRIMARY KEY,
                    sid TEXT,
                    login TEXT NOT NULL,
                    token TEXT NOT NULL,
                    remember_me INTEGER DEFAULT 1,
                    active INTEGER DEFAULT 1,
                    created_at TEXT,
                    updated_at TEXT,
                    expires_at TEXT
                )
                """
            )
            conn.execute("CREATE INDEX IF NOT EXISTS idx_persistent_sid ON persistent_sessions(sid)")
            conn.commit()
    except Exception as e:
        print(f"⚠️ ensure_persistent_sessions_table error: {e}", flush=True)


def get_client_id():
    # 0) si déjà en mémoire, on ne change JAMAIS
    v = str(st.session_state.get("client_id") or "").strip()
    if v:
        return v

    # 1) query param
    try:
        q = st.query_params.get(LOGIN_CLIENT_KEY)
        if isinstance(q, list):
            q = q[0] if q else None
        q = str(q or "").strip()
        if q:
            st.session_state["client_id"] = q
            return q
    except Exception:
        pass

    # 2) cookies (si dispo)
    try:
        ctx = getattr(st, "context", None)
        cookies = getattr(ctx, "cookies", None)
        if cookies:
            c = cookies.get(LOGIN_CLIENT_KEY)
            c = str(c or "").strip()
            if c:
                st.session_state["client_id"] = c
                return c
    except Exception:
        pass

    # 3) fallback : générer une fois
    cid = "cid-" + uuid.uuid4().hex[:16]
    st.session_state["client_id"] = cid
    return cid


def bootstrap_login_persistence():
    """
    Initialise la persistance login côté navigateur :
    - crée un client_id stable si absent
    - relit la session persistée
    - écrit cookies + localStorage
    - pousse les valeurs dans l'URL
    - force un reload unique si nécessaire
    """
    components.html(
        f"""
        <script>
        (function() {{
            const sessionKey = {_js_quote(LOGIN_PERSIST_KEY)};
            const clientStorageKey = {_js_quote(LOGIN_CLIENT_STORAGE_KEY)};
            const clientParam = {_js_quote(LOGIN_CLIENT_KEY)};
            const bootFlag = {_js_quote(LOGIN_BOOTSTRAP_FLAG)};
            const maxAge = {int(LOGIN_PERSIST_HOURS * 3600)};

            function getBestWin() {{
                try {{
                    if (window.top) return window.top;
                }} catch (e) {{}}
                try {{
                    if (window.parent) return window.parent;
                }} catch (e) {{}}
                return window;
            }}

            function cookieFlags(w) {{
                try {{
                    if (w.location && w.location.protocol === "https:") {{
                        return "; path=/; max-age=" + maxAge + "; SameSite=None; Secure";
                    }}
                }} catch (e) {{}}
                return "; path=/; max-age=" + maxAge + "; SameSite=Lax";
            }}

            function safeGetLocal(w, key) {{
                try {{
                    return w.localStorage.getItem(key) || "";
                }} catch (e) {{
                    return "";
                }}
            }}

            function safeSetLocal(w, key, val) {{
                try {{
                    w.localStorage.setItem(key, val);
                }} catch (e) {{}}
            }}

            function safeSetCookie(w, key, val) {{
                try {{
                    w.document.cookie = key + "=" + encodeURIComponent(val) + cookieFlags(w);
                }} catch (e) {{}}
            }}

            function ensureClientId(w) {{
                let cid = safeGetLocal(w, clientStorageKey);

                if (!cid) {{
                    cid = "cid-" + Math.random().toString(36).slice(2) + Date.now().toString(36);
                    safeSetLocal(w, clientStorageKey, cid);
                }}

                return cid || "";
            }}

            const w = getBestWin();
            const cid = ensureClientId(w);
            const sess = safeGetLocal(w, sessionKey);

            if (cid) {{
                safeSetCookie(w, clientParam, cid);
                safeSetLocal(w, clientStorageKey, cid);
            }}

            if (sess) {{
                safeSetCookie(w, sessionKey, sess);
                safeSetLocal(w, sessionKey, sess);
            }}

            let url;
            try {{
                url = new URL(w.location.href);
            }} catch (e) {{
                try {{
                    url = new URL(window.location.href);
                }} catch (e2) {{
                    return;
                }}
            }}

            let changed = false;

            if (cid && url.searchParams.get(clientParam) !== cid) {{
                url.searchParams.set(clientParam, cid);
                changed = true;
            }}

            if (sess && url.searchParams.get(sessionKey) !== sess) {{
                url.searchParams.set(sessionKey, sess);
                changed = true;
            }}

            const guardKey = bootFlag + ":" + (cid || "none");
            let already = false;

            try {{
                already = w.sessionStorage.getItem(guardKey) === "1";
            }} catch (e) {{}}

            if (changed && !already) {{
                try {{
                    w.sessionStorage.setItem(guardKey, "1");
                }} catch (e) {{}}

                try {{
                    w.location.replace(url.toString());
                    return;
                }} catch (e) {{}}
            }}

            if (!changed) {{
                try {{
                    w.sessionStorage.removeItem(guardKey);
                }} catch (e) {{}}
            }}
        }})();
        </script>
        """,
        height=0,
    )
def save_persistent_session(login: str, token: str, remember_me: bool = True):
    client_id = get_client_id()
    if not client_id or not login or not token or not remember_me:
        return False

    ensure_persistent_sessions_table()
    now = datetime.now()
    exp = now + timedelta(hours=LOGIN_PERSIST_HOURS)

    sid = str(uuid.uuid4())

    try:
        with get_connection() as conn:
            conn.execute(
                """
                INSERT INTO persistent_sessions (client_id, sid, login, token, remember_me, active, created_at, updated_at, expires_at)
                VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?)
                ON CONFLICT(client_id) DO UPDATE SET
                    sid=excluded.sid,
                    login=excluded.login,
                    token=excluded.token,
                    remember_me=excluded.remember_me,
                    active=1,
                    updated_at=excluded.updated_at,
                    expires_at=excluded.expires_at
                """,
                (
                    client_id,
                    sid,
                    str(login).strip().lower(),
                    str(token).strip(),
                    1 if remember_me else 0,
                    now.isoformat(timespec="seconds"),
                    now.isoformat(timespec="seconds"),
                    exp.isoformat(timespec="seconds"),
                ),
            )
            conn.commit()

        st.session_state["last_sid"] = sid
        return True
    except Exception as e:
        print(f"⚠️ save_persistent_session error: {e}", flush=True)
        return False


def clear_persistent_session():
    client_id = get_client_id()
    if not client_id:
        return
    try:
        ensure_persistent_sessions_table()
        with get_connection() as conn:
            conn.execute(
                "UPDATE persistent_sessions SET active=0, updated_at=? WHERE client_id=?",
                (datetime.now().isoformat(timespec='seconds'), client_id),
            )
            conn.commit()
    except Exception as e:
        print(f"⚠️ clear_persistent_session error: {e}", flush=True)


def get_login_cookie():
    # 1) Query params (si présents)
    try:
        val = st.query_params.get(LOGIN_PERSIST_KEY)
        if isinstance(val, list):
            val = val[0] if val else None
        if val:
            return str(val).strip()
    except Exception:
        pass

    # 2) Cookies (souvent indispensable sur l'app Streamlit)
    try:
        ctx = getattr(st, "context", None)
        cookies = getattr(ctx, "cookies", None)
        if cookies:
            val = cookies.get(LOGIN_PERSIST_KEY)
            if val:
                return str(val).strip()
    except Exception:
        pass

    return None


def load_persistent_session_from_server():
    client_id = str(get_client_id() or "").strip()
    if not client_id:
        return None

    ensure_persistent_sessions_table()

    try:
        with get_connection() as conn:
            row = conn.execute(
                """
                SELECT login, token, expires_at
                FROM persistent_sessions
                WHERE client_id = ?
                  AND active = 1
                  AND remember_me = 1
                LIMIT 1
                """,
                (client_id,),
            ).fetchone()
    except Exception as e:
        print(f"⚠️ load_persistent_session_from_server error: {e}", flush=True)
        return None

    if not row:
        return None

    login, token, expires_at = row

    try:
        if expires_at and datetime.fromisoformat(str(expires_at)) < datetime.now():
            clear_persistent_session()
            return None
    except Exception:
        pass

    login = str(login or "").strip().lower()
    token = str(token or "").strip()

    if not login or not token:
        return None

    return f"{login}|{token}"
def remember_trusted_device(token: str, username: str, role: str, chauffeur_code: str | None = None):
    try:
        device_key = _device_fingerprint()
        if not device_key:
            return

        data = _load_trusted_devices()
        now_ts = int(time.time())

        # Format nouveau: data[device_key] = {"entries": {username: {...}}, "updated_at": ...}
        row = data.get(device_key)
        if isinstance(row, dict) and "entries" in row and isinstance(row.get("entries"), dict):
            entries = row["entries"]
        elif isinstance(row, dict) and row.get("token"):
            # compat ancien format (1 seul user)
            old_u = str(row.get("username") or "").strip() or "unknown"
            entries = {
                old_u: {
                    "token": str(row.get("token") or "").strip(),
                    "username": old_u,
                    "role": str(row.get("role") or "").strip(),
                    "chauffeur_code": str(row.get("chauffeur_code") or "").strip(),
                    "exp": int(row.get("exp", 0) or 0),
                    "updated_at": int(row.get("updated_at", 0) or 0),
                }
            }
        else:
            entries = {}

        u = str(username or "").strip().lower()
        entries[u] = {
            "token": str(token or "").strip(),
            "username": u,
            "role": str(role or "").strip(),
            "chauffeur_code": str(chauffeur_code or "").strip(),
            "exp": now_ts + int(AL_SESSION_MAX_AGE),
            "updated_at": now_ts,
        }

        # nettoyage: garder seulement entrées valides
        cleaned = {}
        for k, v in entries.items():
            if not isinstance(v, dict):
                continue
            if int(v.get("exp", 0) or 0) < now_ts:
                continue
            tok = str(v.get("token", "") or "").strip()
            if not tok:
                continue
            if not parse_login_token(tok):
                continue
            cleaned[k] = v

        if cleaned:
            data[device_key] = {"entries": cleaned, "updated_at": now_ts}
        else:
            data.pop(device_key, None)

        _save_trusted_devices(data)

    except Exception:
        return
def get_trusted_device_token() -> str | None:
    try:
        device_key = _device_fingerprint()
        if not device_key:
            return None

        data = _load_trusted_devices()
        row = data.get(device_key) if isinstance(data, dict) else None
        if not isinstance(row, dict):
            return None

        now_ts = int(time.time())

        # ✅ Nouveau format
        if "entries" in row and isinstance(row.get("entries"), dict):
            entries = row["entries"]

            # nettoyage local
            valid = []
            for u, v in entries.items():
                if not isinstance(v, dict):
                    continue
                if int(v.get("exp", 0) or 0) < now_ts:
                    continue
                tok = str(v.get("token", "") or "").strip()
                if not tok:
                    continue
                if not parse_login_token(tok):
                    continue
                valid.append(tok)

            # 0 entrée => rien
            if not valid:
                data.pop(device_key, None)
                _save_trusted_devices(data)
                return None

            # ✅ Anti-mix: si plusieurs sessions possibles pour la même empreinte, on NE RESTORE PAS
            # (on oblige login manuel)
            if len(valid) > 1:
                return None

            return valid[0]

        # ✅ Ancien format (1 seule entrée)
        if int(row.get("exp", 0) or 0) < now_ts:
            data.pop(device_key, None)
            _save_trusted_devices(data)
            return None

        tok = str(row.get("token", "") or "").strip()
        if not tok or not parse_login_token(tok):
            data.pop(device_key, None)
            _save_trusted_devices(data)
            return None

        return tok

    except Exception:
        return None

def set_login_cookie(token: str):
    token = str(token or "").strip()
    if not token:
        return

    token_js = _js_quote(token)
    max_age = int(LOGIN_PERSIST_HOURS * 3600)

    # ✅ APP Streamlit: cookie + localStorage uniquement
    components.html(
        f"""
        <script>
        (function() {{
            const sessionKey = "{LOGIN_PERSIST_KEY}";
            const clientKey = "{LOGIN_CLIENT_STORAGE_KEY}";
            const clientParam = "{LOGIN_CLIENT_KEY}";
            const sessionValue = {token_js};
            const maxAge = {max_age};

            function getBestWin() {{
                try {{ if (window.top) return window.top; }} catch(e) {{}}
                try {{ if (window.parent) return window.parent; }} catch(e) {{}}
                return window;
            }}
            const w = getBestWin();

            function cookieFlags() {{
                try {{
                    if (w.location && w.location.protocol === "https:") {{
                        return "; path=/; max-age=" + maxAge + "; SameSite=None; Secure";
                    }}
                }} catch (e) {{}}
                return "; path=/; max-age=" + maxAge + "; SameSite=Lax";
            }}

            function safeGetLocal(key) {{
                try {{ return w.localStorage.getItem(key) || ""; }} catch(e) {{ return ""; }}
            }}
            function safeSetLocal(key, val) {{
                try {{ w.localStorage.setItem(key, val); }} catch(e) {{}}
            }}
            function safeSetCookie(key, val) {{
                try {{ w.document.cookie = key + "=" + encodeURIComponent(val) + cookieFlags(); }} catch(e) {{}}
            }}

            // client_id stable par appareil (localStorage)
            let cid = safeGetLocal(clientKey);
            if (!cid) {{
                cid = "cid-" + Math.random().toString(36).slice(2) + Date.now().toString(36);
                safeSetLocal(clientKey, cid);
            }}

            // session
            safeSetLocal(sessionKey, sessionValue);
            safeSetCookie(sessionKey, sessionValue);

            // client_id
            safeSetLocal(clientKey, cid);
            safeSetCookie(clientParam, cid);
        }})();
        </script>
        """,
        height=0,
    )
def clear_login_cookie():
    components.html(
        f"""
        <script>
        (function() {{
            const sessionKey = {_js_quote(LOGIN_PERSIST_KEY)};
            const clientParam = {_js_quote(LOGIN_CLIENT_KEY)};
            const w = (function() {{ try {{ return window.parent || window; }} catch(e) {{ return window; }} }})();

            function clearCookie(k) {{
                try {{
                    // On tente de supprimer en mode https (Secure) et en mode http
                    if (w.location && w.location.protocol === "https:") {{
                        w.document.cookie = k + "=; path=/; max-age=0; SameSite=None; Secure";
                    }}
                }} catch (e) {{}}
                try {{
                    w.document.cookie = k + "=; path=/; max-age=0; SameSite=Lax";
                }} catch (e) {{}}
            }}

            clearCookie(sessionKey);
            clearCookie(clientParam);

            try {{ w.localStorage.removeItem(sessionKey); }} catch (e) {{}}
            try {{
                const url = new URL(w.location.href);
                url.searchParams.delete(sessionKey);
                url.searchParams.delete(clientParam);
                w.history.replaceState({{}}, "", url.toString());
            }} catch (e) {{}}
        }})();
        </script>
        """,
        height=0,
    )

def restore_login_from_cookie():
    if st.session_state.get("logged_in"):
        return False

    raw = get_login_cookie()
    if not raw:
        raw = load_persistent_session_from_server()
    if not raw:
        return False

    try:
        raw = unquote(str(raw).strip())
    except Exception:
        raw = str(raw).strip()

    if not raw:
        return False

    try:
        login, token = raw.split("|", 1)
    except Exception:
        return False

    login = str(login or "").strip().lower()
    token = str(token or "").strip()
    if not login or not token:
        return False

    # ✅ Anti-mix : si cet appareil est déjà lié à un autre login, on refuse la restauration auto
    bound = get_device_bound_login()
    if bound and bound != login:
        # Optionnel: on peut nettoyer le cookie session pour éviter une boucle de mauvais restore
        # clear_login_cookie()
        return False

    user = USERS.get(login)
    if not user:
        return False

    st.session_state.logged_in = True
    st.session_state.username = login
    st.session_state.role = user.get("role")
    st.session_state.chauffeur_code = user.get("chauffeur_code")
    st.session_state.session_token = token
    st.session_state.remember_me = True

    # ✅ On "bind" l'appareil à ce login après une restauration réussie
    set_device_bound_login(login)

    return True
# ============================================================
#   LOGIN SCREEN
# ============================================================

from datetime import datetime

def login_screen():
    st.title("🚐 Airports-Lines — Planning chauffeurs (DB)")
    st.subheader("Connexion")

    login = st.text_input("Login", key="login_name")
    pwd = st.text_input("Mot de passe", type="password", key="login_pass")
    remember_me = st.checkbox(
        "Se rappeler de moi sur cet appareil",
        value=bool(st.session_state.get("remember_me", True)),
        key="remember_me_widget",
    )

    if st.button("Se connecter"):
        login_norm = str(login or "").strip().lower()
        user = USERS.get(login_norm)

        if user and user["password"] == pwd:
            token = str(uuid.uuid4())

            # ✅ Force un client_id stable pour cet appareil
            cid = str(get_client_id() or "").strip()

            st.session_state.logged_in = True
            st.session_state.username = login_norm
            st.session_state.role = user["role"]
            st.session_state.chauffeur_code = user.get("chauffeur_code")
            st.session_state.session_token = token
            st.session_state.remember_me = bool(remember_me)
            st.session_state["_persist_state_saved"] = False

            if remember_me:
                # ✅ 1) Sauvegarde serveur PAR APPAREIL (anti-mix)
                save_persistent_session(login_norm, token, True)

                # ✅ 2) Sauvegarde locale APP Streamlit : cookie + localStorage uniquement
                # (important: pas de query params pour éviter les sessions qui se mélangent)
                set_login_cookie(f"{login_norm}|{token}")

                # ✅ 3) Optionnel (fortement recommandé) : marquer l'appareil comme associé à ce login
                # (sert à refuser un restore si l'appareil essaie de reprendre un autre chauffeur)
                st.session_state["device_bound_login"] = login_norm
                st.session_state["device_bound_cid"] = cid

            else:
                clear_persistent_session()
                clear_login_cookie()
                st.session_state.pop("device_bound_login", None)
                st.session_state.pop("device_bound_cid", None)

            st.rerun()
        else:
            st.error("Identifiants incorrects.")

FLIGHT_ALERT_DELAY_MIN = 30  # seuil d’alerte retard (modifiable)

def init_all_db_once():
    if st.session_state.get("all_db_init_done"):
        return

    # 🧱 tables
    init_indispo_table()
    init_chauffeur_ack_table()
    init_flight_alerts_table()
    init_time_rules_table()
    init_actions_table()

    # 🧠 mémoire (mail / prix / alias)
    try:
        from database import init_price_memory_table, init_requester_memory_table, init_location_aliases_table
        init_price_memory_table()
        init_requester_memory_table()
        init_location_aliases_table()
    except Exception:
        pass

    # 🧱 colonnes
    ensure_planning_updated_at_column()
    ensure_km_time_columns()
    ensure_flight_alerts_time_columns()
    ensure_ack_columns()
    ensure_caisse_columns()
    ensure_urgence_columns()

    st.session_state["all_db_init_done"] = True

def extract_positive_int(val):
    """
    Retourne un entier > 0 si val contient un chiffre valide,
    sinon retourne None.
    """
    if val is None:
        return None

    s = str(val).strip()

    if not s:
        return None

    # On garde uniquement les chiffres
    if s.isdigit():
        n = int(s)
        return n if n > 0 else None

    return None

# ============================================================
#   COULEURS EXCEL -> FLAGS DB (GROUPAGE / PARTAGE / ATTENTE)
# ============================================================

YELLOW_RGBS = {"FFFFFF00", "FFFF00", "00FFFF00"}

def _cell_is_yellow(cell) -> bool:
    """
    Détecte le jaune Excel (fill, theme, indexed).
    Compatible Excel réel (pas théorique).
    """
    try:
        fill = cell.fill
        if fill is None or fill.patternType is None:
            return False

        fg = fill.fgColor
        if fg is None:
            return False

        # RGB direct
        if fg.type == "rgb" and fg.rgb:
            rgb = fg.rgb.upper()
            return rgb.endswith("FFFF00") or rgb in {"FFFFFF00", "00FFFF00"}

        # Indexed color (Excel ancien)
        if fg.type == "indexed":
            return fg.indexed in {5, 6}  # jaunes courants Excel

        # Theme color (Excel moderne)
        if fg.type == "theme":
            return True  # on considère thème = volontaire

        return False
    except Exception:
        return False

# ============================================================
#   BADGES VISUELS NAVETTES
# ============================================================

def navette_badges(row) -> str:
    badges = []

    def _flag(val) -> bool:
        try:
            return int(val or 0) == 1
        except Exception:
            return False

    if _flag(row.get("IS_GROUPAGE")):
        badges.append("🟡 Groupage")

    if _flag(row.get("IS_PARTAGE")):
        badges.append("🔵 Partage")

    if _flag(row.get("IS_ATTENTE")):
        badges.append("⭐ Attente")

    if _flag(row.get("IS_INDISPO")):
        badges.append("🚫 Indispo")

    if _flag(row.get("IS_NEW")):
        badges.append("🆕 Modifié")

    return " • ".join(badges)
def has_new_chauffeur_reply(row) -> bool:
    """
    True si le chauffeur a répondu et que la navette
    n'est pas encore confirmée par un admin.
    """
    try:
        return bool(row.get("ACK_AT")) and int(row.get("CONFIRMED") or 0) == 0
    except Exception:
        return False
def is_new_ack(prev_ack_at, new_ack_at) -> bool:
    """
    True si une réponse chauffeur vient d'arriver
    (ACK_AT était vide et devient remplie)
    """
    if not prev_ack_at and new_ack_at:
        return True
    return False
@st.cache_data
def get_real_chauffeurs_fast():
    with get_connection() as conn:
        df = pd.read_sql_query(
            "SELECT DISTINCT INITIALE FROM chauffeurs",
            conn,
        )
    return (
        df["INITIALE"]
        .dropna()
        .astype(str)
        .str.strip()
        .str.upper()
        .unique()
        .tolist()
    )


@st.cache_data
def load_planning_for_period(start, end):
    return get_planning(
        start_date=start,
        end_date=end,
        max_rows=5000,
        source="full",
    )


@st.cache_data
def get_ack_map_safe(chauffeurs):
    ack_map = {}
    for ch in chauffeurs:
        try:
            ack_map[ch] = get_chauffeur_last_ack(ch)
        except Exception:
            ack_map[ch] = None
    return ack_map


def was_modified_after_confirmation(row) -> bool:
    """
    True si la navette a été modifiée APRÈS confirmation admin.
    """
    try:
        if not row.get("CONFIRMED_AT") or not row.get("updated_at"):
            return False
        return row["updated_at"] > row["CONFIRMED_AT"]
    except Exception:
        return False


# ============================================================
# 🔁 SYNCHRONISATION AUTOMATIQUE INVISIBLE (PLANNING FUTUR)
# ============================================================

import time

if "last_auto_sync" not in st.session_state:
    # Empêche toute synchro auto au premier chargement
    st.session_state.last_auto_sync = time.time()


from concurrent.futures import ThreadPoolExecutor
import time

# ============================================================
#   🔄 BACKGROUND SYNC (Excel Dropbox -> DB) — non bloquant
# ============================================================
_SYNC_EXECUTOR = ThreadPoolExecutor(max_workers=1)

def _launch_background_excel_sync():
    """Lance une synchronisation Excel -> DB en arrière-plan (silencieuse).
    Ne bloque jamais l'UI.
    """
    try:
        fut = st.session_state.get("_bg_excel_sync_future")
        if fut is not None and not fut.done():
            return  # déjà en cours
    except Exception:
        pass

    def _job():
        try:
            # sync_planning_from_today fait déjà le check 'excel_last_modified' via meta.
            sync_planning_from_today(ui=False)
        except Exception as e:
            # silencieux côté UI ; log console uniquement
            print(f"⚠️ Background Excel sync error: {e}", flush=True)

    try:
        st.session_state["_bg_excel_sync_future"] = _SYNC_EXECUTOR.submit(_job)
        st.session_state["_bg_excel_sync_started_at"] = time.time()
    except Exception as e:
        print(f"⚠️ Unable to start background sync: {e}", flush=True)

def auto_sync_planning_if_needed():
    debug_print('⛔ auto_sync_planning_if_needed DISABLED (DEBUG MODE)')
    return



import os, json
from io import BytesIO
import pandas as pd
import requests
import streamlit as st

DROPBOX_FILE_PATH = "/Goldenlines/Planning 2026.xlsx"

# 🔗 Lien public Dropbox (édition Excel Online)
PLANNING_ONLINE_URL = "https://www.dropbox.com/scl/fi/3s712xzr85tijl3mcw5sy/Planning-2022.xlsx?rlkey=6zl8px1zzjz6lhcp9dutz25n3&st=sqs73z3s&dl=0"


import os
import requests

def load_planning_from_dropbox(sheet_name: str | None = None) -> pd.DataFrame:
    from utils import get_dropbox_excel_cached

    content = get_dropbox_excel_cached()
    if not content:
        return pd.DataFrame()


    bio = BytesIO(content)

    if sheet_name != "Feuil1":
        return pd.read_excel(bio, sheet_name=sheet_name, header=0, engine="openpyxl").fillna("")

    # 🔒 AUTO-DÉTECTION EN-TÊTE Feuil1
    df_raw = pd.read_excel(
        bio,
        sheet_name="Feuil1",
        header=None,
        engine="openpyxl",
    )

    header_row = None
    for i in range(min(10, len(df_raw))):
        row_vals = (
            df_raw.iloc[i]
            .astype(str)
            .str.strip()
            .str.upper()
            .tolist()
        )
        if "DATE" in row_vals and "HEURE" in row_vals:
            header_row = i
            break

    if header_row is None:
        st.error("❌ Impossible de détecter la ligne d’en-tête (DATE / HEURE).")
        return pd.DataFrame()

    df = pd.read_excel(
        BytesIO(content),
        sheet_name="Feuil1",
        header=header_row,
        engine="openpyxl",
    )

    return df.fillna("")




def get_dropbox_file_last_modified() -> datetime | None:
    try:
        token = os.environ.get("DROPBOX_TOKEN")
        if not token:
            return None

        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json",
        }

        data = {
            "path": "/Goldenlines/Planning 2026.xlsx"
        }

        r = requests.post(
            "https://api.dropboxapi.com/2/files/get_metadata",
            headers=headers,
            json=data,
            timeout=20,
        )
        r.raise_for_status()

        info = r.json()
        return datetime.fromisoformat(
            info["server_modified"].replace("Z", "+00:00")
        )

    except Exception:
        return None
from utils import download_dropbox_excel_bytes as _download_dropbox_excel_bytes
from utils import upload_dropbox_excel_bytes as _upload_dropbox_excel_bytes


def download_dropbox_excel_bytes(path: str = "/Goldenlines/Planning 2026.xlsx") -> bytes | None:
    """Wrapper: télécharge le fichier Excel depuis Dropbox (bytes)."""
    try:
        return _download_dropbox_excel_bytes(path)
    except Exception as e:
        print(f"⚠️ Dropbox download error: {e}", flush=True)
        return None


def upload_dropbox_excel_bytes(content: bytes, path: str = "/Goldenlines/Planning 2026.xlsx") -> bool:
    """Wrapper: upload (overwrite) le fichier Excel sur Dropbox."""
    try:
        _upload_dropbox_excel_bytes(content, path)
        return True
    except Exception as e:
        print(f"⚠️ Dropbox upload error: {e}", flush=True)
        return False


# ============================================================
#   📤 EXPORT DB -> EXCEL (Dropbox) — SANS CONFLIT
# ============================================================
def export_db_changes_to_excel_dropbox(row_ids: list[int] | None = None) -> bool:
    """Répercute certaines modifications DB vers l'Excel Dropbox, sans écraser l'existant.

    Stratégie SAFE :
    - Télécharger la dernière version Excel
    - Identifier les lignes Feuil1 via row_key (même algo que DB)
    - Appliquer uniquement :
        * CAISSE_PAYEE : couleur de la cellule 'Caisse' (vert si payé, rouge si non)
    - Upload overwrite (fichier complet mis à jour)
    """
    try:
        from openpyxl import load_workbook
        from openpyxl.styles import PatternFill
    except Exception:
        return False

    # 1) Lire DB (lignes concernées)
    with get_connection() as conn:
        if row_ids:
            q = ",".join(["?"] * len(row_ids))
            rows = conn.execute(f"SELECT * FROM planning WHERE id IN ({q})", [int(x) for x in row_ids]).fetchall()
            cols = [d[0] for d in conn.execute("SELECT * FROM planning LIMIT 1").description] if rows else []
        else:
            rows = conn.execute(
                """
                SELECT *
                FROM planning
                WHERE COALESCE(CAISSE_PAYEE,0) = 1
                  AND COALESCE(DATE_ISO,'') >= date('now','-60 day')
                """
            ).fetchall()
            cols = [d[0] for d in conn.execute("SELECT * FROM planning LIMIT 1").description] if rows else []

    if not rows or not cols:
        return False

    db_rows = [dict(zip(cols, r)) for r in rows]

    # 2) Télécharger Excel
    content = download_dropbox_excel_bytes()
    if not content:
        return False

    wb = load_workbook(BytesIO(content))
    if "Feuil1" not in wb.sheetnames:
        return False

    ws = wb["Feuil1"]

    # 3) Détecter header row (DATE/HEURE)
    header_row = None
    for i in range(1, 15):
        vals = [str(ws.cell(row=i, column=c).value or "").strip().upper() for c in range(1, 40)]
        if "DATE" in vals and "HEURE" in vals:
            header_row = i
            break
    if header_row is None:
        return False

    # mapping col name -> col idx
    headers = {}
    for c in range(1, 60):
        v = ws.cell(row=header_row, column=c).value
        if v is None:
            continue
        name = str(v).strip()
        if name:
            headers[name.upper()] = c

    # Colonnes attendues
    c_date = headers.get("DATE")
    c_heure = headers.get("HEURE")
    c_bdc = headers.get("NUM BDC") or headers.get("NUM_BDC") or headers.get("BDC") or headers.get("NUM. BDC")
    c_nom = headers.get("NOM") or headers.get("CLIENT")
    c_caisse = headers.get("CAISSE") or headers.get("Caisse".upper()) or headers.get("MONTANT") or headers.get("MONTANT €")

    if not (c_date and c_heure and c_nom):
        return False

    # 4) Construire index Excel : row_key -> excel_row_index
    from database import make_row_key_from_row

    excel_index = {}
    max_row = ws.max_row
    for r in range(header_row + 1, max_row + 1):
        date_v = ws.cell(row=r, column=c_date).value
        heure_v = ws.cell(row=r, column=c_heure).value
        nom_v = ws.cell(row=r, column=c_nom).value
        bdc_v = ws.cell(row=r, column=c_bdc).value if c_bdc else None

        # ligne vide
        if date_v is None and heure_v is None and nom_v is None:
            continue

        rk = make_row_key_from_row({
            "DATE": date_v,
            "HEURE": heure_v,
            "Num BDC": bdc_v,
            "NOM": nom_v,
        })
        if rk:
            excel_index[rk] = r

    # 5) Appliquer changements
    fill_green = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    fill_red = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

    changed = 0
    for row in db_rows:
        rk = row.get("row_key")
        if not rk:
            continue
        xr = excel_index.get(rk)
        if not xr:
            continue

        # CAISSE PAYEE -> couleur cellule Caisse
        try:
            payee = int(row.get("CAISSE_PAYEE") or 0)
        except Exception:
            payee = 0

        if c_caisse:
            cell = ws.cell(row=xr, column=c_caisse)
            cell.fill = fill_green if payee == 1 else fill_red
            changed += 1

    if changed <= 0:
        return False

    # 6) Upload
    out = BytesIO()
    wb.save(out)
    ok = upload_dropbox_excel_bytes(out.getvalue())

    # 7) Meta (optionnel)
    try:
        set_meta("excel_last_exported_at", datetime.now().isoformat())
    except Exception:
        pass

    return bool(ok)
def _extract_flight_number(text: str) -> str:
    """
    Extrait un numéro de vol type KL1703, QR193, SN1234
    """
    if not text:
        return ""
    m = re.search(r"\b([A-Z]{2}\s?\d{3,4})\b", text.upper())
    if m:
        return m.group(1).replace(" ", "")
    return ""
# ============================================================
#   DB — COLONNES FLAGS COULEURS (AUTO)
# ============================================================
def ensure_planning_color_columns():
    """
    Ajoute dans la table planning les colonnes de flags si elles n'existent pas.
    Compatible avec une DB déjà existante.
    """
    wanted = [
        "IS_GROUPAGE",
        "IS_PARTAGE",
        "IS_ATTENTE",
        "IS_NEW",        # 🟠 modif (CH orange Excel)
        "CH_COLOR",      # debug/trace
        "CAISSE_COLOR",  # debug/trace
    ]

    with get_connection() as conn:
        cur = conn.cursor()
        cur.execute("PRAGMA table_info(planning)")
        existing = {row[1] for row in cur.fetchall()}

        for col in wanted:
            if col not in existing:
                conn.execute(f'ALTER TABLE planning ADD COLUMN "{col}" TEXT')
        conn.commit()
# ============================================================
# NORMALISATION DES CODES CHAUFFEURS (FA, FA*, FADO, NPFA...)
# ============================================================

def normalize_ch_code(ch_raw: str) -> str:
    """
    Normalise un code chauffeur pour les emails / identifications simples.
    - FA*   -> FA
    - FA1*  -> FA1
    - FADO  -> FA
    - NPFA  -> NP
    - FAAD  -> FA
    """
    if not ch_raw:
        return ""

    code = str(ch_raw).upper().replace("*", "").strip()

    # Cas composés connus → chauffeur principal
    priority = ["FA1", "FA", "NP", "DO", "AD", "GG", "MA", "OM"]

    for ch in priority:
        if code.startswith(ch):
            return ch

    return code



def render_excel_modified_indicator():
    """Affiche un indicateur 'Excel modifié depuis X min' (source Dropbox)."""
    try:
        dt = get_dropbox_file_last_modified()
        if not dt:
            return
        # dt peut être timezone-aware; on le convertit en minutes
        now = datetime.now(dt.tzinfo) if getattr(dt, "tzinfo", None) else datetime.now()
        delta = now - dt.replace(tzinfo=now.tzinfo) if getattr(dt, "tzinfo", None) else now - dt
        mins = int(delta.total_seconds() // 60)
        if mins < 1:
            txt = "à l’instant"
        elif mins < 60:
            txt = f"il y a {mins} min"
        else:
            h = mins // 60
            m = mins % 60
            txt = f"il y a {h}h{m:02d}"
        st.caption(f"📄 Excel Dropbox modifié {txt} (source : Planning 2026.xlsx)")
    except Exception:
        pass

def render_last_sync_info():
    ts = st.session_state.get("last_auto_sync", 0)
    if not ts:
        return

    txt = datetime.fromtimestamp(ts).strftime("%H:%M")
    st.caption(f"🕒 Dernière synchro : {txt}")

def rebuild_db_fast(status):
    import os
    import shutil
    from database import ensure_indexes

    NEW_DB = "airportslines_NEW.db"
    MAIN_DB = "airportslines.db"
    BACKUP_DIR = "db_backups"

    status.update(label="📦 Bascule vers la nouvelle base…")

    os.makedirs(BACKUP_DIR, exist_ok=True)

    if os.path.exists(MAIN_DB):
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.move(
            MAIN_DB,
            os.path.join(BACKUP_DIR, f"airportslines_{ts}.db")
        )

    os.rename(NEW_DB, MAIN_DB)

    ensure_indexes()

    status.update(label="🎉 Base active remplacée", state="complete")
def format_navette_full_details(row, chauffeur_code: str) -> str:
    """
    Mail ADMIN – détail complet navette
    - IMMAT affichée si non vide
    - REH / SIÈGE affichés uniquement si chiffre > 0
    - Pas de Waze / Google Maps
    """
    import pandas as pd

    # =========================
    # DATE
    # =========================
    dv = row.get("DATE")
    if isinstance(dv, (datetime, date)):
        date_txt = dv.strftime("%d/%m/%Y")
    else:
        dtmp = pd.to_datetime(dv, dayfirst=True, errors="coerce")
        date_txt = dtmp.strftime("%d/%m/%Y") if not pd.isna(dtmp) else ""

    # =========================
    # HEURE
    # =========================
    heure_txt = normalize_time_string(row.get("HEURE")) or "??:??"

    # =========================
    # CLIENT / TRAJET
    # =========================
    nom = str(row.get("NOM", "") or "").strip()
    adr_full = build_full_address_from_row(
        pd.Series(row) if not isinstance(row, pd.Series) else row
    )
    tel_client = get_client_phone_from_row(
        pd.Series(row) if not isinstance(row, pd.Series) else row
    )

    # =========================
    # INFOS NAVETTE
    # =========================
    def g(*cols):
        for c in cols:
            v = row.get(c, "")
            if v is None:
                continue
            s = str(v).strip()
            if s and s.lower() != "nan":
                return s
        return ""

    route = g("DE/VERS", "Unnamed: 8", "DESTINATION", "ROUTE")
    vol = extract_vol_val(row, list(row.keys())) if hasattr(row, "keys") else ""
    pax = g("PAX")
    num_bdc = g("NUM_BDC", "NUM BDC", "BDC")
    paiement = g("PAIEMENT", "Paiement")
    caisse = g("CAISSE", "Caisse", "MONTANT", "Montant")

    # =========================
    # CHAUFFEUR
    # =========================
    ch_raw = str(row.get("CH", "") or "").strip()
    ch_norm = normalize_ch_for_phone(ch_raw)

    # =========================
    # VÉHICULE (RÈGLES STRICTES)
    # =========================
    immat = g("IMMAT", "PLAQUE", "IMMATRICULATION")
    reh_n = extract_positive_int(row.get("REH"))
    siege_n = extract_positive_int(row.get("SIEGE", "SIÈGE"))

    # =========================
    # CONSTRUCTION MAIL
    # =========================
    lines = []
    lines.append("📌 NAVETTE — DÉTAIL ADMIN")
    lines.append(f"📆 Date : {date_txt}")
    lines.append(f"⏱ Heure : {heure_txt}")

    if route:
        lines.append(f"🧭 Trajet : {route}")
    if vol:
        lines.append(f"✈️ Vol : {vol}")
    if pax:
        lines.append(f"👥 PAX : {pax}")
    if num_bdc:
        lines.append(f"🧾 BDC : {num_bdc}")

    lines.append("")
    lines.append(f"👨‍✈️ Chauffeur : {ch_raw}")

    if immat or reh_n or siege_n:
        lines.append("")
        lines.append("🚘 Véhicule :")
        if immat:
            lines.append(f"- Plaque : {immat}")
        if siege_n:
            lines.append(f"- Siège enfant : {siege_n}")
        if reh_n:
            lines.append(f"- REH : {reh_n}")

    lines.append("")
    lines.append(f"🧑 Client : {nom or '—'}")
    lines.append(f"📍 Adresse : {adr_full or '—'}")
    lines.append(f"📞 Client : {tel_client or '—'}")

    if paiement or caisse:
        lines.append("")
        lines.append("💳 Paiement :")
        if paiement:
            lines.append(f"- Type : {paiement}")
        if caisse:
            lines.append(f"- Montant caisse : {caisse}")

    return "\n".join(lines).strip()
# ======================================================
# 🔔 Compteur réponses chauffeur en attente (ADMIN)
# ======================================================
def count_pending_confirmations():
    df = get_planning(source="7j")
    if df is None or df.empty:
        return 0

    df = df[(df["IS_INDISPO"] == 0) & (df["CONFIRMED"] != 1)]
    return int(df["ACK_AT"].notna().sum())



# (render_tab_confirmation_chauffeur supprimé: doublon)


def format_navette_ack(row, ch_selected, trajet, probleme):
    import pandas as pd

    # =========================
    # DATE
    # =========================
    dv = row.get("DATE")
    if isinstance(dv, (datetime, date)):
        date_txt = dv.strftime("%d/%m/%Y")
    else:
        dtmp = pd.to_datetime(dv, dayfirst=True, errors="coerce")
        date_txt = dtmp.strftime("%d/%m/%Y") if not pd.isna(dtmp) else ""

    # =========================
    # HEURE
    # =========================
    heure_txt = normalize_time_string(row.get("HEURE")) or "??:??"

    # =========================
    # SENS + LIEU
    # =========================
    sens = str(row.get("Unnamed: 8", "") or "").strip().upper()
    if sens not in ("DE", "VERS"):
        sens = "DE"

    lieu = str(row.get("DESIGNATION", "") or "").strip()
    lieu = resolve_client_alias(lieu)

    sens_txt = f"{sens} ({lieu})" if lieu else sens

    # =========================
    # CLIENT
    # =========================
    nom = str(row.get("NOM", "") or "").strip()
    adr_full = build_full_address_from_row(row)
    tel_client = get_client_phone_from_row(row)

    # =========================
    # VÉHICULE (RÈGLES STRICTES)
    # =========================
    immat = str(row.get("IMMAT", "") or "").strip()
    reh_n = extract_positive_int(row.get("REH"))
    siege_n = extract_positive_int(row.get("SIEGE", "SIÈGE"))

    vehicule_lines = []
    if immat:
        vehicule_lines.append(f"Plaque : {immat}")
    if siege_n:
        vehicule_lines.append(f"Siège enfant : {siege_n}")
    if reh_n:
        vehicule_lines.append(f"REH : {reh_n}")

    vehicule_block = ""
    if vehicule_lines:
        vehicule_block = "\n🚘 Véhicule :\n" + "\n".join(vehicule_lines)

    # =========================
    # MAIL FINAL (SANS WAZE / MAPS)
    # =========================
    return f"""📆 {date_txt} | ⏱ {heure_txt}
👨‍✈️ Chauffeur : {ch_selected}
🚗 Sens : {sens_txt}

🧑 Client : {nom}
📍 Adresse : {adr_full}
📞 Client : {tel_client or "—"}{vehicule_block}

📝 Infos chauffeur :
Trajet : {trajet or "—"}
Problème : {probleme or "—"}
"""


def send_planning_confirmation_email(chauffeur: str, row, trajet: str, commentaire: str):
    """
    Mail admin = DÉTAIL COMPLET navette + en dessous la réponse du chauffeur.
    """

    subject = f"[CONFIRMATION PLANNING] {chauffeur}"

    navette_full = format_navette_full_details(row, chauffeur)

    # Réponse chauffeur (en dessous)
    ts = datetime.now().strftime("%d/%m/%Y %H:%M")
    reponse = f"""✅ RÉPONSE DU CHAUFFEUR
Horodatage : {ts}
Chauffeur : {chauffeur}

Trajet compris : {trajet or "—"}
Commentaire / problème : {commentaire or "—"}
"""

    body = navette_full + "\n\n" + reponse + "\nMessage envoyé depuis l’application Airports Lines."

    send_mail_admin(subject, body)


def is_navette_confirmed(row):
    """
    Une navette est confirmée si ACK_AT est renseigné en DB
    """
    return bool(row.get("ACK_AT"))


def _is_excel_cell_green(color_flag):
    """
    Retourne True si la cellule Excel est verte.
    color_flag est ce que tu stockes déjà via add_excel_color_flags_from_dropbox.
    """
    if not color_flag:
        return False

    color = str(color_flag).lower()

    return any(
        key in color
        for key in ("green", "vert", "#00", "lime")
    )
def format_chauffeur_colored(ch, confirmed, row=None):
    """
    Retourne le chauffeur avec icône selon l'état :
    🟢 confirmé
    🟠 en attente
    🟠🆕 nouvelle réponse chauffeur
    """
    ch = str(ch or "").strip().upper()

    if row is not None and has_new_chauffeur_reply(row):
        return f"🟠🆕 {ch}"

    if confirmed == 1:
        return f"🟢 {ch}"

    return f"🟠 {ch}"

def sync_planning_from_today(excel_sync_ts: str | None = None, *, ui: bool = True):
    """
    🔄 Synchronisation SAFE depuis aujourd’hui
    - ZÉRO doublon (row_key + INSERT OR IGNORE)
    - MAIS si Excel modifie une navette (date/heure/chauffeur/destination...) :
        ➜ l’ancienne version est supprimée/masquée
    - Congés / indispos détectés par HEURE -> HEURE_FIN
    - Dates Excel FR ("samedi 24 janvier 2026") supportées
    - Compatible DB existante
    """
    # --------------------------------------------------
    # 🔒 UI SAFE LOGGING (évite st.* en background thread)
    # --------------------------------------------------
    def _ui_warn(msg: str):
        if ui:
            st.warning(msg)
        else:
            print(f"⚠️ {msg}", flush=True)

    def _ui_error(msg: str):
        if ui:
            st.error(msg)
        else:
            print(f"❌ {msg}", flush=True)

    def _ui_info(msg: str):
        if ui:
            st.info(msg)
        else:
            print(f"ℹ️ {msg}", flush=True)
    # 🆔 Assure ROW_KEY UUID dans Excel (colonne ZX, masquée)
    try:
        from utils import ensure_excel_row_key_column
        ensure_excel_row_key_column(dropbox_path=DROPBOX_FILE_PATH, sheet_name="Feuil1", target_col_letter="ZX")
    except Exception as e:
        print(f"⚠️ ensure ROW_KEY Excel failed: {e}", flush=True)


    from datetime import date, datetime, timedelta
    import pandas as pd
    import re

    today_iso = date.today().strftime("%Y-%m-%d")
    now_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # ======================================================
    # 🔍 CHECK : Excel Dropbox a-t-il changé ?
    # ======================================================
    excel_dt = get_dropbox_file_last_modified()
    if not excel_dt:
        pass
    else:
        last_excel_dt = get_meta("excel_last_modified")
        if last_excel_dt:
            try:
                last_excel_dt = datetime.fromisoformat(last_excel_dt)
                if excel_dt <= last_excel_dt:
                    return 0
            except Exception:
                pass

    # ======================================================
    # 0️⃣ SÉCURITÉ DB : colonnes nécessaires
    # ======================================================
    with get_connection() as conn:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(planning)").fetchall()]

        if "IS_INDISPO" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "IS_INDISPO" INTEGER DEFAULT 0')
        if "DATE_ISO" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "DATE_ISO" TEXT')
        if "updated_at" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "updated_at" TEXT')
        if "IS_SUPERSEDED" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "IS_SUPERSEDED" INTEGER DEFAULT 0')
        if "EXCEL_UID" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "EXCEL_UID" TEXT')
        if "IS_BUREAU" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "IS_BUREAU" INTEGER DEFAULT 0')
        if "INDISPO_REASON" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "INDISPO_REASON" TEXT')
        if "LOCKED_BY_APP" not in cols:
            conn.execute('ALTER TABLE planning ADD COLUMN "LOCKED_BY_APP" INTEGER DEFAULT 0')


        conn.commit()

    # ======================================================
    # 1️⃣ Charger Excel Dropbox (Feuil1)
    # ======================================================
    df_excel = load_planning_from_dropbox("Feuil1")
    if df_excel is None or df_excel.empty:
        _ui_warn("Planning Dropbox vide.")
        return 0

    # ======================================================
    # 2️⃣ Couleurs Excel
    # ======================================================
    df_excel = add_excel_color_flags_from_dropbox(df_excel, "Feuil1")
    ensure_planning_color_columns()

    # ======================================================
    # 3️⃣ Normalisation DATE Excel (support FR + dd/mm + iso)
    # ======================================================
    if "DATE" not in df_excel.columns:
        _ui_error("❌ Colonne DATE absente.")
        return 0

    def _normalize_excel_date_to_iso(val):
        if val is None:
            return None

        if isinstance(val, (datetime, date)):
            try:
                return val.strftime("%Y-%m-%d")
            except Exception:
                pass

        try:
            if isinstance(val, (int, float)) and not pd.isna(val):
                if 20000 <= float(val) <= 60000:
                    dt = pd.to_datetime(float(val), unit="D", origin="1899-12-30", errors="coerce")
                    if not pd.isna(dt):
                        return dt.strftime("%Y-%m-%d")
        except Exception:
            pass

        s = str(val).strip()
        if not s or s.lower() in ("nan", "none"):
            return None

        s_low = s.lower().strip()

        try:
            if len(s_low) == 10 and s_low[4] == "-" and s_low[7] == "-":
                dt = pd.to_datetime(s_low, format="%Y-%m-%d", errors="coerce")
                if not pd.isna(dt):
                    return dt.strftime("%Y-%m-%d")
        except Exception:
            pass

        for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%d/%m/%y", "%d-%m-%y"):
            try:
                dt = datetime.strptime(s_low, fmt)
                if fmt.endswith("%y") and dt.year < 100:
                    dt = dt.replace(year=dt.year + (2000 if dt.year < 50 else 1900))
                return dt.strftime("%Y-%m-%d")
            except Exception:
                pass

        try:
            dt = pd.to_datetime(s_low, dayfirst=True, errors="coerce")
            if not pd.isna(dt):
                return dt.strftime("%Y-%m-%d")
        except Exception:
            pass

        s2 = re.sub(r"^(lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche)\s+", "", s_low).strip()

        months = {
            "janvier": "01", "février": "02", "fevrier": "02", "mars": "03", "avril": "04",
            "mai": "05", "juin": "06", "juillet": "07", "août": "08", "aout": "08",
            "septembre": "09", "octobre": "10", "novembre": "11", "décembre": "12", "decembre": "12",
        }

        m = re.match(r"^(\d{1,2})\s+([a-zéûôîàç]+)\s+(\d{4})$", s2)
        if m:
            jj = int(m.group(1))
            mois_txt = m.group(2)
            aa = int(m.group(3))
            mm = months.get(mois_txt)
            if mm:
                return f"{aa:04d}-{mm}-{jj:02d}"

        return None

    df_excel["DATE_ISO"] = df_excel["DATE"].apply(_normalize_excel_date_to_iso)
    df_excel = df_excel[df_excel["DATE_ISO"].notna()].copy()
    df_excel["DATE"] = pd.to_datetime(df_excel["DATE_ISO"], errors="coerce").dt.strftime("%d/%m/%Y")

    # ======================================================
    # 4️⃣ Normalisation HEURE + HEURE_FIN
    # ======================================================
    df_excel["HEURE"] = (
        df_excel.get("HEURE", "")
        .apply(normalize_time_string)
        .fillna("")
    )

    heure_fin_col = None
    for cand in ["HEURE_FIN", "HEURE FIN", "HEURE2", "HEURE 2", "²²²²"]:
        if cand in df_excel.columns:
            heure_fin_col = cand
            break

    if heure_fin_col:
        df_excel["_HEURE_FIN"] = (
            df_excel[heure_fin_col]
            .apply(normalize_time_string)
            .fillna("")
        )
    else:
        df_excel["_HEURE_FIN"] = ""
    try:
        if "CH" in df_excel.columns:
            df_excel["CH"] = df_excel["CH"].astype(str).str.strip()
    except Exception:
        pass
    # ======================================================
    # 4️⃣ BIS — INTERPRÉTATION ROBUSTE DE ²²²² (STATUT vs HEURE_FIN)
    # ======================================================
    def _norm_txt(v):
        return str(v or "").strip()

    def _norm_up(v):
        return str(v or "").strip().upper()

    def _looks_like_time(v) -> str:
        """
        Retourne 'HH:MM' si v est une heure valide, sinon ''.
        Les codes MA / VA / CP / OFF / etc. ne sont JAMAIS des heures.
        """
        s = _norm_txt(v)
        if not s:
            return ""
        if s.isalpha():
            return ""
        try:
            return normalize_time_string(s) or ""
        except Exception:
            return ""

    # Colonnes
    col_2222 = "²²²²" if "²²²²" in df_excel.columns else None

    # Colonnes de travail (toujours créées)
    df_excel["_HEURE_FIN"] = ""
    df_excel["_STATUT"] = ""

    if col_2222:
        df_excel["_HEURE_FIN"] = df_excel[col_2222].apply(_looks_like_time).fillna("")
        df_excel["_STATUT"] = df_excel[col_2222].apply(
            lambda v: _norm_up(v) if _looks_like_time(v) == "" else ""
        ).fillna("")

    # Sécuriser HEURE (début)
    df_excel["HEURE"] = (
        df_excel.get("HEURE", "")
        .apply(normalize_time_string)
        .fillna("")
    )

    # ======================================================
    # 5️⃣ DÉTECTION CONGÉ / INDISPO — LOGIQUE EXCEL RÉELLE
    # ======================================================
    h1 = df_excel["HEURE"].fillna("").astype(str)
    h2 = df_excel["_HEURE_FIN"].fillna("").astype(str)

    immat = (
        df_excel.get("IMMAT", "")
        .fillna("")
        .astype(str)
        .str.strip()
        .str.upper()
    )

    sens = (
        df_excel.get("Unnamed: 8", "")
        .fillna("")
        .astype(str)
        .str.strip()
        .str.upper()
    )

    statut = (
        df_excel["_STATUT"]
        .fillna("")
        .astype(str)
        .str.strip()
        .str.upper()
    )

    # 1) Congé jour complet : 00:00 → 00:00
    is_conge_0000 = (h1 == "00:00") & (h2 == "00:00")

    # 2) Congé via IMMAT : chiffre seul OU code (MA / VA / CP / CO)
    is_conge_immat = (
        immat.str.fullmatch(r"\d{1,2}", na=False)
        | immat.isin(["MA", "VA", "CP", "CO"])
    )

    # 3) Congé / motif via STATUT (²²²² texte)
    conge_codes = {
        "MA", "VA", "CP", "CO",
        "OFF", "RECUP", "CONGE", "VAC",
        "MAL", "MALADE",
    }
    is_conge_code = statut.isin(conge_codes)

    # 4) Indispo plage horaire
    is_indispo_plage = (
        (h1 != "") & (h2 != "") & (h1 != h2) & (~is_conge_0000)
    )

    # 5) Tâche (présence texte mais pas de véhicule)
    is_task = (sens != "") & (immat == "") & (~is_indispo_plage)

    # 6) Bureau explicite
    is_bureau = sens.str.contains(
        r"\bBUREAU\b",
        case=False,
        regex=True,
        na=False,
    )

    # Décision finale indispo
    df_excel["IS_INDISPO"] = (
        is_conge_0000
        | is_indispo_plage
        | is_conge_code
        | (is_conge_immat & (h1 == "00:00"))
    ).astype(int)

    # Bureau / tâche = jamais indispo
    df_excel["IS_BUREAU"] = is_bureau.astype(int)
    df_excel.loc[df_excel["IS_BUREAU"] == 1, "IS_INDISPO"] = 0
    df_excel.loc[is_task, "IS_INDISPO"] = 0

    # Raison (debug / affichage)
    df_excel["INDISPO_REASON"] = ""
    df_excel.loc[is_indispo_plage, "INDISPO_REASON"] = "INDISPO_PLAGE"
    df_excel.loc[statut == "MA", "INDISPO_REASON"] = "MALADE"
    df_excel.loc[
        is_conge_0000 | (is_conge_immat & (h1 == "00:00")),
        "INDISPO_REASON",
    ] = "CONGE"
    df_excel.loc[
        (is_conge_code) & (statut != "MA"),
        "INDISPO_REASON",
    ] = statut

    # ======================================================
    # 6️⃣ CONFIRMATION / CAISSE DEPUIS EXCEL (INCHANGÉ)
    # ======================================================
    if "CONFIRMED" not in df_excel.columns:
        df_excel["CONFIRMED"] = 0
    if "CONFIRMED_AT" not in df_excel.columns:
        df_excel["CONFIRMED_AT"] = None
    if "CAISSE_PAYEE" not in df_excel.columns:
        df_excel["CAISSE_PAYEE"] = 0

    if "CH_COLOR" in df_excel.columns:
        df_excel["CONFIRMED"] = df_excel["CH_COLOR"].apply(
            lambda c: 1 if _is_excel_cell_green(c) else 0
        )
        df_excel["CONFIRMED_AT"] = df_excel["CONFIRMED"].apply(
            lambda v: now_iso if v == 1 else None
        )

    if "CAISSE_COLOR" in df_excel.columns:

        def _calc_caisse_payee(row):
            paiement = str(row.get("PAIEMENT", "")).lower().strip()
            montant = row.get("Caisse")
            if paiement != "caisse":
                return 0
            try:
                montant = float(montant)
            except Exception:
                return 0
            if montant <= 0:
                return 0
            return 1 if _is_excel_cell_green(row.get("CAISSE_COLOR")) else 0

        df_excel["CAISSE_PAYEE"] = df_excel.apply(_calc_caisse_payee, axis=1)

    # ======================================================
    # 7️⃣ FILTRE “DEPUIS AUJOURD’HUI”
    # ======================================================
    df_excel = df_excel[df_excel["DATE_ISO"] >= today_iso].copy()
    if df_excel.empty:
        _ui_info("Aucune donnée à synchroniser.")
        return 0

    # ======================================================
    # 7️⃣ BIS — EXCEL_UID (CLÉ STABLE)
    # ======================================================
    def _norm_txt_uid(v):
        return str(v or "").strip().lower()

    def _make_excel_uid(row):
        num_bdc = _norm_txt_uid(row.get("Num BDC") or row.get("NUM BDC") or row.get("BDC"))
        vol = _norm_txt_uid(row.get("N° Vol") or row.get("N°Vol") or row.get("N Vol") or row.get("VOL"))
        nom = _norm_txt_uid(row.get("NOM"))
        adresse = _norm_txt_uid(row.get("ADRESSE"))
        cp = _norm_txt_uid(row.get("CP"))
        loc = _norm_txt_uid(row.get("Localité") or row.get("LOCALITE"))
        designation = _norm_txt_uid(row.get("DESIGNATION") or row.get("DESTINATION"))
        sens_uid = _norm_txt_uid(row.get("Unnamed: 8"))

        if num_bdc:
            return f"BDC|{num_bdc}|{nom}"
        if vol:
            return f"VOL|{vol}|{nom}"

        return "|".join(
            [
                "FALLBACK",
                nom,
                adresse,
                cp,
                loc,
                designation,
                sens_uid,
            ]
        )

    df_excel["EXCEL_UID"] = df_excel.apply(_make_excel_uid, axis=1)

    # ======================================================
    # 8️⃣ ROW_KEY UNIQUE (ANTI-DÉDUP CONGÉS)
    # ======================================================
    def _make_row_key_safe(row):
        base = make_row_key_from_row(row.to_dict())

        if int(row.get("IS_INDISPO", 0) or 0) == 1:
            ch = str(row.get("CH", "") or "").strip().upper()
            date_iso = str(row.get("DATE_ISO", "") or "")
            hh1 = str(row.get("HEURE", "") or "")
            hh2 = str(row.get("_HEURE_FIN", "") or "")
            imm = str(row.get("IMMAT", "") or "")
            reason = str(row.get("INDISPO_REASON", "") or "")
            return f"INDISPO|{date_iso}|{ch}|{hh1}|{hh2}|{imm}|{reason}"

        return base

    df_excel["row_key"] = df_excel.apply(_make_row_key_safe, axis=1)
    df_excel = df_excel.drop_duplicates(subset=["row_key"]).copy()

    # ======================================================
    # 9️⃣ RESET COMPLET DES NAVETTES FUTURES
    # 👉 Excel est SOURCE DE VÉRITÉ ABSOLUE
    # 👉 même les confirmées sont remplacées si modifiées
    # ======================================================

    # ✅ Sécurise DATE_ISO (sinon le reset ne touche rien si DB ancienne)
    try:
        from database import ensure_date_iso_populated
        ensure_date_iso_populated()
    except Exception:
        pass

    # ======================================================
    # 9️⃣ RESET COMPLET DES NAVETTES À PARTIR D’AUJOURD’HUI
    # 👉 Excel = SOURCE DE VÉRITÉ
    # 👉 filtre date robuste (DATE_ISO ou DATE dd/mm)
    # ======================================================
    with get_connection() as conn:
        conn.execute(
            '''
            DELETE FROM planning
            WHERE date(
                CASE
                    WHEN COALESCE(DATE_ISO,'') != '' THEN DATE_ISO
                    WHEN LENGTH(DATE) = 10 AND substr(DATE,3,1)='/' THEN
                        substr(DATE,7,4)||'-'||substr(DATE,4,2)||'-'||substr(DATE,1,2)
                    ELSE DATE
                END
            ) >= date(?)
              AND (LOCKED_BY_APP IS NULL OR LOCKED_BY_APP=0)
            ''',
            (today_iso,),
        )
        conn.commit()

    # ======================================================
    # 🔟 INSERTION SAFE (PRÉSERVE CONFIRMATION / ACK si même row_key)
    # ✅ UNE SEULE CONNEXION SQLite
    # ======================================================
    inserts = 0
    planning_cols = get_planning_table_columns()

    EXCEL_TO_DB_COLS = {
        "N° Vol": "N° Vol",
        "NUM BDC": "Num BDC",
        "Num BDC": "Num BDC",
        "BDC": "Num BDC",
        "Paiement": "PAIEMENT",
        "Caisse": "Caisse",
        "GO": "GO",
        "Reh": "Reh",
        "REH": "Reh",
        "Siège": "Siège",
    }

    with get_connection() as conn:
        cur = conn.cursor()

        for _, row in df_excel.iterrows():
            rk = row.get("row_key")
            if not rk:
                continue

            data: dict = {}

            if excel_sync_ts and "EXCEL_SYNC_TS" in planning_cols:
                data["EXCEL_SYNC_TS"] = excel_sync_ts

            for col in df_excel.columns:
                if col in planning_cols and col not in ("id",):
                    val = row.get(col)
                    if val not in (None, "", "nan"):
                        data[col] = sqlite_safe(val)

            for excel_col, db_col in EXCEL_TO_DB_COLS.items():
                if excel_col in df_excel.columns and db_col in planning_cols:
                    val = row.get(excel_col)
                    if val not in (None, "", "nan"):
                        data[db_col] = sqlite_safe(val)

            # ✅ EXCEL_UID + superseded
            if "EXCEL_UID" in planning_cols:
                data["EXCEL_UID"] = sqlite_safe(row.get("EXCEL_UID"))
            data["IS_SUPERSEDED"] = 0

            # Sécurité congé / indispo
            if int(row.get("IS_INDISPO", 0) or 0) == 1:
                data["CONFIRMED"] = 0
                data["CONFIRMED_AT"] = None
                data["ACK_AT"] = None
                data["ACK_TEXT"] = None
                data["CAISSE_PAYEE"] = 0
                data["IS_INDISPO"] = 1
                data["IS_SUPERSEDED"] = 0

            # Préserver l'état métier si même row_key existe
            cur.execute(
                """
                SELECT CONFIRMED, CONFIRMED_AT, ACK_AT, ACK_TEXT
                FROM planning
                WHERE row_key = ?
                """,
                (rk,),
            )
            prev = cur.fetchone()
            if prev:
                data["CONFIRMED"] = prev[0]
                data["CONFIRMED_AT"] = prev[1]
                data["ACK_AT"] = prev[2]
                data["ACK_TEXT"] = prev[3]

            data["row_key"] = rk
            data["updated_at"] = now_iso

            cols_ins = [c for c in data.keys() if c in planning_cols]
            if not cols_ins:
                continue

            col_sql = ", ".join([f'"{c}"' for c in cols_ins])
            placeholders = ", ".join(["?"] * len(cols_ins))
            values = [data[c] for c in cols_ins]

            try:
                # ✅ UPSERT SAFE : met à jour si row_key existe, sinon insère
                cur.execute("SELECT id FROM planning WHERE row_key = ?", (rk,))
                existing = cur.fetchone()

                if existing:
                    # UPDATE (sans toucher row_key)
                    cols_up = [c for c in cols_ins if c != "row_key"]
                    if cols_up:
                        set_sql = ", ".join([f'"{c}" = ?' for c in cols_up])
                        vals_up = [data[c] for c in cols_up] + [rk]
                        cur.execute(
                            f"UPDATE planning SET {set_sql} WHERE row_key = ?",
                            vals_up,
                        )
                else:
                    cur.execute(
                        f"INSERT INTO planning ({col_sql}) VALUES ({placeholders})",
                        values,
                    )
                    if cur.rowcount and cur.rowcount > 0:
                        inserts += 1
            except Exception:
                pass

        conn.commit()

    # ======================================================
    # 11️⃣ Rebuild vues
    # ======================================================
    rebuild_planning_views()

    # ======================================================
    # 12️⃣ Feuil2 → chauffeurs
    # ======================================================
    df_ch = load_planning_from_dropbox("Feuil2")
    if df_ch is not None and not df_ch.empty:
        with get_connection() as conn:
            conn.execute("DROP TABLE IF EXISTS chauffeurs")
            conn.commit()

        cols = [c for c in df_ch.columns if c]
        col_defs = ", ".join(f'"{c}" TEXT' for c in cols)
        cols_sql = ",".join(f'"{c}"' for c in cols)
        placeholders = ",".join("?" for _ in cols)

        with get_connection() as conn:
            conn.execute(f'CREATE TABLE chauffeurs ({col_defs})')
            for _, r in df_ch.iterrows():
                conn.execute(
                    f'INSERT INTO chauffeurs ({cols_sql}) VALUES ({placeholders})',
                    [sqlite_safe(r.get(c)) for c in cols],
                )
            conn.commit()

    # ======================================================
    # 13️⃣ Feuil3
    # ======================================================
    df_f3 = load_planning_from_dropbox("Feuil3")
    if df_f3 is not None and not df_f3.empty:
        with get_connection() as conn:
            conn.execute("DROP TABLE IF EXISTS feuil3")
            conn.commit()

        cols3 = [c for c in df_f3.columns if c]
        col_defs3 = ", ".join(f'"{c}" TEXT' for c in cols3)
        cols_sql3 = ",".join(f'"{c}"' for c in cols3)
        placeholders3 = ",".join("?" for _ in cols3)

        with get_connection() as conn:
            conn.execute(f'CREATE TABLE feuil3 ({col_defs3})')
            for _, r in df_f3.iterrows():
                conn.execute(
                    f'INSERT INTO feuil3 ({cols_sql3}) VALUES ({placeholders3})',
                    [sqlite_safe(r.get(c)) for c in cols3],
                )
            conn.commit()

    # ======================================================
    # 14️⃣ Cache / UI (silencieux + ciblé)
    # ======================================================
    if ui:
        st.session_state["last_sync_time"] = datetime.now().strftime("%H:%M")
    # Clear ciblé si possible (évite de tout casser)
    try:
        get_planning.clear()
    except Exception:
        pass

    if excel_dt:
        set_meta("excel_last_modified", excel_dt.isoformat())

    return inserts








def sync_planning_from_uploaded_file(uploaded_file):
    """
    Synchronisation DB depuis un fichier Excel uploadé manuellement
    (mode secours si Dropbox indisponible)
    """
    try:
        # 🔹 Lire le fichier uploadé en mémoire
        content = uploaded_file.getbuffer()

        # 🔹 Monkey-patch temporaire : on remplace le downloader Dropbox
        def _mock_download_dropbox_excel_bytes(path=None):
            return content

        # Sauvegarde de la fonction originale
        original_download = download_dropbox_excel_bytes

        # Remplacement temporaire
        globals()["download_dropbox_excel_bytes"] = _mock_download_dropbox_excel_bytes

        # 🔁 Réutilise EXACTEMENT la même logique que Dropbox
        inserted = sync_planning_from_today()

        # 🔙 Restauration fonction originale
        globals()["download_dropbox_excel_bytes"] = original_download

        return inserted

    except Exception as e:
        st.error(f"❌ Erreur synchronisation fichier manuel : {e}")
        return 0

def rebuild_planning_db_from_dropbox_full() -> int:
    """
    🔥 Reconstruction COMPLÈTE de la DB planning depuis Dropbox
    (2025 + 2026) — SANS DOUBLONS
    """

    from datetime import datetime
    import pandas as pd

    # ======================================================
    # 1️⃣ Charger Excel Dropbox (Feuil1)
    # ======================================================
    df_excel = load_planning_from_dropbox("Feuil1")
    if df_excel.empty:
        return 0

    # ======================================================
    # 2️⃣ Couleurs Excel
    # ======================================================
    df_excel = add_excel_color_flags_from_dropbox(df_excel, "Feuil1")
    ensure_planning_color_columns()

    # ======================================================
    # 3️⃣ Normalisation DATE
    # ======================================================
    df_excel["DATE"] = pd.to_datetime(
        df_excel["DATE"],
        dayfirst=True,
        errors="coerce",
    ).dt.date

    df_excel = df_excel[df_excel["DATE"].notna()].copy()

    # ======================================================
    # 4️⃣ Normalisation HEURE
    # ======================================================
    df_excel["HEURE"] = df_excel["HEURE"].apply(normalize_time_string)

    # ======================================================
    # 5️⃣ row_key + suppression doublons Excel
    # ======================================================
    df_excel["row_key"] = df_excel.apply(
        lambda r: make_row_key_from_row(r.to_dict()),
        axis=1,
    )

    df_excel = df_excel.drop_duplicates(subset=["row_key"])

    # ======================================================
    # 6️⃣ PURGE TOTALE DE LA TABLE planning
    # ======================================================
    with get_connection() as conn:
        conn.execute("DELETE FROM planning")
        conn.commit()

    # ======================================================
    # 7️⃣ INSERTION PROPRE (OR IGNORE)
    # ======================================================
    inserts = 0

    for _, row in df_excel.iterrows():
        data = {
            col: sqlite_safe(row.get(col))
            for col in df_excel.columns
            if col not in ("id",)
        }

        # Format DATE attendu par ta table
        data["DATE"] = row["DATE"].strftime("%Y-%m-%d")
        data["HEURE"] = row["HEURE"]
        data["row_key"] = row["row_key"]
        data["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        try:
            insert_planning_row(
                data,
                ignore_conflict=True,  # 🔒 sécurité ultime
            )
            inserts += 1
        except Exception:
            pass

    # ======================================================
    # 8️⃣ Recréer les vues SQL
    # ======================================================
    rebuild_planning_views()

    # ======================================================
    # 9️⃣ Import Feuil2 → chauffeurs
    # ======================================================
    df_ch = load_planning_from_dropbox("Feuil2")
    if not df_ch.empty:
        with get_connection() as conn:
            conn.execute("DROP TABLE IF EXISTS chauffeurs")
            conn.commit()

        cols = [c for c in df_ch.columns if c]
        col_defs = ", ".join(f'"{c}" TEXT' for c in cols)
        cols_sql = ",".join(f'"{c}"' for c in cols)
        placeholders = ",".join("?" for _ in cols)

        with get_connection() as conn:
            conn.execute(f'CREATE TABLE chauffeurs ({col_defs})')
            conn.commit()

        for _, r in df_ch.iterrows():
            values = [sqlite_safe(r.get(c)) for c in cols]
            with get_connection() as conn:
                conn.execute(
                    f'INSERT INTO chauffeurs ({cols_sql}) VALUES ({placeholders})',
                    values,
                )
                conn.commit()

    # ======================================================
    # 🔟 Import Feuil3 → feuil3
    # ======================================================
    df_f3 = load_planning_from_dropbox("Feuil3")
    if not df_f3.empty:
        with get_connection() as conn:
            conn.execute("DROP TABLE IF EXISTS feuil3")
            conn.commit()

        cols3 = [c for c in df_f3.columns if c]
        col_defs3 = ", ".join(f'"{c}" TEXT' for c in cols3)
        cols_sql3 = ",".join(f'"{c}"' for c in cols3)
        placeholders3 = ",".join("?" for _ in cols3)

        with get_connection() as conn:
            conn.execute(f'CREATE TABLE feuil3 ({col_defs3})')
            conn.commit()

        for _, r in df_f3.iterrows():
            values = [sqlite_safe(r.get(c)) for c in cols3]
            with get_connection() as conn:
                conn.execute(
                    f'INSERT INTO feuil3 ({cols_sql3}) VALUES ({placeholders3})',
                    values,
                )
                conn.commit()

    return inserts



from database import make_row_key_from_row, get_latest_ch_overrides_map

def apply_actions_overrides(df: pd.DataFrame) -> pd.DataFrame:
    """
    Applique les overrides chauffeur (CH) sur le DataFrame.
    ⚡ Optimisée :
    - PAS de recalcul de row_key
    - DB lue UNE SEULE FOIS
    - Logique métier identique
    """

    if df is None or df.empty:
        return df

    # ⛔ Ne jamais recalculer row_key en UI
    if "row_key" not in df.columns:
        return df

    df = df.copy()

    # ==================================================
    # 🔑 Charger les overrides UNE SEULE FOIS
    # ==================================================
    row_keys = (
        df["row_key"]
        .dropna()
        .astype(str)
        .unique()
        .tolist()
    )

    if not row_keys:
        df["_needs_excel_update"] = 0
        return df

    mp = get_latest_ch_overrides_map(row_keys)

    if not mp:
        df["_needs_excel_update"] = 0
        return df

    # ==================================================
    # ⚡ Application rapide des overrides
    # ==================================================
    df["_CH_ORIG"] = df.get("CH", "")

    ch_series = df["CH"] if "CH" in df.columns else pd.Series("", index=df.index)

    df["CH"] = (
        df["row_key"]
        .map(mp)
        .combine_first(ch_series)
    )

    df["_needs_excel_update"] = (
        df["row_key"]
        .isin(mp.keys())
        .astype(int)
    )

    return df


import requests
def flight_badge(status: str, delay_min: int = 0) -> str:
    status = (status or "").upper()
    delay_min = int(delay_min or 0)

    if status == "ON_TIME":
        return "🟢 À l’heure"
    if status == "DELAYED":
        if delay_min >= 30:
            return f"🔴 Retard {delay_min} min"
        return f"🟠 Retard {delay_min} min"
    if status == "CANCELLED":
        return "🔴 Annulé"
    if status == "LANDED":
        return "✅ Atterri"
    return "⚪ Statut inconnu"
def extract_vol_val(row, columns):
    """
    Extrait le numéro de vol depuis une ligne,
    robuste aux variantes de nom de colonne.
    """
    for col in ["N° Vol", "N° Vol ", "Num Vol", "VOL", "Vol"]:
        if col in columns:
            v = str(row.get(col, "") or "").strip()
            if v:
                return v
    return ""
AVIATIONSTACK_KEY = "e5cb6733f9d69693e880c982795ba27d"
import requests
import streamlit as st

@st.cache_data(ttl=600, show_spinner=False)
def get_flight_status_cached(flight_number: str):
    """
    Retourne TOUJOURS un tuple :
    (status, delay_min, sched_dt, est_dt)
    sched_dt / est_dt = datetime pandas (ou None)
    """
    if not flight_number:
        return "", 0, None, None

    try:
        r = requests.get(
            "http://api.aviationstack.com/v1/flights",
            params={"access_key": AVIATIONSTACK_KEY, "flight_iata": flight_number},
            timeout=5,
        )
        data = r.json()

        if not data.get("data"):
            return "", 0, None, None

        f = data["data"][0]
        status_raw = (f.get("flight_status") or "").lower()

        # mapping statut
        if status_raw in ("scheduled", "active"):
            status = "ON_TIME"
        elif status_raw == "delayed":
            status = "DELAYED"
        elif status_raw == "cancelled":
            status = "CANCELLED"
        elif status_raw == "landed":
            status = "LANDED"
        else:
            status = ""

        # ⚠️ on prend ici ARRIVAL (arrivée) : scheduled / estimated
        sched = f.get("arrival", {}).get("scheduled")
        est = f.get("arrival", {}).get("estimated")

        sched_dt = pd.to_datetime(sched) if sched else None
        est_dt = pd.to_datetime(est) if est else None

        delay_min = 0
        if sched_dt is not None and est_dt is not None:
            delay_min = int((est_dt - sched_dt).total_seconds() / 60)

        return status, delay_min, sched_dt, est_dt

    except Exception:
        return "", 0, None, None

def _parse_ddmmyyyy_date(val):
    try:
        s = str(val or "").strip()
        if not s:
            return None
        # support YYYY-MM-DD too
        if re.match(r"^\d{4}-\d{2}-\d{2}$", s):
            return datetime.strptime(s, "%Y-%m-%d").date()
        return datetime.strptime(s, "%d/%m/%Y").date()
    except Exception:
        return None

def _parse_hhmm_time(val):
    try:
        t = normalize_time_string(val)
        if not t:
            return None
        return datetime.strptime(t, "%H:%M").time()
    except Exception:
        return None

def should_query_flight_status(row) -> bool:
    """Limite les requêtes vol pour éviter de brûler l'API.
    Règle : on vérifie surtout à partir de J-1 (24h avant) et jusqu'à quelques heures après.
    """
    try:
        d = _parse_ddmmyyyy_date(row.get("DATE"))
        if not d:
            return False

        # Time de vol prioritaire : Décollage (si présent), sinon H South, sinon heure navette
        t = _parse_hhmm_time(row.get("Décollage")) or _parse_hhmm_time(row.get("H South")) or _parse_hhmm_time(row.get("HEURE"))
        if t:
            planned = datetime.combine(d, t)
        else:
            # pas d'heure → on fait simple : seulement si aujourd'hui / demain
            planned = datetime.combine(d, datetime.min.time())

        now = datetime.now()

        # fenêtre : de 24h avant à +6h après
        delta_h = (planned - now).total_seconds() / 3600.0

        if -6 <= delta_h <= 24:
            return True

        # fallback : si DATE = aujourd'hui ou demain, on autorise
        if d in {now.date(), (now.date() + timedelta(days=1))}:
            return True

        return False
    except Exception:
        return False


# ============================================================
#   MAPPING ABRÉVIATIONS CLIENTS / SITES
# ============================================================

CLIENT_ALIASES = {
    "KI HQ": {
        "name": "Knauf Insulation",
        "site": "Headquarters",
        "city": "Visé",
    },
    "JCO": {
        "name": "John Cockerill",
        "site": "Site industriel",
        "city": "Seraing",
    },
    "JCC": {
        "name": "John Cockerill",
        "site": "Site château",
        "city": "Seraing",
    },
}


# ==========================
#  KM / TEMPS (OpenRouteService)
# ==========================
ORS_API_KEY = "5b3ce3597851110001cf62480ac03479d6074e1ebda549044ad14608"

AIRPORT_ALIASES = {
    "CRL": "Brussels South Charleroi Airport, Belgium",
    "CHARLEROI": "Brussels South Charleroi Airport, Belgium",
    "BRU": "Brussels Airport, Zaventem, Belgium",
    "BRUXELLES": "Brussels Airport, Zaventem, Belgium",
    "ZAVENTEM": "Brussels Airport, Zaventem, Belgium",
    "LUX": "Luxembourg Airport, Luxembourg",
    "LUXEMBOURG": "Luxembourg Airport, Luxembourg",
}

def _pick_first(row, candidates):
    for c in candidates:
        if c in row.index:
            v = str(row.get(c, "") or "").strip()
            if v and v.lower() != "nan":
                return v
    return ""

def _clean_address_piece(val) -> str:
    if val is None:
        return ""
    s = str(val).strip()
    if not s:
        return ""
    # évite les 4000.0 / 12.0 venant d'Excel
    if s.endswith('.0'):
        try:
            f = float(s)
            if f.is_integer():
                s = str(int(f))
        except Exception:
            pass
    return re.sub(r"\s+", " ", s).strip(" ,;")


def build_full_address_from_row(row: pd.Series) -> str:
    # Essaye de reconstruire "Adresse + CP + Ville"
    adr = _clean_address_piece(_pick_first(row, ["ADRESSE", "Adresse", "ADRESSE RDV", "Adresse RDV", "RUE", "Rue"]))
    cp  = _clean_address_piece(_pick_first(row, ["CP", "Code postal", "CODE POSTAL", "Postal", "ZIP"]))
    vil = _clean_address_piece(_pick_first(row, ["Localité", "LOCALITE", "Ville", "VILLE", "COMMUNE"]))
    parts = [p for p in [adr, cp, vil] if p]
    return " ".join(parts).strip()


def build_navigation_address_from_row(row: pd.Series) -> str:
    """
    Construit une adresse GPS plus fiable que l'adresse d'affichage.
    Format privilégié : "Rue numéro, CP Ville, Belgique".
    Fallback sur la destination si l'adresse est absente.
    """
    adr = _clean_address_piece(_pick_first(row, ["ADRESSE", "Adresse", "ADRESSE RDV", "Adresse RDV", "RUE", "Rue"]))
    cp = _clean_address_piece(_pick_first(row, ["CP", "Code postal", "CODE POSTAL", "Postal", "ZIP"]))
    vil = _clean_address_piece(_pick_first(row, ["Localité", "LOCALITE", "Ville", "VILLE", "COMMUNE"]))

    street_line = adr
    city_line = " ".join([p for p in [cp, vil] if p]).strip()

    parts = [p for p in [street_line, city_line] if p]
    if parts:
        country = "Belgique"
        city_up = vil.upper()
        if city_up in {"LUXEMBOURG", "LUXEMBURG", "LUX AIRPORT"}:
            country = "Luxembourg"
        return ", ".join(parts + [country])

    dest = resolve_destination_text(row)
    dest = resolve_client_alias(dest)
    dest_up = str(dest or "").upper()
    airport_map = {
        "BRU": "Brussels Airport, Zaventem, Belgique",
        "ZAV": "Brussels Airport, Zaventem, Belgique",
        "CRL": "Brussels South Charleroi Airport, Charleroi, Belgique",
        "LUX": "Luxembourg Airport, Luxembourg",
        "LGG": "Liège Airport, Grâce-Hollogne, Belgique",
        "DUS": "Düsseldorf Airport, Düsseldorf, Allemagne",
        "MIDI": "Gare de Bruxelles-Midi, Bruxelles, Belgique",
    }
    for code, full_addr in airport_map.items():
        if code in dest_up:
            return full_addr

    return dest or build_full_address_from_row(row)


def resolve_destination_text(row: pd.Series) -> str:
    # Colonne destination/route dans ton fichier : tu utilises déjà "DE/VERS" et parfois "Unnamed: 8"
    dest = _pick_first(row, ["DE/VERS", "DESTINATION", "Destination", "Unnamed: 8", "ROUTE"])
    if not dest:
        return ""
    key = dest.strip().upper()
    for k, full in AIRPORT_ALIASES.items():
        if k in key:
            return full
    return dest

@st.cache_data
def ors_route_km_min(origin_text: str, dest_text: str):
    """
    Retourne (km, minutes) via ORS directions.
    Cache 24h pour éviter de brûler la clé.
    """
    if not ORS_API_KEY:
        return None, None
    if not origin_text or not dest_text:
        return None, None

    # ORS: on passe par géocodage Nominatim-like ? => ORS a aussi /geocode/search.
    # Pour rester simple et robuste: ORS Geocode puis Directions.
    try:
        # 1) Geocode origin
        r1 = requests.get(
            "https://api.openrouteservice.org/geocode/search",
            params={"api_key": ORS_API_KEY, "text": origin_text},
            timeout=8
        ).json()
        if not r1.get("features"):
            return None, None
        o_lon, o_lat = r1["features"][0]["geometry"]["coordinates"]

        # 2) Geocode dest
        r2 = requests.get(
            "https://api.openrouteservice.org/geocode/search",
            params={"api_key": ORS_API_KEY, "text": dest_text},
            timeout=8
        ).json()
        if not r2.get("features"):
            return None, None
        d_lon, d_lat = r2["features"][0]["geometry"]["coordinates"]

        # 3) Directions driving-car
        r3 = requests.post(
            "https://api.openrouteservice.org/v2/directions/driving-car",
            headers={"Authorization": ORS_API_KEY, "Content-Type": "application/json"},
            json={"coordinates": [[o_lon, o_lat], [d_lon, d_lat]]},
            timeout=10
        ).json()

        feat = (r3.get("features") or [None])[0]
        if not feat:
            return None, None

        seg = feat["properties"]["segments"][0]
        dist_m = float(seg.get("distance", 0.0))
        dur_s  = float(seg.get("duration", 0.0))

        km = round(dist_m / 1000.0, 1)
        minutes = int(round(dur_s / 60.0))
        return km, minutes
    except Exception:
        return None, None
# ============================================================
#   CONFIG STREAMLIT
# ============================================================

st.set_page_config(
    page_title="Airports-Lines – Planning chauffeurs",
    layout="wide",
)


def get_chauffeurs_for_ui() -> List[str]:
    """
    Liste des codes CH pour les listes déroulantes :

    - on part des chauffeurs officiels (Feuil2 → get_chauffeurs())
    - on ajoute tous les codes distincts trouvés dans la colonne CH
      du planning (FA*, FANP, FADO, …)
    - on ne modifie rien dans la DB, ni dans le XLSX
    - les codes sont affichés exactement comme dans le planning
      (on enlève juste les espaces autour)
    """
    # Base : chauffeurs officiels (Feuil2 / table chauffeurs)
    try:
        base = get_chauffeurs()  # ex: FA, FA1, DO, NP, ...
    except Exception:
        base = []

    # Valeurs réelles présentes dans la colonne CH du planning
    extra: List[str] = []
    try:
        df_all = get_planning(
            start_date=None,
            end_date=None,
            chauffeur=None,
            type_filter=None,
            search=None,
            max_rows=None,  # pas de limite
        )
        if not df_all.empty and "CH" in df_all.columns:
            extra = (
                df_all["CH"]
                .astype(str)
                .map(lambda x: x.strip() if x is not None else "")
                .replace("", pd.NA)
                .dropna()
                .unique()
                .tolist()
            )
    except Exception:
        df_all = None

    # Union des deux listes, sans doublons, sans changer la casse
    all_codes = []
    seen = set()
    for code in (base + extra):
        if code is None:
            continue
        c = str(code).strip()
        if not c:
            continue
        if c not in seen:
            seen.add(c)
            all_codes.append(c)

    # Tri alphabétique simple
    all_codes = sorted(all_codes, key=lambda x: x.upper())
    return all_codes
from database import split_chauffeurs

def receive_chauffeur_planning(chauffeur: str, texte: str, canal="whatsapp"):
    with get_connection() as conn:
        conn.execute(
            """
            INSERT INTO chauffeur_messages (ts, chauffeur, canal, contenu)
            VALUES (?, ?, ?, ?)
            """,
            (
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                chauffeur,
                canal,
                texte,
            ),
        )
        conn.commit()

    # 🔔 notif admin
    st.session_state["admin_notif"] = {
        "type": "PLANNING",
        "chauffeur": chauffeur,
    }

def send_email_to_chauffeurs_from_row(row, subject: str, body: str):
    """
    Envoie un email à TOUS les chauffeurs réels concernés par la navette.
    Gère FA*DO, FADONP, FADO*NP*, etc.
    """

    ch_code = row.get("CH")
    if not ch_code:
        return [], []

    chauffeurs = split_chauffeurs(ch_code)

    emails_sent = []
    emails_missing = []

    for ch in dict.fromkeys(chauffeurs):  # anti-doublon
        email = get_chauffeur_email(ch)

        if not email:
            emails_missing.append(ch)
            continue

        send_mail(
            to=email,
            subject=subject,
            body=body,
        )

        emails_sent.append(ch)

    return emails_sent, emails_missing

# ===========================
#  CONFIG NOTIFICATIONS EMAIL
# ===========================

SMTP_HOST = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_USER = "airportslinesbureau@gmail.com"
FROM_EMAIL = SMTP_USER
ADMIN_NOTIFICATION_EMAIL = "airportslinesbureau@gmail.com"


def get_smtp_password():
        """
        Récupère le mot de passe SMTP de façon SAFE :
        - secrets.toml si présent
        - sinon variable d’environnement
        - sinon chaîne vide (ne plante jamais)
        """
        try:
                return st.secrets["SMTP_PASSWORD"]
        except Exception:
                return os.environ.get("SMTP_PASSWORD", "")


SMTP_PASSWORD = get_smtp_password()

# ============================================================
#   HELPERS — NORMALISATION DES HEURES
# ============================================================

def normalize_time_string(val):
    """
    Nettoie et convertit une heure vers HH:MM (FORMAT UNIQUE).
    Retourne "" si invalide / vide.
    """
    if val is None:
        return ""

    s = str(val).strip()
    if not s or s == "0":
        return ""

    # Remplacer H / h par :
    s = s.replace("H", ":").replace("h", ":").strip()

    # Cas datetime / pandas Timestamp
    try:
        if hasattr(val, "hour") and hasattr(val, "minute"):
            return f"{int(val.hour):02d}:{int(val.minute):02d}"
    except Exception:
        pass

    # Format HHMM → HH:MM
    if s.isdigit():
        try:
            if len(s) <= 2:
                h = int(s)
                m = 0
            else:
                h = int(s[:-2])
                m = int(s[-2:])
            if 0 <= h <= 23 and 0 <= m <= 59:
                return f"{h:02d}:{m:02d}"
            return ""
        except Exception:
            return ""

    # Format H:M, HH:M, H:MM, HH:MM, HH:MM:SS
    if ":" in s:
        try:
            parts = s.split(":")
            h = int(parts[0])
            m = int(parts[1])
            if 0 <= h <= 23 and 0 <= m <= 59:
                return f"{h:02d}:{m:02d}"
            return ""
        except Exception:
            return ""

    return ""

def format_sens_ar(val: str) -> str:
    """
    Normalise la colonne Unnamed: 8 SANS écraser un texte libre.

    - Si la cellule contient UNIQUEMENT un sens (DE / VERS) et éventuellement A/R → on normalise.
    - Si la cellule contient autre chose (ex: 'BUREAU', 'COURSE', 'RDV ...') → on renvoie le texte tel quel.
    """
    if val is None:
        return ""

    raw = str(val).strip()
    if not raw or raw.lower() == "nan":
        return ""

    txt = raw.upper().strip().replace("—", "-").replace("–", "-")

    # Si ça ressemble à un "sens" simple, on normalise
    # Exemples acceptés :
    #   DE
    #   VERS
    #   A/R
    #   DE A/R
    #   VERS A/R
    simple = txt.replace(" ", "")
    simple = simple.replace("-", "")
    if simple in {"DE", "VERS", "AR", "A/R", "DEA/R", "VERSA/R", "DEAR", "VERSAR"}:
        has_ar = "A/R" in txt or txt.replace("/", "").replace(" ", "") == "AR" or "AR" == simple[-2:]
        if "DE" in txt:
            sens = "DE"
        elif "VERS" in txt:
            sens = "VERS"
        else:
            sens = ""
        if has_ar and sens:
            return f"{sens} – A/R"
        if has_ar:
            return "A/R"
        return sens

    # Sinon : texte libre → on garde
    return raw


def build_printable_html_planning(df: pd.DataFrame, ch: str):
    html_table = df.to_html(index=False)

    return f"""
    <html>
    <head>
        <style>
            @page {{
                size: A4 landscape;
                margin: 10mm;
            }}
            body {{
                font-family: Arial, sans-serif;
                zoom: 0.85;
            }}
            h2 {{
                margin-bottom: 10px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                font-size: 10px;
            }}
            th, td {{
                border: 1px solid #ccc;
                padding: 4px;
                text-align: left;
            }}
            th {{
                background-color: #f0f0f0;
            }}
        </style>
    </head>
    <body onload="window.print();">
        <h2>Planning chauffeur — {ch}</h2>
        {html_table}
    </body>
    </html>
    """

def resolve_client_alias(text: str) -> str:
    """
    Remplace une abréviation connue par sa description complète.
    (Pour affichage : vue chauffeur, vue mobile, PDF, WhatsApp, etc.)
    """
    if not text:
        return ""

    raw = str(text).strip()
    key = raw.upper()

    info = CLIENT_ALIASES.get(key)
    if not info:
        return raw

    parts = [info.get("name", "").strip()]
    if info.get("site"):
        parts.append(str(info["site"]).strip())
    if info.get("city"):
        parts.append(str(info["city"]).strip())

    parts = [p for p in parts if p]
    return " – ".join(parts) if parts else raw

# ============================================================
#   HELPERS – BOOL FLAG
# ============================================================

def bool_from_flag(x) -> bool:
    """Convertit 1, TRUE, x, oui, Yes, etc. en bool."""
    if x is None:
        return False
    s = str(x).strip().lower()
    return s in ["1", "true", "x", "oui", "yes"]
# ============================================================
#   📊 HISTORIQUE DES ENVOIS — DB
# ============================================================

def ensure_send_log_table():
    with get_connection() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS send_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ts DATETIME DEFAULT CURRENT_TIMESTAMP,
                chauffeur TEXT,
                canal TEXT,
                periode TEXT,
                statut TEXT,
                message TEXT
            )
        """)
        conn.commit()


def log_send(chauffeur, canal, periode, statut, message):
    with get_connection() as conn:
        conn.execute(
            """
            INSERT INTO send_log (chauffeur, canal, periode, statut, message)
            VALUES (?, ?, ?, ?, ?)
            """,
            (chauffeur, canal, periode, statut, message),
        )
        conn.commit()



# ============================================================
#   DÉTECTION INDISPO CHAUFFEUR
#   (Feuil1 = NP 12:00 … etc.)
# ============================================================

def is_indispo_row(row, columns):
    """
    Détecte une indispo chauffeur.
    ⚠️ Les congés (00:00 + IMMAT = 1/2) NE SONT PAS des indispos.
    """

    heure = normalize_time_string(row.get("HEURE"))
    immat = str(row.get("IMMAT", "") or "").strip()

    # ✅ CONGÉ CHAUFFEUR → JAMAIS filtré
    if heure == "00:00" and immat.isdigit():
        return False

    # ❌ logique indispo existante (inchangée)
    if int(row.get("IS_INDISPO", 0) or 0) == 1:
        return True

    return False


def print_html_popup(html: str):
    # Ouvre une nouvelle fenêtre et lance l'impression
    html_safe = (html or "").replace("`", "\\`")
    popup = f"""
    <script>
      (function() {{
        var w = window.open("", "_blank");
        if (!w) {{
          alert("Popup bloquée par le navigateur. Autorise les popups puis réessaie.");
          return;
        }}
        w.document.open();
        w.document.write(`{html_safe}`);
        w.document.close();
        w.focus();
        setTimeout(function() {{
          w.print();
        }}, 400);
      }})();
    </script>
    """
    components.html(popup, height=0, width=0)

# ============================================================
#   HELPERS — PHONE / WHATSAPP / MAIL
# ============================================================

def clean_phone(phone: str) -> str:
    if phone is None:
        return ""
    return "".join(ch for ch in str(phone) if ch.isdigit())


def phone_to_whatsapp_number(phone: str) -> str:
    digits = clean_phone(phone)
    if not digits:
        return ""
    if digits.startswith("0"):
        return "32" + digits[1:]
    return digits


def build_whatsapp_link(phone: str, message: str) -> str:
    import urllib.parse
    num = phone_to_whatsapp_number(phone)
    if not num:
        return "#"
    text = urllib.parse.quote(str(message or ""))
    return f"https://wa.me/{num}?text={text}"


def build_waze_link(address: str) -> str:
    """Construit un lien Waze vers une adresse texte (plus fiable en webview)."""
    import urllib.parse
    import re

    addr = re.sub(r"\s+", " ", str(address or "")).strip(" ,;")
    if not addr:
        return "#"

    query = urllib.parse.quote(addr)
    return f"https://www.waze.com/ul?query={query}&navigate=yes"

def build_google_maps_link(address: str) -> str:
    import urllib.parse
    if not address:
        return "#"
    return (
        "https://www.google.com/maps/search/?api=1&query="
        + urllib.parse.quote(address)
    )


def build_mailto_link(to_email: str, subject: str, body: str) -> str:
    import urllib.parse
    if not to_email:
        return "#"
    return (
        "mailto:"
        + to_email
        + "?subject="
        + urllib.parse.quote(subject)
        + "&body="
        + urllib.parse.quote(body)
    )
def send_mail_admin(subject: str, body: str):
    """Envoie un mail texte simple à l'admin."""
    try:
        if not SMTP_PASSWORD:
            raise RuntimeError("SMTP_PASSWORD manquant (définis-le dans les variables d'environnement ou st.secrets)")
        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = SMTP_USER
        msg["To"] = ADMIN_NOTIFICATION_EMAIL

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as s:
            s.starttls()
            s.login(SMTP_USER, SMTP_PASSWORD)
            s.send_message(msg)
    except Exception as e:
        print("Erreur envoi mail:", e)
def build_planning_mail_body(
    df_ch: pd.DataFrame,
    ch: str,
    from_date: date,
    to_date: date | None,
):
    cols = df_ch.columns.tolist()
    lines: list[str] = []

    # =============================
    # EN-TÊTE
    # =============================
    periode = (
        from_date.strftime("%d/%m/%Y")
        if not to_date or from_date == to_date
        else f"{from_date.strftime('%d/%m/%Y')} → {to_date.strftime('%d/%m/%Y')}"
    )

    lines.append(f"🚖 Planning — Chauffeur : {ch}")
    lines.append(f"📆 Période : {periode}")
    lines.append("")

    # =============================
    # BOUCLE NAVETTES
    # =============================
    for _, row in df_ch.iterrows():

        # ===================================================
        # 🚖 NAVETTE — BLOC COMPLET (MAIL)
        # ===================================================

        # ------------------
        # Flags groupage / partage / attente
        # ------------------
        is_groupage = int(row.get("IS_GROUPAGE", 0) or 0) == 1
        is_partage = int(row.get("IS_PARTAGE", 0) or 0) == 1
        is_attente = int(row.get("IS_ATTENTE", 0) or 0) == 1

        prefix = ""
        if is_groupage:
            prefix += "[GROUPÉE] "
        elif is_partage:
            prefix += "[PARTAGÉE] "
        if is_attente:
            prefix += "⭐ "

        # ------------------
        # Chauffeur
        # ------------------
        ch_code = str(row.get("CH", "") or ch).strip()
        lines.append(f"👨‍✈️ {ch_code}")

        # ------------------
        # Confirmation
        # ------------------
        if is_navette_confirmed(row):
            lines.append("✅ Navette confirmée")
        else:
            lines.append("🕒 À confirmer")

        # ------------------
        # Date / Heure
        # ------------------
        dv = row.get("DATE")
        if isinstance(dv, date):
            date_txt = dv.strftime("%d/%m/%Y")
        else:
            dtmp = pd.to_datetime(dv, dayfirst=True, errors="coerce")
            date_txt = dtmp.strftime("%d/%m/%Y") if not pd.isna(dtmp) else ""

        heure_txt = normalize_time_string(row.get("HEURE")) or "??:??"
        lines.append(f"{prefix}📆 {date_txt} | ⏱ {heure_txt}")

        # ------------------
        # Sens / Destination
        # ------------------
        sens_txt = format_sens_ar(row.get("Unnamed: 8"))
        dest = resolve_client_alias(str(row.get("DESIGNATION", "") or "").strip())
        if sens_txt or dest:
            lines.append(f"➡ {sens_txt} ({dest})".strip())

        # ------------------
        # Client
        # ------------------
        nom = str(row.get("NOM", "") or "").strip()
        if nom:
            lines.append(f"🧑 {nom}")

        # ------------------
        # 👥 PAX
        # ------------------
        pax = row.get("PAX")
        if pax not in ("", None, 0, "0"):
            try:
                pax_i = int(pax)
                if pax_i > 0:
                    lines.append(f"👥 {pax_i} pax")
            except Exception:
                lines.append(f"👥 {pax} pax")

        # ------------------
        # 🚘 Véhicule
        # ------------------
        if row.get("IMMAT"):
            lines.append(f"🚘 Plaque : {row.get('IMMAT')}")

        siege_bebe = extract_positive_int(row.get("SIEGE", row.get("SIÈGE")))
        if siege_bebe:
            lines.append(f"🍼 Siège bébé : {siege_bebe}")

        reh_n = extract_positive_int(row.get("REH"))
        if reh_n:
            lines.append(f"🪑 Rehausseur : {reh_n}")

        # ------------------
        # Adresse / Tel
        # ------------------
        adr = build_full_address_from_row(row)
        if adr:
            lines.append(f"📍 {adr}")

        tel = get_client_phone_from_row(row)
        if tel:
            lines.append(f"📞 {tel}")

        # ------------------
        # Paiement
        # ------------------
        paiement = str(row.get("PAIEMENT", "") or "").lower().strip()
        caisse = row.get("Caisse")

        if paiement == "facture":
            lines.append("🧾 FACTURE")
        elif paiement == "caisse" and caisse:
            lines.append(f"💶 {caisse} € (CASH)")
        elif paiement == "bancontact" and caisse:
            lines.append(f"💳 {caisse} € (BANCONTACT)")

        # ------------------
        # Vol + statut
        # ------------------
        vol = extract_vol_val(row, cols)
        if vol:
            lines.append(f"✈️ Vol {vol}")
            if should_query_flight_status(row):
                status, delay_min, *_ = get_flight_status_cached(vol)
                badge = flight_badge(status, delay_min)
                if badge:
                    lines.append(f"📡 {badge}")
            else:
                lines.append("📡 Statut vol : vérif J-1 / J")

        # ------------------
        # GO
        # ------------------
        go_val = str(row.get("GO", "") or "").strip()
        if go_val:
            lines.append(f"🟢 {go_val}")

        # ------------------
        # 🧾 BDC
        # ------------------
        for cand in ["NUM BDC", "Num BDC", "NUM_BDC", "BDC"]:
            if cand in cols and row.get(cand):
                lines.append(f"🧾 BDC : {row.get(cand)}")
                break

        # ------------------
        # Séparation navettes
        # ------------------
        lines.append("")


    return "\n".join(lines).strip()




def get_client_phone_from_row(row: pd.Series) -> str:
    """
    Récupère le numéro GSM du client.
    Ta colonne dans l'Excel s'appelle 'Tél'.
    On ajoute aussi des variantes au cas où.
    """
    candidate_cols = [
        "Tél",          # ta colonne principale
        "TEL",          # variantes possibles
        "Tel",
        "Téléphone",
        "GSM",
        "N° GSM",
        "N°GSM",
        "TEL CLIENT",
        "TEL_CLIENT",
        "PHONE",
    ]

    for col in candidate_cols:
        if col in row.index:
            val = row.get(col)
            if val is not None and str(val).strip():
                return str(val).strip()

    return ""
def normalize_ch_for_phone(ch_code: str) -> str:
    """
    Normalise le code chauffeur pour retrouver son GSM / MAIL dans Feuil2.

    Règles métier finales :
      - 'DO*'   -> 'DO'
      - 'DOFA'  -> 'DO'
      - 'FADO'  -> 'DO'
      - 'FA*'   -> 'FA'
      - 'FA1*'  -> 'FA1'
      - 'AD*'   -> 'AD'
      - 'NP*'   -> 'NP'
    """
    if not ch_code:
        return ""

    code = str(ch_code).strip().upper()

    # Supprimer les étoiles
    code = code.replace("*", "")

    # 🔥 PRIORITÉ ABSOLUE À DO
    if "DO" in code:
        return "DO"

    # Liste des chauffeurs connus (Feuil2)
    try:
        known = [c.strip().upper() for c in get_chauffeurs()]
    except Exception:
        known = []

    # Code exact connu
    if code in known:
        return code

    # Préfixe connu (FA*, NPX → FA / NP)
    if not code[-1].isdigit():
        for k in known:
            if code.startswith(k):
                return k

    return code

def build_client_sms(row: pd.Series, tel_chauffeur: str) -> str:
    """
    Construit le message SMS/WhatsApp envoyé au client
    pour confirmer son transfert.
    """
    # DATE
    d_val = row.get("DATE", "")
    if isinstance(d_val, date):
        d_txt = d_val.strftime("%d/%m/%Y")
    else:
        try:
            d_txt = pd.to_datetime(d_val, dayfirst=True, errors="coerce").strftime("%d/%m/%Y")
        except Exception:
            d_txt = str(d_val or "").strip()

    # HEURE
    heure = normalize_time_string(row.get("HEURE", "")) or "??:??"

    # NOM client (si dispo)
    nom_client = str(row.get("NOM", "") or "").strip()
    if nom_client:
        bonjour = f"Bonjour Mr / Mme {nom_client}, c'est Airports-Lines."
    else:
        bonjour = "Bonjour, c'est Airports-Lines."

    # Code chauffeur (CH)
    ch_code = str(row.get("CH", "") or "").strip()

    return (
        f"{bonjour}\n"
        f"Votre transfert du {d_txt} à {heure} est confirmé.\n"
        f"Votre chauffeur sera {ch_code} (GSM {tel_chauffeur}).\n"
        f"Merci pour votre confiance."
    )
def build_client_sms_from_driver(row: pd.Series, ch_code: str, tel_chauffeur: str) -> str:
    """
    Message WhatsApp envoyé par le chauffeur au client,
    SANS mentionner l'adresse du point de rendez-vous.
    """

    # DATE
    d_val = row.get("DATE", "")
    if isinstance(d_val, date):
        d_txt = d_val.strftime("%d/%m/%Y")
    else:
        try:
            d_txt = pd.to_datetime(d_val, dayfirst=True, errors="coerce").strftime("%d/%m/%Y")
        except Exception:
            d_txt = str(d_val or "").strip()

    # HEURE
    heure = normalize_time_string(row.get("HEURE", "")) or "??:??"

    # Nom du client
    nom_client = str(row.get("NOM", "") or "").strip()
    if nom_client:
        bonjour = f"Bonjour Mr / Mme {nom_client}, c'est votre chauffeur {ch_code} pour Airports-Lines."
    else:
        bonjour = f"Bonjour, c'est votre chauffeur {ch_code} pour Airports-Lines."

    # Message SANS adresse
    lignes = [
        bonjour,
        f"Je serai bien à l'heure prévue le {d_txt} à {heure}.",
    ]

    if tel_chauffeur:
        lignes.append(f"Voici mon numéro : {tel_chauffeur}.")

    lignes.append("En cas de problème, n’hésitez pas à me prévenir.")

    return "\n".join(lignes)

def show_client_messages_for_period(df_base: pd.DataFrame, start: date, nb_days: int):
    """
    Prépare et affiche la liste des messages clients (WhatsApp/SMS)
    pour une période donnée à partir du planning, avec diagnostics.
    """
    end = start + timedelta(days=nb_days - 1)

    df = df_base.copy()
    if "DATE" not in df.columns:
        st.warning("La colonne DATE est manquante dans le planning, impossible de filtrer.")
        return

    # Normalisation des dates en objets date
    try:
        df["DATE_TMP"] = pd.to_datetime(df["DATE"], dayfirst=True, errors="coerce").dt.date
    except Exception:
        df["DATE_TMP"] = pd.NaT

    mask = df["DATE_TMP"].notna() & (df["DATE_TMP"] >= start) & (df["DATE_TMP"] <= end)
    df = df[mask].copy()
    df.drop(columns=["DATE_TMP"], inplace=True, errors="ignore")

    if df.empty:
        st.info("Aucune navette client sur cette période (planning vide).")
        return

    st.markdown(
        f"#### Messages clients pour la période du "
        f"{start.strftime('%d/%m/%Y')} au {end.strftime('%d/%m/%Y')}"
    )

    st.caption(f"{len(df)} ligne(s) dans le planning sur cette période (avant filtrage).")

    cols = df.columns.tolist()
    lignes_indispo = 0
    lignes_sans_tel = 0
    lignes_sans_ch_phone = 0
    lignes_affichees = 0

    for _, row in df.iterrows():
        # 1) On ignore les lignes d'indisponibilité
        if is_indispo_row(row, cols):
            lignes_indispo += 1
            continue

        # 2) Numéro client
        client_phone = get_client_phone_from_row(row)
        if not client_phone:
            lignes_sans_tel += 1
            continue

        # 3) GSM chauffeur (si absent, on affiche quand même mais sans lien WhatsApp fonctionnel)
        raw_ch_code = str(row.get("CH", "") or "").strip()

        # On normalise le code pour retrouver le bon chauffeur dans Feuil2
        norm_ch_code = normalize_ch_for_phone(raw_ch_code)
        tel_ch, _mail_ch = get_chauffeur_contact(norm_ch_code) if norm_ch_code else ("", "")
        if not tel_ch:
            lignes_sans_ch_phone += 1

        # Construire le texte du message
        msg = build_client_sms(row, tel_ch or "??")
        wa_url = build_whatsapp_link(client_phone, msg) if tel_ch else None

        # Affichage : date / heure / nom client
        date_val = row.get("DATE", "")
        if isinstance(date_val, date):
            d_txt = date_val.strftime("%d/%m/%Y")
        else:
            try:
                d_txt = pd.to_datetime(date_val, dayfirst=True, errors="coerce").strftime("%d/%m/%Y")
            except Exception:
                d_txt = str(date_val or "").strip()

        heure = normalize_time_string(row.get("HEURE", "")) or "??:??"
        nom_client = str(row.get("NOM", "") or "").strip()
        label_client = nom_client if nom_client else "(client sans nom)"

        if wa_url:
            st.markdown(
                f"- **{d_txt} {heure}** – {label_client} – CH {raw_ch_code} → "
                f"[Envoyer WhatsApp au client]({wa_url})"
            )
        else:
            st.markdown(
                f"- **{d_txt} {heure}** – {label_client} – CH {raw_ch_code} "
                f"⚠ pas de GSM chauffeur configuré (Feuil2)."
            )

        lignes_affichees += 1

    # Résumé des filtres
    st.markdown("---")
    st.caption(
        f"Résumé : {lignes_affichees} navette(s) affichée(s) • "
        f"{lignes_indispo} indispo(s) ignorée(s) • "
        f"{lignes_sans_tel} sans numéro client ('Tél') • "
        f"{lignes_sans_ch_phone} sans GSM chauffeur."
    )
    st.caption(
        "⚠ Les messages ne partent pas automatiquement : "
        "clique sur chaque lien WhatsApp pour les envoyer."
    )

import time

def silent_tab_refresh(tab_key: str, interval_sec: int = 60):
    """
    Rafraîchissement silencieux par onglet.
    Ne touche PAS à la session login.
    """
    now = time.time()

    last = st.session_state["tab_refresh"].get(tab_key, 0)

    if now - last >= interval_sec:
        st.session_state["tab_refresh"][tab_key] = now
        return True  # on recharge les données

    return False

# ============================================================
#   HELPERS — ENVOI SMTP
# ============================================================

def send_email_smtp(to_email: str, subject: str, body: str) -> bool:
    """Envoie un e-mail texte simple via SMTP. Retourne True si OK."""
    if not to_email:
        return False

    try:
        msg = MIMEText(body or "", "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = FROM_EMAIL
        msg["To"] = to_email

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)

        return True

    except Exception as e:
        st.error(f"Erreur en envoyant le mail à {to_email} : {e}")
        return False


def send_email_smtp_with_attachments(
    to_email: str,
    subject: str,
    body: str,
    attachments: list[tuple[str, bytes, str]] | None = None,
) -> bool:
    """Envoie un e-mail via SMTP avec pièces jointes.
    attachments = [(filename, content_bytes, mime_subtype)] ; ex: ('planning.pdf', pdf_bytes, 'pdf')
    """
    if not to_email:
        return False

    try:
        msg = MIMEMultipart()
        msg["Subject"] = subject
        msg["From"] = FROM_EMAIL
        msg["To"] = to_email

        msg.attach(MIMEText(body or "", "plain", "utf-8"))

        if attachments:
            for (filename, content, subtype) in attachments:
                if content is None:
                    continue
                part = MIMEApplication(content, _subtype=(subtype or "octet-stream"))
                part.add_header("Content-Disposition", "attachment", filename=filename)
                msg.attach(part)

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)

        return True

    except Exception as e:
        st.error(f"Erreur en envoyant le mail à {to_email} : {e}")
        return False


import urllib.parse


def build_outlook_mailto(to, subject, body):
    subject = urllib.parse.quote(subject, safe="")
    body = urllib.parse.quote(body, safe="")
    return f"mailto:{to}?subject={subject}&body={body}"



# ============================================================
#   RÔLES — RESTRICTION GO/GL (Pour LEON)
# ============================================================

def role_allows_go_gl_only() -> bool:
    return st.session_state.get("role") == "restricted"


def leon_allowed_for_row(go_val: str) -> bool:
    """Leon ne peut agir QUE sur GO / GL."""
    if not role_allows_go_gl_only():
        return True
    val = (go_val or "").upper().strip()
    return val in ["GO", "GL"]


# ============================================================
#   LOGOUT (DÉCONNEXION PROPRE ET SÉCURISÉE)
# ============================================================

def logout():
    """Déconnexion volontaire uniquement sur cet appareil."""
    clear_persistent_session()
    clear_login_cookie()
    for k in (
        "logged_in",
        "username",
        "role",
        "chauffeur_code",
        "session_token",
        "_persist_state_saved",
    ):
        st.session_state.pop(k, None)

    try:
        st.cache_data.clear()
    except Exception:
        pass

    st.info("Déconnecté.")
    st.rerun()

# ============================================================

def _send_planning_next_3_days_to_all(*, want_pdf: bool = True) -> dict:
    """
    Envoie le planning des 3 prochains jours (J à J+2) à tous les chauffeurs actifs sur la période.
    Retourne un petit récap {sent, skipped_empty, missing_email, errors}.
    """
    from datetime import date, timedelta

    recap = {"sent": 0, "skipped_empty": 0, "missing_email": 0, "errors": 0}

    # 🔐 sécurité: réservé admin
    if st.session_state.get("role") != "admin":
        st.warning("⛔ Envoi planning réservé au bureau (admin).")
        return recap

    today = date.today()
    d_start = today + timedelta(days=1)
    d_end = today + timedelta(days=3)

    # init table log si dispo
    try:
        ensure_send_log_table()
    except Exception:
        pass

    # chauffeurs actifs sur période (SQL direct comme l’onglet Bureau)
    active_chauffeurs = set()
    try:
        with get_connection() as conn:
            df_chcol = pd.read_sql_query(
                """
                SELECT CH
                FROM planning
                WHERE COALESCE(IS_INDISPO,0)=0
                  AND COALESCE(IS_SUPERSEDED,0)=0
                  AND DATE_ISO BETWEEN ? AND ?
                  AND COALESCE(CH,'') <> ''
                """,
                conn,
                params=(d_start.isoformat(), d_end.isoformat()),
            )
        if not df_chcol.empty:
            for raw in df_chcol["CH"].dropna().astype(str):
                for c in split_chauffeurs(raw):
                    if c:
                        active_chauffeurs.add(c)
    except Exception:
        active_chauffeurs = set()

    chauffeurs = sorted(active_chauffeurs)
    if not chauffeurs:
        st.info("Aucun chauffeur trouvé sur les 3 prochains jours.")
        return recap

    periode_label = "3 prochains jours"
    for ch in dict.fromkeys(chauffeurs):
        try:
            _tel, mail = get_chauffeur_contact(ch)
            if not mail:
                recap["missing_email"] += 1
                try:
                    log_send(ch, "MAIL", periode_label, "KO", "Email manquant")
                except Exception:
                    pass
                continue

            df_ch = get_chauffeur_planning(
                chauffeur=ch,
                from_date=d_start,
                to_date=d_end,
            )

            if df_ch is None or df_ch.empty:
                recap["skipped_empty"] += 1
                try:
                    log_send(ch, "MAIL", periode_label, "OK", "Aucune navette (pas d'envoi)")
                except Exception:
                    pass
                continue

            # corps mail (identique)
            try:
                body = build_planning_mail_body(
                    df_ch=df_ch,
                    ch=ch,
                    from_date=d_start,
                    to_date=d_end,
                )
            except Exception:
                body = f"Planning {ch} — {periode_label}."

            subject = f"[PLANNING] {ch} — {periode_label}"

            if want_pdf:
                planning_cols_driver = [
                    "DATE","HEURE","CH","²²²²","IMMAT","PAX","Reh","Siège",
                    "Unnamed: 8","DESIGNATION","H South","Décollage","N° Vol","Origine",
                    "GO","Num BDC","NOM","ADRESSE","CP","Localité","Tél",
                    "Type Nav","PAIEMENT","Caisse"
                ]
                df_table = df_ch.copy()
                for c in planning_cols_driver:
                    if c not in df_table.columns:
                        df_table[c] = ""
                df_table = df_table[planning_cols_driver]

                # ✅ Affichage plage horaire (HEURE début + ²²²² = fin) quand ²²²² contient une vraie heure
                def _looks_like_time_local(v):
                    try:
                        t = normalize_time_string(v)
                        return t if t else ""
                    except Exception:
                        return ""

                if "HEURE" in df_table.columns and "²²²²" in df_table.columns:
                    h_deb = df_table["HEURE"].astype(str).map(lambda x: normalize_time_string(x) or str(x))
                    h_fin = df_table["²²²²"].astype(str).map(_looks_like_time_local)

                    # Si fin valide et différente → on affiche HEURE sous forme "HH:MM–HH:MM"
                    df_table["HEURE"] = [
                        (f"{a}–{b}" if (a and b and a != b) else a)
                        for a, b in zip(h_deb, h_fin)
                    ]

                # ✅ IMPORTANT : on ne normalise PAS Unnamed: 8 ici (texte libre = on garde)


                pdf_buf = export_chauffeur_planning_table_pdf(df_table, ch)
                pdf_bytes = pdf_buf.getvalue() if hasattr(pdf_buf, "getvalue") else bytes(pdf_buf)

                fname = f"Planning_{ch}_{d_start.isoformat()}_{d_end.isoformat()}.pdf"

                ok = send_email_smtp_with_attachments(
                    to_email=mail,
                    subject=subject,
                    body=body,
                    attachments=[(fname, pdf_bytes, "pdf")],
                )
            else:
                ok = send_email_smtp(
                    to_email=mail,
                    subject=subject,
                    body=body,
                )

            if not ok:
                raise RuntimeError("SMTP : envoi échoué")

            recap["sent"] += 1
            try:
                log_send(ch, "MAIL", periode_label, "OK", "Envoyé")
            except Exception:
                pass

        except Exception as e:
            recap["errors"] += 1
            try:
                log_send(ch, "MAIL", periode_label, "KO", str(e))
            except Exception:
                pass

    return recap


def _topbar_sync_db_and_refresh(*, also_send_next3: bool = False):
    """Bouton top-bar : MAJ DB depuis Dropbox puis rafraîchit l'UI.
    Option: envoi aussi le planning des 3 prochains jours.
    """
    from datetime import datetime

    try:
        with st.spinner("🔄 Synchronisation Dropbox → DB…"):
            # force un refresh même si Dropbox n’a pas changé : on passe un ts unique
            try:
                sync_planning_from_today(excel_sync_ts=datetime.now().isoformat(), ui=True)
            except TypeError:
                sync_planning_from_today(ui=True)

        if also_send_next3:
            with st.spinner("📧 Envoi planning (3 prochains jours)…"):
                recap = _send_planning_next_3_days_to_all(want_pdf=True)
                st.success(
                    f"📧 Envoi terminé — envoyés: {recap['sent']} | vides: {recap['skipped_empty']} | emails manquants: {recap['missing_email']} | erreurs: {recap['errors']}"
                )

        # refresh UI
        try:
            st.cache_data.clear()
        except Exception:
            pass
        st.rerun()
    except Exception as e:
        st.error(f"❌ Erreur MAJ DB Dropbox: {e}")

#   TOP BAR (INFORMATIONS UTILISATEUR + DECONNEXION)
# ============================================================


def render_top_bar():
    # 4 colonnes : user | maj db | maj+envoi | déconnexion
    col1, col2, col3, col4 = st.columns([5, 2, 2, 1])

    role = st.session_state.get("role")

    # -------------------------------
    # 👤 Utilisateur connecté
    # -------------------------------
    with col1:
        user = st.session_state.get("username")

        if user:
            if role == "admin":
                label = "Admin"
            elif role == "restricted":
                label = "Restreint"
            elif role == "driver":
                ch = st.session_state.get("chauffeur_code")
                label = f"Chauffeur {ch}"
            else:
                label = role or ""

            st.markdown(f"👤 **{user}** — {label}")

    # -------------------------------
    # 🔄 MAJ DB Dropbox + Vue (ADMIN ONLY)
    # -------------------------------
    with col2:
        if role == "admin":
            if st.button("🔄 Maj DB (Dropbox) + vue", use_container_width=True, key="topbar_sync_db"):
                _topbar_sync_db_and_refresh(also_send_next3=False)
        elif role == "driver":
            if st.button("🔄 MAJ planning", use_container_width=True, key="topbar_driver_sync"):
                _topbar_sync_db_and_refresh(also_send_next3=False)
        else:
            st.empty()

    # -------------------------------
    # 📧 MAJ + Envoi planning 3 jours (ADMIN ONLY, PDF joint)
    # -------------------------------
    with col3:
        if role == "admin":
            if st.button("📧 Maj + envoyer planning (3j)", use_container_width=True, key="topbar_sync_send_3j"):
                _topbar_sync_db_and_refresh(also_send_next3=True)
        else:
            st.empty()

    # -------------------------------
    # 🔓 Déconnexion
    # -------------------------------
    with col4:
        if st.button("🔓 Déconnexion", use_container_width=True, key="topbar_logout"):
            logout()


#   STYLE PLANNING

#   STYLE PLANNING — TOUTES LES COULEURS (FINAL SAFE)
# ============================================================

def style_groupage_partage(styler):
    df = styler.data

    def style_row(row):
        styles = [""] * len(row)

        def _flag(val):
            try:
                return int(val or 0) == 1
            except Exception:
                return False

        # ======================================================
        # 🟦 Congé chauffeur (HEURE 00:00 + IMMAT numérique)
        # ======================================================
        try:
            heure = _normalize_heure_str(row.get("HEURE"))
        except Exception:
            heure = str(row.get("HEURE", "") or "").strip()

        immat = str(row.get("IMMAT", "") or "").strip()
        is_conge = (heure == "00:00") and immat.isdigit()

        if is_conge:
            return ["background-color: #e3f2fd"] * len(row)

        # ======================================================
        # 🔴 Indisponibilité (logique existante)
        # ======================================================
        if is_indispo_row(row, df.columns.tolist()):
            return ["background-color: #f8d7da"] * len(row)

        # ======================================================
        # 🟡 Groupage / Partage
        # ======================================================
        if _flag(row.get("IS_GROUPAGE")):
            return ["background-color: #fff3cd"] * len(row)

        if _flag(row.get("IS_PARTAGE")) and "HEURE" in df.columns:
            styles[df.columns.get_loc("HEURE")] = "background-color: #fff3cd"

        # ======================================================
        # 🟠 / 🟢 Couleur CH depuis Excel (si dispo)
        # ======================================================
        if "CH_COLOR" in df.columns and "CH" in df.columns:
            try:
                ch_color = str(row.get("CH_COLOR", "") or "").strip().lower()
            except Exception:
                ch_color = ""

            if ch_color == "orange":
                styles[df.columns.get_loc("CH")] = "background-color: #ffe0b2; font-weight: 700"
            elif ch_color == "green":
                styles[df.columns.get_loc("CH")] = "background-color: #c8e6c9; font-weight: 700"

        return styles

    return styler.apply(style_row, axis=1)
def style_indispo(styler):
    def _red(row):
        if row.get("IS_INDISPO", 0) == 1:
            return ["background-color: #ffb3b3"] * len(row)
        return [""] * len(row)

    return styler.apply(_red, axis=1)

def style_chauffeur_confirmation(styler):
    df = styler.data

    if "CH" not in df.columns:
        return styler

    def _style_ch(confirmed, is_new):
        try:
            is_new_i = int(is_new or 0)
        except Exception:
            is_new_i = 0

        try:
            confirmed_i = int(confirmed or 0)
        except Exception:
            confirmed_i = 0

        # 🟠 modifié (Excel CH orange)
        if is_new_i == 1:
            return "background-color: #fff3cd; font-weight: bold"

        # 🟢 confirmé (Excel CH vert)
        if confirmed_i == 1:
            return "background-color: #d1e7dd; font-weight: bold"

        # ⚪ normal = OK
        return ""

    def apply_col(col):
        if col.name != "CH":
            return [""] * len(col)

        out = []
        for i in range(len(col)):
            confirmed = df.iloc[i].get("CONFIRMED")
            is_new = df.iloc[i].get("IS_NEW")
            out.append(_style_ch(confirmed, is_new))
        return out

    return styler.apply(apply_col, axis=0)



def style_caisse_payee(styler):
    df = styler.data

    if "Caisse" not in df.columns or "PAIEMENT" not in df.columns:
        return styler

    def style_row(row):
        styles = [""] * len(row)
        idx = df.columns.get_loc("Caisse")

        if str(row.get("PAIEMENT", "")).lower() == "caisse":
            if int(row.get("CAISSE_PAYEE", 0) or 0) == 1:
                styles[idx] = "background-color:#d1e7dd;font-weight:bold"
            else:
                styles[idx] = "background-color:#f8d7da;font-weight:bold"

        return styles

    return styler.apply(style_row, axis=1)



def format_chauffeur_ui(ch, confirmed):
    """
    Retourne le chauffeur avec couleur + icône.
    """
    ch = str(ch or "").strip().upper()

    if confirmed == 1:
        return f"🟢 <b>{ch}</b>"
    return f"🟠 <b>{ch}</b>"


def format_caisse_ui(paiement, montant, caisse_payee):
    """
    Retourne l'affichage paiement caisse avec couleur.
    """
    try:
        montant = float(montant)
    except Exception:
        montant = None

    paiement = str(paiement or "").lower().strip()

    if paiement == "caisse" and montant:
        if caisse_payee == 1:
            return (
                "<span style='color:#2e7d32;font-weight:700;'>"
                f"💶 {montant:.2f} € (PAYÉ)</span>"
            )
        return (
            "<span style='color:#d32f2f;font-weight:700;'>"
            f"💶 {montant:.2f} € (NON PAYÉ)</span>"
        )

    return ""



# ============================================================
#   PDF CHAUFFEUR – FEUILLE DE ROUTE
# ============================================================

def create_chauffeur_pdf(df_ch: pd.DataFrame, ch_selected: str, day_label: str) -> bytes:
    """
    Génère une feuille PDF claire pour le chauffeur.
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    y = height - 2 * cm

    c.setFont("Helvetica-Bold", 14)
    c.drawString(2 * cm, y, f"Feuille chauffeur — {ch_selected} — {day_label}")
    y -= 1 * cm
    c.setFont("Helvetica", 10)

    cols = df_ch.columns.tolist()

    for _, row in df_ch.iterrows():

        if y < 3 * cm:
            c.showPage()
            y = height - 2 * cm
            c.setFont("Helvetica-Bold", 14)
            c.drawString(2 * cm, y, f"Feuille chauffeur — {ch_selected} — {day_label}")
            y -= 1 * cm
            c.setFont("Helvetica", 10)

        # Indisponibilité
        if is_indispo_row(row, cols):
            heure = normalize_time_string(row.get("HEURE", ""))
            fin = normalize_time_string(row.get("²²²²", ""))
            c.drawString(2 * cm, y, f"{heure or '??:??'} → {fin or '??:??'} — 🚫 Indisponible")
            y -= 1 * cm
            continue

        # Heure
        heure = normalize_time_string(row.get("HEURE", "")) or "??:??"

        # Destination
        designation = str(row.get("DESIGNATION", "") or "").strip()
        route = ""
        for cnd in ["Unnamed: 8", "DESIGNATION"]:
            if cnd in cols and row.get(cnd):
                route = str(row[cnd]).strip()
                break

        if route and designation and designation not in route:
            dest = f"{route} ({designation})"
        else:
            dest = route or designation or "Navette"

        # Groupage / Partage
        g = bool_from_flag(row.get("GROUPAGE", "0"))
        p = bool_from_flag(row.get("PARTAGE", "0"))
        prefix = "[GRP] " if g else "[PARTAGE] " if p else ""

        # Ligne principale
        ligne1 = f"{prefix}{heure} – {dest}"

        # Nom client
        nom = str(row.get("NOM", "") or "")
        if nom:
            ligne1 += f" – {nom}"

        c.drawString(2 * cm, y, ligne1)
        y -= 0.5 * cm

        # Adresse
        adresse = str(row.get("ADRESSE", "") or "").strip()
        cp = str(row.get("CP", "") or "").strip()
        loc = str(row.get("Localité", "") or row.get("LOCALITE", "") or "").strip()
        adr_full = " ".join(x for x in [adresse, cp, loc] if x)

        if adr_full:
            c.drawString(2 * cm, y, adr_full)
            y -= 0.5 * cm

        # Vol
        infos_vol = []
        if row.get("N° Vol"): infos_vol.append(f"Vol {row.get('N° Vol')}")
        if row.get("Origine"): infos_vol.append(f"Origine {row.get('Origine')}")
        if row.get("Décollage"): infos_vol.append(f"Décollage {row.get('Décollage')}")
        if row.get("H South"): infos_vol.append(f"H SO {row.get('H South')}")
        if infos_vol:
            c.drawString(2 * cm, y, " | ".join(infos_vol))
            y -= 0.5 * cm
        # ✈️ Numéro de vol (PDF)
        vol_val = ""
        for col in ["N° Vol", "N° Vol ", "Num Vol", "VOL", "Vol"]:
            if col in df_ch.columns:
                v = str(row.get(col, "") or "").strip()
                if v:
                    vol_val = v
                    break
        
        if vol_val:
            if should_query_flight_status(row):
                status, delay_min, sched_dt, est_dt = get_flight_status_cached(vol_val)
                badge = flight_badge(status, delay_min)
            else:
                status, delay_min, sched_dt, est_dt = "", 0, None, None
                badge = ""




        # Paiement / caisse
        infos_pay = []
        if row.get("PAX"): infos_pay.append(f"PAX {row.get('PAX')}")
        if row.get("PAIEMENT"): infos_pay.append(f"Paiement {row.get('PAIEMENT')}")
        if row.get("Caisse"): infos_pay.append(f"Caisse : {row.get('Caisse')} €")
        if infos_pay:
            c.drawString(2 * cm, y, " | ".join(infos_pay))

        y -= 1 * cm

    c.save()
    pdf = buffer.getvalue()
    buffer.close()
    return pdf


# ============================================================
#   MESSAGES POUR WHATSAPP / MAIL — VUE CHAUFFEUR
# ============================================================

def build_chauffeur_day_message(df_ch: pd.DataFrame, ch_selected: str, day_label: str) -> str:
    cols = df_ch.columns.tolist()
    lines = []
    lines.append(f"🚖 Planning du {day_label} — Chauffeur : {ch_selected}")
    lines.append("")

    for _, row in df_ch.iterrows():

        if is_indispo_row(row, cols):
            h1 = normalize_time_string(row.get("HEURE", ""))
            h2 = normalize_time_string(row.get("²²²²", ""))
            lines.append(f"⏱ {h1} → {h2} — 🚫 Indisponible")
            lines.append("")
            continue

        heure = normalize_time_string(row.get("HEURE", "")) or "??:??"

        designation = str(row.get("DESIGNATION", "") or "").strip()
        route = ""
        for cnd in ["Unnamed: 8", "DESIGNATION"]:
            if cnd in cols and row.get(cnd):
                route = str(row[cnd]).strip()
                break

        if route and designation and designation not in route:
            dest = f"{route} ({designation})"
        else:
            dest = route or designation or "Navette"

        dest = resolve_client_alias(dest)

        nom = str(row.get("NOM", "") or "")

        # Groupage
        g = bool_from_flag(row.get("GROUPAGE", "0"))
        p = bool_from_flag(row.get("PARTAGE", "0"))
        prefix = "[GRP] " if g else "[PARTAGE] " if p else ""

        line = f"{prefix}➡ {heure} — {dest}"
        if nom:
            line += f" — {nom}"
        lines.append(line)

        # Adresse
        adr = " ".join(
            x for x in [
                str(row.get("ADRESSE", "") or "").strip(),
                str(row.get("CP", "") or "").strip(),
                str(row.get("Localité", "") or row.get("LOCALITE", "") or "").strip(),
            ] if x
        )
        if adr:
            lines.append(f"📍 {adr}")

        # Extras
        extra = []
        if row.get("PAX"): extra.append(f"{row.get('PAX')} pax")
        if row.get("PAIEMENT"): extra.append(f"Paiement {row.get('PAIEMENT')}")
        if row.get("Caisse"): extra.append(f"Caisse {row.get('Caisse')} €")
        if extra:
            lines.append(" | ".join(extra))

        if g: lines.append("🔶 Groupage")
        if p: lines.append("🟨 Navette partagée")

        lines.append("")

    return "\n".join(lines).strip()

# ============================================================
#   ONGLET 📅 PLANNING — VUE RAPIDE AVEC COULEURS (EXCEL-LIKE)
#   + BADGE STATUT en colonne dédiée
# ============================================================

def render_tab_planning():
    st.subheader("📅 Planning — vue rapide")

    # ===================================================
    # 🔄 Rafraîchissement manuel UNIQUEMENT
    # ===================================================
    consume_soft_refresh("planning")
    if st.button("🔄 Rafraîchir la vue planning", key="btn_refresh_planning"):
        request_soft_refresh("planning")

    today = date.today()

    # ===================================================
    # 📆 PÉRIODE (SOURCE UNIQUE)
    # ===================================================
    if "planning_start" not in st.session_state:
        st.session_state.planning_start = today
        st.session_state.planning_end = today

    colb1, colb2, colb3, colb4 = st.columns(4)

    with colb1:
        if st.button("📆 Aujourd’hui"):
            st.session_state.planning_start = today
            st.session_state.planning_end = today

    with colb2:
        if st.button("📆 Demain"):
            d = today + timedelta(days=1)
            st.session_state.planning_start = d
            st.session_state.planning_end = d

    with colb3:
        if st.button("📆 Cette semaine"):
            lundi = today - timedelta(days=today.weekday())
            dimanche = lundi + timedelta(days=6)
            st.session_state.planning_start = lundi
            st.session_state.planning_end = dimanche

    with colb4:
        if st.button("📆 7 prochains jours"):
            st.session_state.planning_start = today
            st.session_state.planning_end = today + timedelta(days=6)

    start_date = st.session_state.planning_start
    end_date = st.session_state.planning_end

    st.caption(
        f"📅 Période : **{start_date.strftime('%d/%m/%Y')} → {end_date.strftime('%d/%m/%Y')}**"
    )

    # ===================================================
    # 🔍 FILTRES UI
    # ===================================================
    colf1, colf2 = st.columns([2, 1])

    with colf1:
        search = st.text_input("🔍 Recherche (client, vol, lieu…)", "")

    with colf2:
        chs = get_chauffeurs_for_ui()
        ch_value = st.selectbox("🚖 Chauffeur", ["(Tous)"] + chs)
        ch_value = None if ch_value == "(Tous)" else ch_value

    # ===================================================
    # 📖 LECTURE DB (période stricte)
    # ===================================================
    with get_connection() as conn:
        df = pd.read_sql_query(
            """
            SELECT *
            FROM planning
            WHERE
                DATE_ISO >= ?
                AND DATE_ISO <= ?
                AND COALESCE(IS_SUPERSEDED,0) = 0
            ORDER BY DATE_ISO, HEURE
            """,
            conn,
            params=(start_date.isoformat(), end_date.isoformat()),
        )

    if df is None or df.empty:
        st.info("Aucune navette ou indisponibilité pour cette période.")
        return

    # ===================================================
    # 🧠 NORMALISATION DATE
    # ===================================================
    if "DATE_ISO" in df.columns:
        df["DATE_OBJ"] = pd.to_datetime(df["DATE_ISO"], errors="coerce").dt.date
    else:
        df["DATE_OBJ"] = pd.to_datetime(df.get("DATE"), dayfirst=True, errors="coerce").dt.date

    # ===================================================
    # 🧹 FILTRE CHAUFFEUR
    # ===================================================
    if ch_value and "CH" in df.columns:
        ch_norm = normalize_ch_code(ch_value)
        df = df[
            df["CH"]
            .fillna("")
            .astype(str)
            .str.upper()
            .str.replace("*", "", regex=False)
            .str.startswith(ch_norm)
        ]

    # ===================================================
    # 🔍 RECHERCHE TEXTE
    # ===================================================
    if search:
        mask = False
        cols_search = [c for c in ["DESIGNATION", "NOM", "ADRESSE", "N° Vol", "VOL", "Localité", "LOCALITE", "REMARQUE", "Unnamed: 8"] if c in df.columns]
        for col in cols_search:
            mask |= (
                df[col]
                .fillna("")
                .astype(str)
                .str.contains(search, case=False, na=False)
            )
        df = df[mask]

    if df.empty:
        st.info("Aucune donnée après filtres.")
        return

    # ===================================================
    # 🚫 BADGE STATUT (colonne dédiée)
    #   - Bureau ≠ indispo
    #   - MA en rouge
    #   - congé en jaune
    # ===================================================
    def _fmt_date(d):
        try:
            return d.strftime("%d/%m/%Y")
        except Exception:
            return ""

    def _hsort(val):
        s = str(val or "").strip()
        if "→" in s:
            s = s.split("→", 1)[0].strip()
        s = s.replace("h", ":").replace("H", ":")
        if ":" in s:
            try:
                hh, mm = s.split(":")[0:2]
                return (int(hh), int(mm))
            except Exception:
                return (99, 99)
        return (99, 99)

    def compute_statut(row) -> str:
        # Bureau jamais indispo
        sens = str(row.get("Unnamed: 8", "") or "").upper()
        if "BUREAU" in sens:
            return ""

        is_ind = int(row.get("IS_INDISPO", 0) or 0) == 1
        if not is_ind:
            return ""

        reason = str(row.get("INDISPO_REASON", "") or "").upper().strip()

        # Si la raison n'existe pas (anciennes DB), fallback sur ²²²² / IMMAT / heures
        if not reason:
            col2222 = str(row.get("²²²²", "") or "").upper().strip()
            immat = str(row.get("IMMAT", "") or "").upper().strip()
            heure = str(row.get("HEURE", "") or "").strip()
            # MA
            if col2222 == "MA" or immat == "MA":
                reason = "MALADE"
            # congé : immat = chiffre OU heure 00:00
            elif immat.isdigit() and len(immat) <= 2:
                reason = "CONGE"
            elif heure == "00:00":
                reason = "CONGE"
            # sinon
            else:
                reason = col2222 or "INDISPO"

        if reason == "MALADE" or reason == "MA":
            return "🟥 MALADIE"
        if reason in ("CONGE", "VAC", "VACANCES"):
            return "🏖 CONGÉ"
        if reason == "INDISPO_PLAGE":
            return "🟧 INDISPO"
        if reason:
            return f"🟨 {reason}"
        return "🟧 INDISPO"

    # Colonnes calculées
    df["DATE"] = df["DATE_OBJ"].apply(_fmt_date)
    df["STATUT"] = df.apply(compute_statut, axis=1)

    # ===================================================
    # 🔃 TRI (DATE + HEURE)
    # ===================================================
    df["_HSORT"] = df["HEURE"].apply(_hsort) if "HEURE" in df.columns else [(99, 99)] * len(df)
    df = df.sort_values(["DATE_OBJ", "_HSORT"], kind="mergesort").drop(columns=["_HSORT"], errors="ignore")

    # ===================================================
    # 🎨 PRÉPARATION AFFICHAGE (STATUT à côté de DATE)
    # ===================================================
    df_display = df.drop(columns=["DATE_OBJ", "id"], errors="ignore").copy()

    # Re-ordonner colonnes : DATE, STATUT, HEURE, CH, ...
    preferred = ["DATE", "STATUT", "HEURE", "CH"]
    cols = list(df_display.columns)
    new_cols = [c for c in preferred if c in cols] + [c for c in cols if c not in preferred]
    df_display = df_display[new_cols]

    if "CH" in df_display.columns:
        df_display["CH"] = df_display.apply(
            lambda r: format_chauffeur_colored(r.get("CH"), r.get("CONFIRMED")),
            axis=1,
        )

    styled = df_display.style
    styled = style_indispo(styled)
    styled = style_groupage_partage(styled)
    styled = style_caisse_payee(styled)

    # ===================================================
    # 🧹 MASQUER COLONNES TECHNIQUES (PANDAS SAFE)
    # ===================================================
    cols_to_hide = [
        "IS_GROUPAGE",
        "IS_PARTAGE",
        "IS_ATTENTE",
        "CONFIRMED",
        "CAISSE_PAYEE",
        "CH_COLOR",
        "IS_INDISPO",
        "IS_SUPERSEDED",
        "DATE_ISO",
        "INDISPO_REASON",
        "IS_BUREAU",
        "LOCKED_BY_APP",
        "EXCEL_UID",
        "EXCEL_SYNC_TS",
    ]
    cols_to_hide = [c for c in cols_to_hide if c in df_display.columns]

    if cols_to_hide:
        try:
            styled = styled.hide(columns=cols_to_hide)
        except TypeError:
            styled = styled.hide(subset=cols_to_hide, axis="columns")

    # ===================================================
    # 📊 AFFICHAGE FINAL
    # ===================================================
    st.dataframe(styled, use_container_width=True, height=520)


def _make_row_key_from_parts(date_iso: str, heure: str, nom: str, adresse: str, vol: str, heure_fin: str = "") -> str:
    def _s(x):
        return ("" if x is None else str(x)).strip().lower()
    h = normalize_time_string(heure) or _s(heure)
    hf = normalize_time_string(heure_fin) or _s(heure_fin)
    return f"{_s(date_iso)}|{h}|{_s(nom)}|{_s(adresse)}|{_s(vol)}|{hf}"

def apply_pending_ch_changes_to_dropbox_excel(sheet_name: str = "Feuil1") -> tuple[int, int]:
    pending = list_pending_actions(limit=2000)
    pending = [p for p in pending if str(p.get("action_type","")) == "CH_CHANGE"]
    if not pending:
        return (0, 0)

    xls_bytes = download_dropbox_excel_bytes()
    if not xls_bytes:
        return (len(pending), 0)

    wb = load_workbook(filename=BytesIO(xls_bytes))
    if sheet_name not in wb.sheetnames:
        return (len(pending), 0)

    ws = wb[sheet_name]

    headers = {}
    for col in range(1, ws.max_column + 1):
        v = ws.cell(row=1, column=col).value
        if v is not None:
            headers[str(v).strip()] = col

    if "CH" not in headers or "DATE" not in headers:
        return (len(pending), 0)

    col_date = headers.get("DATE")
    col_heure = headers.get("HEURE")
    col_nom = headers.get("NOM")
    col_adresse = headers.get("ADRESSE")
    col_vol = headers.get("VOL")
    col_hf = headers.get("HEURE_FIN") or headers.get("HEURE FIN") or headers.get("HEURE2") or headers.get("HEURE 2")

    rowkey_to_excelrow = {}
    for r in range(2, ws.max_row + 1):
        date_val = ws.cell(row=r, column=col_date).value
        try:
            if isinstance(date_val, (datetime, date)):
                date_iso = date_val.strftime("%Y-%m-%d")
            else:
                date_iso = _normalize_excel_date_to_iso(date_val)
        except Exception:
            date_iso = None
        if not date_iso:
            continue

        heure = ws.cell(row=r, column=col_heure).value if col_heure else ""
        nom = ws.cell(row=r, column=col_nom).value if col_nom else ""
        adr = ws.cell(row=r, column=col_adresse).value if col_adresse else ""
        vol = ws.cell(row=r, column=col_vol).value if col_vol else ""
        hf = ws.cell(row=r, column=col_hf).value if col_hf else ""

        rk = _make_row_key_from_parts(date_iso, str(heure or ""), str(nom or ""), str(adr or ""), str(vol or ""), str(hf or ""))
        rowkey_to_excelrow[rk] = r

    applied = 0
    done_ids = []
    unlocked_keys = []
    ch_col = headers["CH"]

    for a in pending:
        rk = str(a.get("row_key") or "").strip()
        new_ch = str(a.get("new_value") or "").strip()
        if not rk or not new_ch:
            continue
        excel_row = rowkey_to_excelrow.get(rk)
        if not excel_row:
            continue
        ws.cell(row=excel_row, column=ch_col).value = new_ch
        applied += 1
        done_ids.append(int(a["id"]))
        unlocked_keys.append(rk)

    if applied == 0:
        return (len(pending), 0)

    out = BytesIO()
    wb.save(out)

    if not upload_dropbox_excel_bytes(out.getvalue()):
        return (len(pending), 0)

    mark_actions_done(done_ids)
    unlock_rows_by_row_keys(unlocked_keys)

    return (len(pending), applied)

def render_tab_quick_day_mobile():
    """Vue jour admin : toutes les navettes du jour (tous chauffeurs) + changement chauffeur + WhatsApp + suivi Excel."""
    st.subheader("⚡ Vue jour (mobile) — Tous chauffeurs")

    # ===================================================
    # 🔁 Soft refresh contrôlé (zéro rerun brutal)
    # ===================================================
    if consume_soft_refresh("quick_day"):
        try:
            get_planning.clear()
        except Exception:
            pass

    sel_date_iso = date.today().strftime("%Y-%m-%d")

    # ===================================================
    # 📌 Actions en attente vers Excel (CH_CHANGE)
    # ===================================================
    pending_map = {}
    pending_list = []
    try:
        from database import list_pending_actions
        actions = list_pending_actions(limit=800)
        for (aid, rk, atype, oldv, newv, usr, created_at) in actions:
            if atype == "CH_CHANGE" and rk:
                pending_map[str(rk)] = {
                    "id": aid,
                    "new_ch": newv,
                    "old_ch": oldv,
                    "user": usr,
                    "created_at": created_at,
                }
                pending_list.append((aid, rk, oldv, newv, usr, created_at))
    except Exception:
        pending_map = {}
        pending_list = []

    # ===================================================
    # 📤 Appliquer toutes les modifs vers Excel (Feuil1)
    # ===================================================
    col_top1, col_top2 = st.columns([2, 1])
    with col_top1:
        if st.button("📤 Appliquer les changements chauffeur dans l’Excel (Feuil1)", key="qd_apply_all_excel"):
            with st.spinner("Mise à jour Excel en cours…"):
                total, applied = apply_pending_ch_changes_to_dropbox_excel("Feuil1")

            if applied > 0:
                st.success(f"✅ {applied}/{total} changement(s) appliqué(s) dans l’Excel")
                request_soft_refresh("quick_day")
            else:
                st.info("Aucun changement chauffeur à appliquer (ou lignes introuvables dans l’Excel).")

    with col_top2:
        if pending_map:
            st.metric("🟡 En attente Excel", len(pending_map))
        else:
            st.metric("🟡 En attente Excel", 0)

    # ===================================================
    # 1️⃣ Charger toute la journée (SAFE)
    # ===================================================
    df = get_planning(
        start_date=sel_date_iso,
        end_date=sel_date_iso,
        chauffeur=None,
        type_filter=None,
        search="",
        max_rows=400,
        source="day",
    )

    if df is None or df.empty:
        st.info("Aucune navette pour cette journée.")
        return

    df = apply_actions_overrides(df)
    df = df.copy()
    cols = df.columns.tolist()

    # ===================================================
    # Séparation navettes / indispos (évite vue vide)
    # ===================================================
    df_navettes = df[df.get("IS_INDISPO", 0) == 0].copy()
    df_indispos = df[df.get("IS_INDISPO", 0) == 1].copy()

    if df_navettes.empty and df_indispos.empty:
        st.info("Aucune navette ni indisponibilité pour cette journée.")
        return

    # ===================================================
    # 2️⃣ Liste chauffeurs
    # ===================================================
    chs_ui = get_chauffeurs_for_ui()
    if not chs_ui:
        chs_ui = get_chauffeurs() or CH_CODES

    # ===================================================
    # 3️⃣ Tri par heure (NAVETTES UNIQUEMENT)
    # ===================================================
    def _key_time(v):
        txt = normalize_time_string(v)
        if not txt:
            return datetime.max.time()
        for fmt in ("%H:%M:%S", "%H:%M"):
            try:
                return datetime.strptime(txt, fmt).time()
            except Exception:
                pass
        return datetime.max.time()

    if "HEURE" in df_navettes.columns:
        df_navettes["_sort_time"] = df_navettes["HEURE"].apply(_key_time)
        df_navettes = df_navettes.sort_values("_sort_time", ascending=True)

    st.markdown("### 📋 Détail des navettes (texte compact)")
    st.caption("Vue admin : toutes les navettes du jour. Les changements sont appliqués en DB immédiatement.")

    # ===================================================
    # 🟡 Bandeau global des modifications à reporter (détaillé)
    # ===================================================
    if pending_list:
        with st.expander("🟡 Modifications à reporter dans Excel (Feuil1)", expanded=True):
            for (aid, rk, oldv, newv, usr, created_at) in pending_list[:200]:
                st.markdown(f"• **{oldv} → {newv}** — row_key: `{rk}`"
                            + (f" — {usr}" if usr else "")
                            + (f" — {created_at}" if created_at else ""))

    # ===================================================
    # 🚫 Indisponibilités (MA, congés, etc.)
    # ===================================================
    if not df_indispos.empty:
        st.markdown("### 🚫 Indisponibilités")
        for _, row in df_indispos.iterrows():
            ch = str(row.get("CH", "") or "").strip()
            h1 = normalize_time_string(row.get("HEURE", "")) or ""
            h2 = normalize_time_string(row.get("_HEURE_FIN", "")) or ""
            reason = str(row.get("INDISPO_REASON", "") or "").strip() or "Indisponible"

            line = f"👤 {ch}"
            if h1 or h2:
                line += f" | ⏱ {h1} → {h2}"
            line += f" | 🚫 {reason}"
            st.markdown(line)

    # ===================================================
    # 📋 AFFICHAGE DES NAVETTES
    # ===================================================
    for _, row in df_navettes.iterrows():

        # ID
        try:
            row_id = int(row.get("id"))
        except Exception:
            continue

        # row_key
        rk = str(row.get("row_key") or "")

        # Date
        date_val = row.get("DATE")
        if isinstance(date_val, (datetime, date)):
            date_txt = date_val.strftime("%d/%m/%Y")
        else:
            dtmp = pd.to_datetime(date_val, dayfirst=True, errors="coerce")
            date_txt = dtmp.strftime("%d/%m/%Y") if not pd.isna(dtmp) else ""

        # Heure
        heure_txt = normalize_time_string(row.get("HEURE", "")) or "??:??"

        # Chauffeur
        ch_current = str(row.get("CH", "") or "").strip()

        # 🛠️ Badge manuel + 🟡 pending excel
        manual_badge = " 🛠️" if int(row.get("CH_MANUAL", 0) or 0) == 1 else ""
        pending_badge = " 🟡" if (rk and rk in pending_map) else ""

        # Destination (route + designation)
        designation = str(row.get("DESIGNATION", "") or "").strip()
        route_txt = str(row.get("Unnamed: 8", "") or "").strip()
        dest = f"{route_txt} ({designation})" if route_txt and designation else (route_txt or designation or "Navette")

        # Client
        nom = str(row.get("NOM", "") or "").strip()

        # Adresse
        adresse = str(row.get("ADRESSE", "") or "").strip()
        cp = str(row.get("CP", "") or "").strip()
        loc = str(row.get("Localité", "") or row.get("LOCALITE", "") or "").strip()
        adr_full = " ".join(x for x in [adresse, cp, loc] if x)

        # Extras
        pax = str(row.get("PAX", "") or "").strip()
        paiement = str(row.get("PAIEMENT", "") or "").strip()
        bdc = str(row.get("Num BDC", "") or "").strip()

        # ✈️ Vol + badge
        vol = extract_vol_val(row, cols)
        badge_vol = ""
        if vol:
            try:
                status, delay_min, sched_dt, est_dt = get_flight_status_cached(vol)
                badge_vol = flight_badge(status, delay_min)
            except Exception:
                badge_vol = ""

        # Ligne affichée
        line = f"📆 {date_txt} | ⏱ {heure_txt} | 👤 {ch_current}{manual_badge}{pending_badge} → {dest}"
        if nom:
            line += f" | 🙂 {nom}"
        if adr_full:
            line += f" | 📍 {adr_full}"
        if vol:
            line += f" | ✈️ {vol} {badge_vol}"
        if paiement:
            line += f" | 💳 {paiement}"
        if bdc:
            line += f" | 📄 BDC: {bdc}"
        if pax:
            line += f" | 👥 {pax} pax"

        with st.container(border=True):
            st.markdown(line)

            colA, colB, colC, colD = st.columns([2, 1, 1, 1])

            # 🔁 Remplacement chauffeur
            with colA:
                new_ch = st.selectbox(
                    "Remplacer chauffeur",
                    chs_ui,
                    index=chs_ui.index(ch_current) if ch_current in chs_ui else 0,
                    key=f"qd_newch_{row_id}",
                )

            # 💾 Sauvegarde DB + action en attente Excel
            with colB:
                if new_ch != ch_current:
                    if st.button("💾 Appliquer", key=f"qd_save_{row_id}"):

                        now_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        user = (
                            st.session_state.get("username")
                            or st.session_state.get("user")
                            or ""
                        )

                        # 1) DB : appliquer immédiatement
                        update_planning_row(
                            row_id,
                            {
                                "CH": new_ch,
                                "CH_MANUAL": 1,
                                "updated_at": now_iso,
                            },
                        )

                        # 2) Audit
                        try:
                            with get_connection() as conn:
                                conn.execute(
                                    """
                                    INSERT INTO planning_audit
                                    (ts, user, action, row_key, details)
                                    VALUES (?, ?, ?, ?, ?)
                                    """,
                                    (
                                        now_iso,
                                        user,
                                        "CH_MANUAL_CHANGE",
                                        rk,
                                        f"{ch_current} → {new_ch}",
                                    ),
                                )
                                conn.commit()
                        except Exception:
                            pass

                        # 3) Pending action vers Excel
                        try:
                            from database import log_ch_change
                            log_ch_change(rk, ch_current, new_ch, user=user)
                        except Exception:
                            pass

                        st.success(
                            "✅ Chauffeur modifié\n"
                            "🛠️ Override manuel actif\n"
                            "🟡 À reporter dans Excel"
                        )

                        # Refresh contrôlé (pas de rerun brutal)
                        request_soft_refresh("quick_day")
                else:
                    st.caption("")

            # 📤 Excel (par ligne si pending)
            with colD:
                if rk and rk in pending_map:
                    if st.button("📤 Excel", key=f"qd_excel_{row_id}"):
                        try:
                            from utils import update_excel_rows_by_row_key
                            from database import mark_actions_done

                            upd = {rk: {"CH": pending_map[rk].get("new_ch")}}
                            cnt = update_excel_rows_by_row_key(upd)
                            mark_actions_done([pending_map[rk].get("id")])

                            st.success(f"✅ Envoyé vers Excel ({cnt})")
                            request_soft_refresh("quick_day")
                        except Exception as e:
                            st.error(f"Erreur Excel : {e}")
                else:
                    st.caption("")

            # 💬 WhatsApp
            with colC:
                norm_ch = normalize_ch_for_phone(new_ch or ch_current)
                tel_ch, _ = get_chauffeur_contact(norm_ch) if norm_ch else ("", "")
                if tel_ch:
                    msg = (
                        f"Bonjour {new_ch or ch_current},\n"
                        f"Navette du {date_txt} à {heure_txt}\n"
                        f"Destination : {dest}\n"
                        + (f"Client : {nom}\n" if nom else "")
                        + (f"Adresse : {adr_full}\n" if adr_full else "")
                        + (f"PAX : {pax}\n" if pax else "")
                        + (f"BDC : {bdc}\n" if bdc else "")
                        + "Merci de confirmer si problème 🙏"
                    )
                    wa = build_whatsapp_link(tel_ch, msg)
                    st.markdown(f"[💬 WhatsApp]({wa})")
                else:
                    st.caption("No GSM")


# ============================================================
#   ONGLET 📊 TABLEAU / ÉDITION — EXCEL ONLINE → DB
# ============================================================
def render_tab_table():
    st.subheader("📊 Planning — Édition Excel Online")

    st.markdown(
        "Le planning s’édite dans **Excel Online**. "
        "La base locale est synchronisée **uniquement à partir d’aujourd’hui**."
    )

    EXCEL_ONLINE_URL = (
        "https://www.dropbox.com/scl/fi/lymuumy8en46l7p0uwjj3/"
        "Planning-2026.xlsx"
        "?rlkey=sgvr0a58ekpr471p5aguqk3k8&dl=0"
    )

    # 🌐 Ouvrir Excel Online
    st.markdown(
        f"""
        <a href="{EXCEL_ONLINE_URL}" target="_blank">
            <button style="
                padding:10px 16px;
                font-size:16px;
                background-color:#0f6cbd;
                color:white;
                border:none;
                border-radius:6px;
                cursor:pointer;
            ">
                🌐 Ouvrir le planning Excel Online
            </button>
        </a>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("---")


# ============================================================
#   ONGLET 🔍 CLIENTS — HISTORIQUE & CRÉATION RAPIDE
# ============================================================

def render_tab_clients():
    st.subheader("🔍 Clients — Historique & création rapide")

    query = st.text_input(
        "Nom du client (ou partie du nom)",
        "",
        key="client_search",
    )

    if not query.strip():
        st.info("Tape un nom de client pour afficher son historique.")
        return

    df = search_client(query, max_rows=500)
    if df.empty:
        st.warning("Aucune navette trouvée pour ce client.")
        return

    if "id" not in df.columns:
        st.error("La table `planning` doit contenir une colonne `id`.")
        return

    # max 40 colonnes
    if df.shape[1] > 40:
        df = df.iloc[:, :40]

    st.markdown(f"#### {len(df)} navette(s) trouvée(s)")

    df_display = df.copy()
    df_display = df_display.drop(columns=["id"])
    st.dataframe(df_display, use_container_width=True, height=400)

    # Sélection d’une navette modèle
    ids = df["id"].tolist()
    df_view = df.drop(columns=["id"]).copy().reset_index(drop=True)
    df_view.insert(0, "_SELECT", False)
    if "KM_EST" in df.columns:
        df_view["_KM_EST"] = df["KM_EST"].fillna("").astype(str)
    if "TEMPS_EST" in df.columns:
        df_view["_TEMPS_EST"] = df["TEMPS_EST"].fillna("").astype(str)
    # --- Affichage KM / TEMPS depuis la DB ---
    if "KM_EST" in df.columns:
        df_view["_KM_EST"] = df["KM_EST"].fillna("").astype(str)
    else:
        df_view["_KM_EST"] = ""

    if "TEMPS_EST" in df.columns:
        df_view["_TEMPS_EST"] = df["TEMPS_EST"].fillna("").astype(str)
    else:
        df_view["_TEMPS_EST"] = ""

    # Injecter KM / MIN si on a déjà calculé
    km_map = st.session_state.get("km_time_by_id", {}) or {}
    km_col = []
    min_col = []
    for rid in ids:
        km, mn = km_map.get(int(rid), (None, None))
        km_col.append("" if km is None else f"{km} km")
        min_col.append("" if mn is None else f"{mn} min")

    # Colonnes d'affichage (préfixe "_" pour éviter confusion avec colonnes Excel)
    df_view["_KM_EST"] = km_col
    df_view["_TEMPS_EST"] = min_col

    st.markdown("#### Sélectionne une navette modèle")
    edited = st.data_editor(
        df_view,
        use_container_width=True,
        height=300,
        num_rows="fixed",
        key="client_editor",
    )
    # ==================================================
    # D) Exécuter le calcul KM / TEMPS (à la demande)
    # ==================================================
    if st.session_state.get("km_time_run"):
        selected_indices = edited.index[edited["_SELECT"] == True].tolist()
        selected_ids = [int(ids[i]) for i in selected_indices]

        mode = st.session_state.get("km_time_last_mode", "✅ Lignes cochées (_SELECT)")
        targets = selected_ids if mode.startswith("✅") else [int(x) for x in ids]

        for rid in targets:
            row = df[df["id"] == rid].iloc[0]

            if row.get("KM_EST") and row.get("TEMPS_EST"):
                continue

            origin = (
                build_full_address_from_row(row)
                or st.session_state.get("km_base_address", "Liège, Belgique")
            )
            dest = resolve_destination_text(row)

            km, mn = ors_route_km_min(origin, dest)
            if km is not None and mn is not None:
                update_planning_row(
                    rid,
                    {
                        "KM_EST": str(km),
                        "TEMPS_EST": str(mn),
                    }
                )

        # ✅ CES LIGNES DOIVENT ÊTRE ICI
        st.session_state["km_time_run"] = False
        st.success("KM et temps calculés et sauvegardés ✅")
        st.rerun()

  
        # 🔒 IMPORTANT : couper le flag AVANT rerun
        st.session_state["km_time_run"] = False
        st.session_state["km_time_last_mode"] = None

        st.success("KM et temps calculés et sauvegardés ✅")

        # rerun propre (une seule fois)
        st.experimental_rerun()




    selected_indices = edited.index[edited["_SELECT"] == True].tolist()
    if selected_indices:
        selected_idx = selected_indices[-1]
    else:
        selected_idx = 0

    selected_id = int(ids[selected_idx])
    base_row = get_row_by_id(selected_id)
    if base_row is None:
        st.error("Navette modèle introuvable.")
        return

    st.markdown("### 📝 Créer / modifier à partir du modèle")

    cols_names = get_planning_columns()
    cols_names = cols_names[:40]

    new_values: Dict[str, Any] = {}
    cL, cR = st.columns(2)
    today = date.today()

    for i, col_name in enumerate(cols_names):
        cont = cL if i % 2 == 0 else cR
        val = base_row.get(col_name)

        # DATE
        if col_name == "DATE":
            default_date = today
            if isinstance(val, str) and val:
                try:
                    default_date = datetime.strptime(val, "%d/%m/%Y").date()
                except Exception:
                    pass
            new_d = cont.date_input(
                "DATE",
                value=default_date,
                key=f"client_DATE_{selected_id}",
            )
            new_values[col_name] = new_d.strftime("%d/%m/%Y")
            continue

        # GROUPAGE / PARTAGE
        if col_name in ["GROUPAGE", "PARTAGE"]:
            b = cont.checkbox(
                "Groupage" if col_name == "GROUPAGE" else "Navette partagée",
                value=bool_from_flag(val),
                key=f"client_{col_name}_{selected_id}",
            )
            new_values[col_name] = "1" if b else "0"
            continue

        # GO
        if col_name == "GO":
            txt = "" if val is None else str(val)
            t2 = cont.text_input(
                "GO (AL / GO / GL)",
                value=txt,
                key=f"client_GO_{selected_id}",
            )
            new_values[col_name] = t2.strip().upper()
            continue

        # HEURE
        if col_name == "HEURE":
            txt = "" if val is None else str(val)
            t2 = cont.text_input(
                "HEURE",
                value=txt,
                key=f"client_HEURE_{selected_id}",
            )
            new_values[col_name] = normalize_time_string(t2)
            continue

        # HEURE FIN (²²²²)
        if col_name == "²²²²":
            txt = "" if val is None else str(val)
            t2 = cont.text_input(
                "Heure fin (²²²²)",
                value=txt,
                key=f"client_2222_{selected_id}",
            )
            new_values[col_name] = normalize_time_string(t2)
            continue

        txt = "" if val is None or str(val).lower() == "nan" else str(val)
        t2 = cont.text_input(col_name, value=txt, key=f"client_{col_name}_{selected_id}")
        new_values[col_name] = t2

    role = st.session_state.role

    c1, c2 = st.columns(2)

    with c1:
        if st.button("➕ Créer une nouvelle navette pour ce client"):
            if role_allows_go_gl_only() and not leon_allowed_for_row(new_values.get("GO")):
                st.error("Utilisateur 'leon' : création autorisée uniquement pour GO / GL.")
            else:
                insert_planning_row(new_values)
                st.success("Nouvelle navette créée.")
                st.rerun()

    with c2:
        if st.button("✅ Mettre à jour la navette existante"):
            if role_allows_go_gl_only() and not leon_allowed_for_row(base_row.get("GO")):
                st.error("Utilisateur 'leon' : modification autorisée uniquement pour GO / GL.")
            else:
                update_planning_row(selected_id, new_values)
                st.success("Navette mise à jour.")
                st.rerun()

    st.markdown("---")
    st.markdown("### 🔁 Créer un RETOUR à partir de ce modèle")

    retour_data = new_values.copy()
    colR1, colR2 = st.columns(2)
    with colR1:
        retour_date = st.date_input(
            "Date du RETOUR",
            value=today,
            key=f"client_retour_DATE_{selected_id}",
        )
    with colR2:
        retour_heure = st.text_input(
            "Heure du RETOUR",
            value="",
            key=f"client_retour_HEURE_{selected_id}",
        )

    retour_data["DATE"] = retour_date.strftime("%d/%m/%Y")
    if "HEURE" in retour_data:
        retour_data["HEURE"] = normalize_time_string(retour_heure)

    if st.button("📋 Créer un RETOUR (copie modifiable)"):
        if role_allows_go_gl_only() and not leon_allowed_for_row(retour_data.get("GO")):
            st.error("Utilisateur 'leon' : création autorisée uniquement pour GO / GL.")
        else:
            insert_planning_row(retour_data)
            st.success("Navette RETOUR créée.")
            st.rerun()

# ============================================================
#   OUTILS CHAUFFEURS — CONTACTS, STATS, TRI
# ============================================================

def get_chauffeur_contact(ch: str):
    """Récupère téléphone + mail du chauffeur via table `chauffeurs` (Feuil2)."""
    tel = ""
    mail = ""
    try:
        with get_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT * FROM chauffeurs WHERE TRIM(INITIALE) = ? LIMIT 1", (ch,))
            row = cur.fetchone()
            if row:
                cols = [d[0] for d in cur.description]
                data = {cols[i]: row[i] for i in range(len(cols))}
                tel = (
                    data.get("TEL_CH")
                    or data.get("TEL")
                    or data.get("Tél")
                    or data.get("PHONE")
                    or ""
                )
                mail = data.get("MAIL") or data.get("Email") or ""
    except Exception:
        pass
    return str(tel or ""), str(mail or "")


def render_chauffeur_stats(df_ch: pd.DataFrame):
    """Affiche navettes / PAX / caisse pour un chauffeur."""
    if df_ch is None or df_ch.empty:
        return

    cols = df_ch.columns
    mask_course = ~df_ch.apply(lambda r: is_indispo_row(r, cols), axis=1)
    df_course = df_ch[mask_course].copy()

    nb_nav = len(df_course)
    pax_total = pd.to_numeric(df_course.get("PAX", 0), errors="coerce").fillna(0).sum()
    caisse_total = pd.to_numeric(df_course.get("Caisse", 0), errors="coerce").fillna(0).sum()

    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("🚐 Navettes (hors indispo)", int(nb_nav))
    with c2:
        st.metric("👥 PAX total", int(pax_total))
    with c3:
        st.metric("💶 Caisse totale", float(caisse_total))
import re

def match_ch_for_mail(cell, ch):
    """
    Retourne True si le chauffeur ch doit recevoir la ligne CH.
    Gère NPFA, FANP, NP*FA, FA-NP, etc.
    """
    if not cell or not ch:
        return False

    s = str(cell).upper()

    # Normalisation
    s = (
        s.replace(" ", "")
         .replace("*", "")
         .replace("/", "")
         .replace("-", "")
         .replace(",", "")
    )

    # Découpage en blocs de 2 lettres
    parts = re.findall(r"[A-Z]{2}", s)

    return ch in parts

# ============================================================
#   ENVOI PLANNING AUX CHAUFFEURS (MAIL + WHATSAPP)
# ============================================================

def send_planning_to_chauffeurs(
    chauffeurs: list[str],
    from_date: date,
    to_date: date | None = None,
    message_type: str = "planning",
):
    """
    Envoie à chaque chauffeur un mail avec SON planning individuel
    et prépare les liens WhatsApp.

    ⚠️ Logique chauffeur STRICTEMENT IDENTIQUE à la vue chauffeur :
    - NPFA / FANP / NP*FA / DOFA → NP et FA reçoivent
    """

    if not chauffeurs:
        st.warning("Aucun chauffeur sélectionné.")
        return

    sent = 0
    no_email: list[str] = []
    wa_links: list[dict] = []

    # ===================================================
    # 🔍 Chargement planning (UNE SEULE FOIS, SANS FILTRE CH)
    # ===================================================
    df_all = get_planning(
        start_date=from_date,
        end_date=to_date,
        chauffeur=None,          # ⚠️ IMPORTANT
        type_filter=None,
        search="",
        max_rows=5000,
        source="full",
    )

    if df_all is None or df_all.empty:
        st.warning("Aucune navette sur la période sélectionnée.")
        return

    # ===================================================
    # 📧 BOUCLE CHAUFFEURS
    # ===================================================
    for ch in chauffeurs:

        ch = str(ch).strip().upper()
        if not ch:
            continue

        tel, mail = get_chauffeur_contact(ch)

        # ===================================================
        # ⚡ FILTRAGE CHAUFFEUR (COPIÉ DE LA VUE CHAUFFEUR)
        # ===================================================
        ch_series = (
            df_all["CH"]
            .fillna("")
            .astype(str)
            .str.upper()
            .str.strip()
        )

        mask_exact = ch_series == ch
        mask_star = ch_series == f"{ch}*"
        mask_contains = ch_series.str.contains(ch, regex=False)
        mask_not_digit_suffix = ~ch_series.str.match(rf"{ch}\d")

        df_ch = df_all[
            (mask_exact | mask_star | mask_contains) & mask_not_digit_suffix
        ].copy()

        if df_ch.empty:
            continue

        # 🔒 Sécurité anti-mails énormes
        if len(df_ch) > 400:
            st.warning(
                f"⚠️ {ch} : trop de lignes ({len(df_ch)}) — envoi ignoré."
            )
            continue

        # ===================================================
        # 📧 CONSTRUCTION DU MAIL
        # ===================================================
        if message_type == "planning":
            subject = f"🚖 Planning — {ch} ({from_date.strftime('%d/%m/%Y')})"
            msg_txt = build_planning_mail_body(
                df_ch=df_ch,
                ch=ch,
                from_date=from_date,
                to_date=to_date,
            )
        else:
            subject = f"📢 Modification planning — {ch}"
            msg_txt = (
                "Bonjour,\n\n"
                "📢 Une modification de planning a été effectuée aujourd’hui.\n"
                "Merci de consulter l’application Airports Lines "
                "et de confirmer la réception.\n\n"
                "— Airports Lines"
            )

        # ===================================================
        # 📧 ENVOI EMAIL
        # ===================================================
        if mail:
            if send_email_smtp(mail, subject, msg_txt):
                sent += 1
        else:
            no_email.append(ch)

        # ===================================================
        # 💬 LIEN WHATSAPP
        # ===================================================
        if tel:
            wa_msg = build_chauffeur_new_planning_message(ch, from_date)
            wa_url = build_whatsapp_link(tel, wa_msg)
            wa_links.append({
                "ch": ch,
                "tel": tel,
                "url": wa_url,
            })

    # ===================================================
    # 📊 RETOUR UI
    # ===================================================
    st.success(f"📧 Emails envoyés pour {sent} chauffeur(s).")

    if no_email:
        st.info(
            "📭 Pas d'adresse email configurée pour : "
            + ", ".join(sorted(set(no_email)))
        )

    if wa_links:
        st.markdown("### 💬 Prévenir les chauffeurs par WhatsApp")
        st.caption("Clique sur un lien pour ouvrir WhatsApp avec le message pré-rempli.")

        for item in wa_links:
            st.markdown(
                f"- {item['ch']} ({item['tel']}) → "
                f"[Envoyer WhatsApp]({item['url']})"
            )



def _sort_df_by_date_heure(df: pd.DataFrame) -> pd.DataFrame:
    """Tri par DATE + HEURE (normalisée)."""
    df = df.copy()

    if "DATE" in df.columns:
        try:
            df["DATE_SORT"] = pd.to_datetime(df["DATE"], errors="coerce")
        except Exception:
            df["DATE_SORT"] = pd.NaT
    else:
        df["DATE_SORT"] = pd.NaT

    if "HEURE" in df.columns:
        def _hs(h):
            h = normalize_time_string(h)
            if not h:
                return (99, 99)
            try:
                parts = h.split(":")
                if len(parts) != 2:
                    return (99, 99)
                return (int(parts[0]), int(parts[1]))
            except Exception:
                return (99, 99)
        df["HEURE_SORT"] = df["HEURE"].apply(_hs)
    else:
        df["HEURE_SORT"] = (99, 99)

    df = df.sort_values(["DATE_SORT", "HEURE_SORT"]).drop(
        columns=["DATE_SORT", "HEURE_SORT"],
        errors="ignore",
    )
    return df


def build_chauffeur_future_message(df: pd.DataFrame, ch_selected: str, from_date: date) -> str:
    lines: List[str] = []
    lines.append(f"🚖 Planning à partir du {from_date.strftime('%d/%m/%Y')} — Chauffeur : {ch_selected}")
    lines.append("")

    df = df.copy()
    if "DATE" in df.columns:
        df["DATE"] = pd.to_datetime(df["DATE"], errors="coerce").dt.date
        df = df[df["DATE"].notna() & (df["DATE"] >= from_date)]

    if df.empty:
        lines.append("Aucune navette planifiée.")
        return "\n".join(lines)

    df = df[df["CH"].astype(str).str.upper() == ch_selected.upper()]
    if df.empty:
        lines.append("Aucune navette pour ce chauffeur.")
        return "\n".join(lines)

    df = _sort_df_by_date_heure(df)
    cols = df.columns.tolist()

    for d, sub in df.groupby("DATE"):
        lines.append(f"📆 {d.strftime('%d/%m/%Y')}")

        for _, row in sub.iterrows():

            if is_indispo_row(row, cols):
                h1 = normalize_time_string(row.get("HEURE"))
                h2 = normalize_time_string(row.get("²²²²"))
                lines.append(f"  ⏱ {h1 or '??:??'} → {h2 or '??:??'} — 🚫 Indisponible")
                lines.append("")
                continue

            heure = normalize_time_string(row.get("HEURE")) or "??:??"

            sens_txt = format_sens_ar(row.get("Unnamed: 8"))
            dest = resolve_client_alias(str(row.get("DESIGNATION", "") or "").strip())
            sens_dest = f"{sens_txt} ({dest})" if sens_txt and dest else dest or sens_txt or "Navette"

            nom = str(row.get("NOM", "") or "").strip()

            lines.append(f"  ➡ {heure} — {sens_dest} — {nom}")

            adr = build_full_address_from_row(row)
            if adr:
                lines.append(f"     📍 {adr}")

            extras = []
            if row.get("PAX"):
                extras.append(f"{row.get('PAX')} pax")

            paiement = str(row.get("PAIEMENT", "") or "").lower()
            caisse = row.get("Caisse")
            if paiement == "facture":
                extras.append("Facture")
            elif paiement in ("caisse", "bancontact"):
                extras.append(f"{paiement} {caisse}€" if caisse else paiement)

            if extras:
                lines.append("     " + " — ".join(extras))

            lines.append("")
        lines.append("")

    return "\n".join(lines).strip()


def build_chauffeur_new_planning_message(ch: str, from_date: date) -> str:
    """
    Petit message WhatsApp pour dire au chauffeur qu'il a un nouveau planning.
    """
    d_txt = from_date.strftime("%d/%m/%Y")
    return (
        f"Bonjour {ch}, c'est Airports-Lines.\n"
        f"Ton planning a été mis à jour à partir du {d_txt}.\n"
        f"Les courses modifiées sont indiquées dans ta vue chauffeur.\n\n"
        f"Merci de te connecter à l'application et de cliquer sur "
        f"« J'ai bien reçu mon planning » pour confirmer. 👍"
    )

def build_chauffeur_day_message(df_ch: pd.DataFrame, ch_selected: str, day_label: str) -> str:
    cols = df_ch.columns.tolist()
    lines = []

    lines.append(f"🚖 Planning à partir du {day_label} — Chauffeur : {ch_selected}")
    lines.append("")

    for _, row in df_ch.iterrows():

        if is_indispo_row(row, cols):
            h1 = normalize_time_string(row.get("HEURE")) or "??:??"
            h2 = normalize_time_string(row.get("²²²²")) or "??:??"
            lines.append(f"⏱ {h1} → {h2} — 🚫 Indisponible")
            lines.append("")
            continue

        heure = normalize_time_string(row.get("HEURE")) or "??:??"

        sens_txt = format_sens_ar(row.get("Unnamed: 8"))
        dest = resolve_client_alias(resolve_destination_text(row))
        sens_dest = f"{sens_txt} ({dest})" if sens_txt and dest else dest or sens_txt or "Navette"

        nom = str(row.get("NOM", "") or "").strip()
        lines.append(f"  ➡ {heure} — {sens_dest} — {nom}")

        adr = build_full_address_from_row(row)
        if adr:
            lines.append(f"     📍 {adr}")

        tel = get_client_phone_from_row(row)
        if tel:
            lines.append(f"     📞 Client : {tel}")

        vol = extract_vol_val(row, cols)
        if vol:
            lines.append(f"     ✈️ Vol : {vol}")

        extras = []
        if row.get("PAX"):
            extras.append(f"{row.get('PAX')} pax")

        paiement = str(row.get("PAIEMENT", "") or "").lower()
        caisse = row.get("Caisse")
        if paiement == "facture":
            extras.append("Facture")
        elif paiement in ("caisse", "bancontact"):
            extras.append(f"{paiement} {caisse}€" if caisse else paiement)

        if extras:
            lines.append("     " + " — ".join(extras))

        lines.append("")

    return "\n".join(lines).strip()

# ============================================================
#   ONGLET 🚖 VUE CHAUFFEUR (PC + GSM)
#   -> DEVENU : ENVOI PLANNING BUREAU (OPTIMISÉ)
# ============================================================

def render_tab_vue_chauffeur(forced_ch=None):
    import pandas as pd
    from datetime import date, timedelta

    st.subheader("📢 Bureau — Envoi planning chauffeurs")

    # =======================================================
    # 🔐 Accès réservé
    # =======================================================
    if st.session_state.get("role") != "admin":
        st.info("Cette page sert uniquement au **bureau** pour envoyer le planning.")
        return

    today = date.today()

    # =======================================================
    # 🧱 Init DB (1x)
    # =======================================================
    if not st.session_state.get("send_log_init_done"):
        ensure_send_log_table()
        st.session_state["send_log_init_done"] = True

    # =======================================================
    # 📅 PÉRIODE
    # =======================================================
    periode = st.radio(
        "📅 Quelle période envoyer ?",
        ["Aujourd’hui", "Demain + 2 jours"],
        horizontal=True,
        key="bureau_send_periode",
    )

    if periode == "Aujourd’hui":
        d_start = today
        d_end = today
        periode_label = "du jour"
    else:
        d_start = today + timedelta(days=1)
        d_end = today + timedelta(days=3)
        periode_label = "de demain à J+3"

    # =======================================================
    # 🧾 FORMAT D'ENVOI
    # =======================================================
    send_format = st.radio(
        "🧾 Format d'envoi par e-mail",
        ["Normal (texte)", "PDF (texte + pièce jointe)"],
        horizontal=True,
        key="bureau_send_format",
        help="Normal = identique à avant. PDF = même message + planning en PDF (comme dans la vue Driver).",
    )
    want_pdf = str(send_format).startswith("PDF")

    # =======================================================
    # 🚖 CHAUFFEURS SUR LA PÉRIODE (SQL DIRECT -> fiable)
    # =======================================================
    active_chauffeurs = set()
    with get_connection() as conn:
        df_chcol = pd.read_sql_query(
            """
            SELECT CH
            FROM planning
            WHERE COALESCE(IS_INDISPO,0)=0
              AND COALESCE(IS_SUPERSEDED,0)=0
              AND DATE_ISO BETWEEN ? AND ?
              AND COALESCE(CH,'') <> ''
            """,
            conn,
            params=(d_start.isoformat(), d_end.isoformat()),
        )

    if not df_chcol.empty:
        for raw in df_chcol["CH"].dropna().astype(str):
            for c in split_chauffeurs(raw):
                if c:
                    active_chauffeurs.add(c)

    chauffeurs_planning = sorted(active_chauffeurs)

    # forced_ch (si tu veux forcer un chauffeur depuis un autre écran)
    if forced_ch:
        forced = str(forced_ch).strip().upper()
        if forced and forced not in chauffeurs_planning:
            chauffeurs_planning = [forced] + chauffeurs_planning

    if not chauffeurs_planning:
        st.warning("Aucun chauffeur trouvé sur la période sélectionnée.")
        return

    # =======================================================
    # 🎯 DESTINATAIRES
    # =======================================================
    ch_choice = st.radio(
        "🚖 Destinataire",
        ["Tous les chauffeurs", "Un chauffeur"],
        horizontal=True,
        key="bureau_send_target",
    )

    if ch_choice == "Un chauffeur":
        multi_mode = st.checkbox("☑️ Sélection multiple de chauffeurs", key="bureau_send_multi")
        if multi_mode:
            target_chauffeurs = st.multiselect(
                "Sélectionner les chauffeurs",
                chauffeurs_planning,
                key="bureau_send_multiselect",
            )
        else:
            one = st.selectbox(
                "Sélectionner le chauffeur",
                chauffeurs_planning,
                key="bureau_send_one",
            )
            target_chauffeurs = [one] if one else []
    else:
        target_chauffeurs = chauffeurs_planning

    target_chauffeurs = [c for c in target_chauffeurs if c]

    if not target_chauffeurs:
        st.warning("Aucun chauffeur sélectionné.")
        return

    col_mail, col_wa = st.columns(2)

    # ===========================
    # 📧 MAIL (TEXTE + PDF optionnel)
    # ===========================
    with col_mail:
        if st.button(
            "📧 Envoyer le planning",
            use_container_width=True,
            key="bureau_send_mail_btn",
        ):
            errors = []
            sent = 0

            # anti-doublon chauffeurs
            for ch in dict.fromkeys(target_chauffeurs):

                try:
                    # -------------------
                    # Contact chauffeur
                    # -------------------
                    _tel, mail = get_chauffeur_contact(ch)
                    if not mail:
                        raise ValueError("Email manquant")

                    # -------------------
                    # Planning chauffeur (source brute)
                    # -------------------
                    df_ch = get_chauffeur_planning(
                        chauffeur=ch,
                        from_date=d_start,
                        to_date=d_end,
                    )

                    if df_ch is None or df_ch.empty:
                        log_send(
                            ch,
                            "MAIL",
                            periode_label,
                            "OK",
                            "Aucune navette (pas d'envoi)",
                        )
                        continue

                    # -------------------
                    # Corps du mail (identique à avant)
                    # -------------------
                    body = build_planning_mail_body(
                        df_ch=df_ch,
                        ch=ch,
                        from_date=d_start,
                        to_date=d_end,
                    )

                    subject = f"[PLANNING] {ch} — {periode_label}"

                    # ===================================================
                    # 📊 CONSTRUIRE LE MÊME TABLEAU QUE LA VUE DRIVER
                    # ===================================================
                    planning_cols_driver = [
                        "DATE","HEURE","CH","²²²²","IMMAT","PAX","Reh","Siège",
                        "Unnamed: 8","DESIGNATION","H South","Décollage","N° Vol","Origine",
                        "GO","Num BDC","NOM","ADRESSE","CP","Localité","Tél",
                        "Type Nav","PAIEMENT","Caisse"
                    ]

                    df_table = df_ch.copy()

                    # sécurité colonnes
                    for c in planning_cols_driver:
                        if c not in df_table.columns:
                            df_table[c] = ""

                    df_table = df_table[planning_cols_driver]

                    # -------------------
                    # Envoi mail (PDF ou non)
                    # -------------------
                    if want_pdf:
                        pdf_buf = export_chauffeur_planning_table_pdf(
                            df_table,
                            ch,
                        )

                        pdf_bytes = (
                            pdf_buf.getvalue()
                            if hasattr(pdf_buf, "getvalue")
                            else bytes(pdf_buf)
                        )

                        fname = (
                            f"Planning_{ch}_"
                            f"{d_start.isoformat()}_"
                            f"{(d_end or d_start).isoformat()}.pdf"
                        )

                        ok = send_email_smtp_with_attachments(
                            to_email=mail,
                            subject=subject,
                            body=body,
                            attachments=[(fname, pdf_bytes, "pdf")],
                        )
                    else:
                        ok = send_email_smtp(
                            to_email=mail,
                            subject=subject,
                            body=body,
                        )

                    if not ok:
                        raise RuntimeError("SMTP : envoi échoué")

                    sent += 1
                    log_send(ch, "MAIL", periode_label, "OK", "Envoyé")

                except Exception as e:
                    log_send(ch, "MAIL", periode_label, "ERREUR", str(e))
                    errors.append((ch, str(e)))

            # -------------------
            # Résultat UI
            # -------------------
            if errors:
                st.error("❌ Certains envois ont échoué.")
                with st.expander("Voir le détail"):
                    for ch, err in errors:
                        st.write(f"- {ch} : {err}")
            else:
                st.success(
                    f"✅ Planning {periode_label} envoyé ({sent} chauffeur(s))."
                )


    # ===========================
    # 💬 WHATSAPP (LIENS)
    # ===========================
    with col_wa:
        if st.button("💬 Envoyer par WhatsApp", use_container_width=True, key="bureau_send_wa_btn"):
            wa_links = []

            for ch in dict.fromkeys(target_chauffeurs):
                tel, _mail = get_chauffeur_contact(ch)
                if not tel:
                    continue

                df_ch = get_chauffeur_planning(
                    chauffeur=ch,
                    from_date=d_start,
                    to_date=d_end,
                )

                if df_ch is None or df_ch.empty:
                    continue

                wa_text = build_planning_mail_body(
                    df_ch=df_ch,
                    ch=ch,
                    from_date=d_start,
                    to_date=d_end,
                )

                wa_links.append({
                    "ch": ch,
                    "tel": tel,
                    "url": build_whatsapp_link(tel, wa_text),
                })

            if not wa_links:
                st.warning("Aucun numéro WhatsApp disponible (ou planning vide).")
            else:
                st.markdown("### 💬 Envoi WhatsApp")
                for item in wa_links:
                    st.markdown(
                        f"- **{item['ch']}** ({item['tel']}) → "
                        f"[📲 Ouvrir WhatsApp]({item['url']})"
                    )

    with get_connection() as conn:
        df_log = pd.read_sql_query(
            """
            SELECT ts, chauffeur, canal, periode, statut, message
            FROM send_log
            ORDER BY ts DESC
            LIMIT 100
            """,
            conn,
        )

    st.markdown("---")
    st.markdown("### 🧾 Historique (100 derniers)")
    st.dataframe(df_log, use_container_width=True, height=260)


def export_chauffeur_planning_pdf(df_ch: pd.DataFrame, ch: str):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)

    margin_x = 1.5 * cm
    y = height - 2 * cm
    line_h = 0.55 * cm

    def new_page():
        nonlocal y
        c.showPage()
        y = height - 2 * cm
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin_x, y, f"Planning chauffeur — {ch}")
        y -= 0.9 * cm
        c.setFont("Helvetica", 10)

    # En-tête page 1
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin_x, y, f"Planning chauffeur — {ch}")
    y -= 0.9 * cm
    c.setFont("Helvetica", 10)

    cols = df_ch.columns.tolist()

    def write_line(txt: str, indent: float = 0.0, bold: bool = False):
        nonlocal y
        if y < 2 * cm:
            new_page()

        c.setFont("Helvetica-Bold" if bold else "Helvetica", 10)
        c.drawString(margin_x + indent, y, txt[:120])
        y -= line_h

    for _, row in df_ch.iterrows():

        # --- Date ---
        dv = row.get("DATE")
        if isinstance(dv, date):
            date_txt = dv.strftime("%d/%m/%Y")
        else:
            try:
                dt = pd.to_datetime(dv, dayfirst=True, errors="coerce")
                date_txt = dt.strftime("%d/%m/%Y") if not pd.isna(dt) else "??/??/????"
            except Exception:
                date_txt = "??/??/????"

        # --- Heure ---
        heure = normalize_time_string(row.get("HEURE")) or "??:??"

        # --- Sens + destination ---
        sens_txt = format_sens_ar(row.get("Unnamed: 8"))
        lieu = resolve_client_alias(str(row.get("DESIGNATION", "") or "").strip())
        sens_dest = f"{sens_txt} ({lieu})" if sens_txt and lieu else (lieu or sens_txt or "Navette")

        # --- Client / tel / adresse ---
        nom = str(row.get("NOM", "") or "").strip()
        tel_client = get_client_phone_from_row(row)
        adr_full = build_full_address_from_row(row)

        # --- NUMÉRO DE BDC (ROBUSTE) ---
        num_bdc = ""
        for cand in ["NUM BDC", "Num BDC", "NUM_BDC", "BDC"]:
            if cand in cols and row.get(cand):
                num_bdc = str(row.get(cand)).strip()
                break

        # --- Véhicule (SIÈGE BÉBÉ / RÉHAUSSEUR) ---
        immat = str(row.get("IMMAT", "") or "").strip()

        # 🍼 Siège bébé (SIEGE / SIÈGE)
        siege_bebe = extract_positive_int(row.get("SIEGE", row.get("SIÈGE")))

        # 🪑 Rehausseur
        reh_n = extract_positive_int(row.get("REH"))


        # --- Paiement / caisse / pax ---
        pax = row.get("PAX")
        paiement = str(row.get("PAIEMENT", "") or "").lower()
        caisse = row.get("Caisse")

        # --- Vol ---
        vol = extract_vol_val(row, cols)

        # --- GO ---
        go_val = str(row.get("GO", "") or "").strip()

        # =======================
        # Impression bloc navette
        # =======================
        write_line(f"📆 {date_txt} | ⏱ {heure} — {sens_dest}", bold=True)

        if nom:
            write_line(f"👤 Client : {nom}", indent=10)

        if num_bdc:
            write_line(f"🧾 BDC : {num_bdc}", indent=10)

        if tel_client:
            write_line(f"📞 Client : {tel_client}", indent=10)

        if adr_full:
            write_line(f"📍 Adresse : {adr_full}", indent=10)

        veh_infos = []

        if immat:
            veh_infos.append(f"Plaque {immat}")

        if siege_bebe:
            veh_infos.append(f"🍼 Siège bébé {siege_bebe}")

        if reh_n:
            veh_infos.append(f"🪑 Rehausseur {reh_n}")

        if veh_infos:
            write_line("🚘 " + " | ".join(veh_infos), indent=10)


        extra = []
        if vol:
            extra.append(f"✈️ {vol}")
        if pax:
            extra.append(f"👥 {pax} pax")

        if paiement == "facture":
            extra.append("🧾 Facture")
        elif paiement in ("caisse", "bancontact"):
            if caisse not in ("", None):
                extra.append(f"💶 {caisse} € ({paiement})")
            else:
                extra.append(f"💶 {paiement}")

        if extra:
            write_line(" — ".join(extra), indent=10)

        if go_val:
            write_line(f"🟢 GO : {go_val}", indent=10)

        write_line("")

    c.save()
    buffer.seek(0)
    return buffer



    # =======================================================
    #   ENVOI DE CONFIRMATION (NAVETTES REMPLIES UNIQUEMENT)
    # =======================================================
    st.markdown("### ✅ Envoyer mes informations au bureau")

    recap_lines = []
    nb_remplies = 0

    for _, row in df_ch.iterrows():
        nav_id = row.get("id")

        trajet = st.session_state.get(f"trajet_nav_{nav_id}", "").strip()
        probleme = st.session_state.get(f"prob_nav_{nav_id}", "").strip()

        # on ignore totalement les navettes vides
        if not trajet and not probleme:
            continue

        nb_remplies += 1

        recap_lines.append(
            format_navette_ack(
                row=row,
                ch_selected=ch_selected,
                trajet=trajet,
                probleme=probleme,
            )
        )

    if nb_remplies == 0:
        st.warning(
            "ℹ️ Aucune information encodée. "
            "Merci de compléter au moins une navette avant l’envoi."
        )

    if st.button(
        "📤 Envoyer mes informations",
        disabled=(nb_remplies == 0),
        key=f"confirm_all_{ch_selected}_{scope}_{sel_date}",
    ):
        send_mail_admin(
            subject=f"[INFOS CHAUFFEUR] {ch_selected}",
            body="\n".join(recap_lines),
        )

        # marquer comme envoyées UNIQUEMENT les navettes remplies
        for _, row in df_ch.iterrows():
            nav_id = row.get("id")

            trajet = st.session_state.get(f"trajet_nav_{nav_id}", "").strip()
            probleme = st.session_state.get(f"prob_nav_{nav_id}", "").strip()

            if trajet or probleme:
                st.session_state[f"sent_nav_{nav_id}"] = True

        set_chauffeur_last_ack(ch_selected)

        st.success(f"✅ {nb_remplies} navette(s) envoyée(s) au bureau.")
        st.rerun()


# ============================================================
#   🚖 ONGLET CHAUFFEUR — MON PLANNING
# ============================================================





def export_chauffeur_planning_pdf_table(df_ch: pd.DataFrame, ch: str):
    """PDF en mode TABLEAU (comme une vue planning), A4 paysage, auto-fit."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)

    margin = 10  # points
    x0 = margin
    y0 = height - margin

    # Titre
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x0, y0 - 20, f"Planning chauffeur — {ch}")
    c.setFont("Helvetica", 9)

    # Colonnes à afficher (ordre préféré)
    preferred = [
        "DATE", "HEURE", "NOM", "DESIGNATION", "ADRESSE", "PAX", "VOL", "GO", "PAIEMENT", "Caisse"
    ]
    cols = [c for c in preferred if c in df_ch.columns]
    # Ajoute d'autres colonnes utiles si présentes
    for extra in ["IMMAT", "REMARQUE", "Num BDC", "NUM_BDC", "NUM BDC"]:
        if extra in df_ch.columns and extra not in cols:
            cols.append(extra)

    if not cols:
        cols = df_ch.columns.tolist()[:10]

    # Prépare les données string
    dfp = df_ch.copy()
    # Normalise date
    if "DATE" in dfp.columns:
        try:
            dfp["DATE"] = dfp["DATE"].apply(lambda v: v.strftime("%d/%m/%Y") if hasattr(v, "strftime") else str(v))
        except Exception:
            pass
    # Limite longueur cellules
    def _cell(v):
        s = "" if v is None else str(v)
        s = s.replace("\n", " ").strip()
        return s[:40]

    data = [[_cell(v) for v in row] for row in dfp[cols].values.tolist()]

    # Dimensions tableau
    top = y0 - 40
    bottom = margin + 20
    table_height = top - bottom
    table_width = width - 2 * margin

    n_rows = len(data) + 1  # header
    n_cols = len(cols)

    # Taille police auto
    font_size = 9
    row_h = max(12, table_height / max(n_rows, 1))
    if row_h < 10:
        font_size = 7
        row_h = 10
    elif row_h > 16:
        row_h = 16

    # Largeurs colonnes auto (basées sur longueur)
    lens = []
    for j, col in enumerate(cols):
        mx = len(str(col))
        for i in range(min(len(data), 80)):
            mx = max(mx, len(str(data[i][j])))
        lens.append(mx)

    total = sum(lens) if sum(lens) > 0 else n_cols
    col_w = [table_width * (l / total) for l in lens]

    # Dessin header
    y = top
    c.setFont("Helvetica-Bold", font_size)
    x = x0
    for j, col in enumerate(cols):
        c.rect(x, y - row_h, col_w[j], row_h, stroke=1, fill=0)
        c.drawString(x + 2, y - row_h + 3, str(col)[:25])
        x += col_w[j]
    y -= row_h

    # Lignes
    c.setFont("Helvetica", font_size)
    for r in data:
        if y - row_h < bottom:
            c.showPage()
            c.setFont("Helvetica-Bold", 14)
            c.drawString(x0, height - margin - 20, f"Planning chauffeur — {ch}")
            c.setFont("Helvetica", 9)
            y = height - margin - 40

            # Redessine header
            c.setFont("Helvetica-Bold", font_size)
            x = x0
            for j, col in enumerate(cols):
                c.rect(x, y - row_h, col_w[j], row_h, stroke=1, fill=0)
                c.drawString(x + 2, y - row_h + 3, str(col)[:25])
                x += col_w[j]
            y -= row_h
            c.setFont("Helvetica", font_size)

        x = x0
        for j, val in enumerate(r):
            c.rect(x, y - row_h, col_w[j], row_h, stroke=1, fill=0)
            c.drawString(x + 2, y - row_h + 3, str(val)[:60])
            x += col_w[j]
        y -= row_h

    c.save()
    buffer.seek(0)
    return buffer


def generate_urgence_mission_pdf_bytes(row: dict) -> bytes:
    """Génère un PDF (A4) pour une mission urgente (1 navette)."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)

    x = 2 * cm
    y = height - 2 * cm

    def line(txt, dy=0.7*cm, bold=False):
        nonlocal y
        if bold:
            c.setFont("Helvetica-Bold", 12)
        else:
            c.setFont("Helvetica", 11)
        c.drawString(x, y, txt)
        y -= dy

    line("AIRPORTS LINES — MISSION URGENTE", bold=True, dy=1.0*cm)
    line(f"Date : {row.get('DATE','')}")
    line(f"Heure : {row.get('HEURE','')}")
    line(f"Chauffeur : {row.get('CH','')}")
    line(f"PAX : {row.get('PAX','')}")
    line(f"Vol : {row.get('VOL','')}")
    line(f"Destination : {row.get('DESIGNATION','')}")
    line(f"Adresse pick-up : {row.get('ADRESSE','')}")
    line(f"Client : {row.get('NOM','')}")
    line(f"Tél : {row.get('Tél', row.get('TEL',''))}")
    line(f"Paiement : {row.get('PAIEMENT','')}")
    line(f"BDC : {row.get('Num BDC', row.get('BDC',''))}")
    line(f"Remarque : {row.get('REMARQUE','')}", dy=1.2*cm)

    c.setFont("Helvetica", 9)
    c.drawString(x, 1.5*cm, f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def format_urgence_message(row: dict) -> str:
    date_txt = row.get("DATE") or row.get("DATE_ISO") or ""
    heure_txt = row.get("HEURE") or ""
    dest = row.get("DESIGNATION") or ""
    nom = row.get("NOM") or ""
    adr = row.get("ADRESSE") or ""
    pax = row.get("PAX") or ""
    vol = row.get("VOL") or row.get("N° Vol") or ""
    bdc = row.get("Num BDC") or row.get("BDC") or ""
    return (
        "🚨 URGENCE — Nouvelle mission\n"
        f"📅 {date_txt} à {heure_txt}\n"
        f"➡️ {dest}\n"
        + (f"👤 Client : {nom}\n" if nom else "")
        + (f"📍 Adresse : {adr}\n" if adr else "")
        + (f"🧳 PAX : {pax}\n" if pax else "")
        + (f"✈️ Vol : {vol}\n" if vol else "")
        + (f"🧾 BDC : {bdc}\n" if bdc else "")
        + "Merci de confirmer immédiatement ✅"
    )


def notify_chauffeur_urgence(row: dict) -> dict:
    """Notifie le(s) chauffeur(s) (mail + WhatsApp link) et log en DB."""
    ch_code = (row.get("CH") or "").strip()
    if not ch_code:
        return {"ok": False, "error": "CH vide"}

    msg = format_urgence_message(row)

    # Email
    emails_sent, emails_missing = send_email_to_chauffeurs_from_row(
        row=row,
        subject="🚨 URGENCE — Nouvelle mission",
        body=msg,
    )

    # WhatsApp links
    wa_links = []
    for ch in split_chauffeurs(ch_code):
        tel, _ = get_chauffeur_contact(ch)
        if tel:
            wa_links.append((ch, build_whatsapp_link(tel, msg)))

    # Log interne chauffeur_messages
    try:
        for ch in split_chauffeurs(ch_code):
            receive_chauffeur_planning(chauffeur=ch, texte=msg, canal="URGENCE")
    except Exception:
        pass

    # DB meta
    try:
        set_urgence_status(
            int(row.get("id")),
            status="EN_COURS",
            notified_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            channel="MAIL+WA",
        )
    except Exception:
        pass

    return {
        "ok": True,
        "emails_sent": emails_sent,
        "emails_missing": emails_missing,
        "wa_links": wa_links,
        "message": msg,
    }



def render_tab_urgences_admin():
    st.subheader("🚨 Urgences")

    # 🔁 Rerun contrôlé (anti-refresh brutal)
    consume_soft_refresh("urgences_admin")

    colA, colB, colC = st.columns([1, 1, 2])
    with colA:
        status = st.selectbox("Statut", ["Toutes", "EN_COURS", "TERMINEE"], index=0, key="urg_status")
    with colB:
        days_back = st.number_input("Jours à afficher", min_value=1, max_value=120, value=30, step=1, key="urg_days_back")
    with colC:
        st.caption("Les urgences = lignes planning où **URGENCE=1**.")

    df = get_urgences(
        status=None if status == "Toutes" else status,
        days_back=int(days_back),
    )

    if df is None or df.empty:
        st.success("✅ Aucune urgence")
        return

    # Colonnes utiles (évite dataframe trop large)
    cols_show = []
    for c in ["DATE", "HEURE", "CH", "IMMAT", "PAX", "DESIGNATION", "NOM", "ADRESSE", "VOL", "URGENCE_STATUS", "URGENCE_NOTIFIED_AT", "id"]:
        if c in df.columns:
            cols_show.append(c)

    st.dataframe(df[cols_show], use_container_width=True, hide_index=True, height=320)

    st.markdown("----")
    st.markdown("### Détails / actions")

    for _, row in df.iterrows():
        rid = int(row.get("id") or 0)
        title = f"{row.get('DATE','')} {row.get('HEURE','')} — {row.get('CH','')} — {row.get('NOM','')}"
        with st.expander(title, expanded=False):
            row_dict = row.to_dict()

            st.write(f"**Destination** : {row.get('DESIGNATION','')}")
            st.write(f"**Adresse** : {row.get('ADRESSE','')}")
            st.write(f"**Vol** : {row.get('VOL','')}")
            st.write(f"**Statut** : {row.get('URGENCE_STATUS','') or 'EN_COURS'}")

            # Conflits rapides (si CH présent)
            ch = (row.get("CH") or "").strip()
            date_iso = row.get("DATE_ISO")
            heure = row.get("HEURE")
            if date_iso and heure and ch:
                df_conf = find_time_conflicts(
                    date_iso=str(date_iso),
                    heure=str(heure),
                    ch=str(split_chauffeurs(ch)[0]) if split_chauffeurs(ch) else ch,
                    window_min=90,
                    exclude_id=rid,
                )
                if df_conf is not None and not df_conf.empty:
                    st.warning("⚠️ Conflit horaire détecté (même chauffeur, +/- 90 min)")
                    show_cols = [c for c in ["DATE", "HEURE", "CH", "NOM", "ADRESSE", "DESIGNATION", "VOL", "id"] if c in df_conf.columns]
                    st.dataframe(df_conf[show_cols], use_container_width=True, hide_index=True)

            c1, c2, c3, c4 = st.columns(4)

            with c1:
                if st.button("🔔 Notifier chauffeur", key=f"urg_notify_{rid}"):
                    res = notify_chauffeur_urgence(row_dict)
                    if not res.get("ok"):
                        st.error(res.get("error", "Erreur notification"))
                    else:
                        st.success("✅ Notification envoyée")
                        if res.get("wa_links"):
                            for ch_code, link in res["wa_links"]:
                                st.markdown(f"[💬 WhatsApp {ch_code}]({link})")

            with c2:
                if st.button("✅ Marquer terminée", key=f"urg_done_{rid}"):
                    try:
                        set_urgence_status(rid, status="TERMINEE")
                        st.success("OK")
                        request_soft_refresh("urgences_admin", clear_cache=True, mute_autosync_sec=10)
                    except Exception as e:
                        st.error(str(e))

            with c3:
                # ⚡ Génération PDF à la demande (évite lenteur à chaque rerun)
                pdf_key = f"_urg_pdf_bytes_{rid}"

                if st.button("⚙️ Générer PDF", key=f"urg_genpdf_{rid}"):
                    try:
                        st.session_state[pdf_key] = generate_urgence_mission_pdf_bytes(row_dict)
                        st.success("PDF prêt ✅")
                    except Exception as e:
                        st.error(f"Erreur PDF : {e}")

                pdf_bytes = st.session_state.get(pdf_key)
                if pdf_bytes:
                    st.download_button(
                        "📄 Télécharger PDF mission",
                        data=pdf_bytes,
                        file_name=f"MISSION_URGENTE_{rid}.pdf",
                        mime="application/pdf",
                        key=f"urg_pdf_{rid}",
                    )
                else:
                    st.caption("Génère le PDF avant téléchargement.")



            with c4:
                if st.button("🧹 Retirer l'urgence", key=f"urg_clear_{rid}"):
                    try:
                        update_planning_row(rid, {"URGENCE": 0, "URGENCE_STATUS": "TERMINEE"})
                        st.success("OK")
                        request_soft_refresh("urgences_admin", clear_cache=True, mute_autosync_sec=10)
                    except Exception as e:
                        st.error(str(e))

def export_chauffeur_planning_table_pdf(df, ch_code):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors
    from io import BytesIO

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=10,
        rightMargin=10,
        topMargin=15,
        bottomMargin=15,
    )

    cols = [
        "DATE","HEURE","CH","²²²²","IMMAT","PAX","Reh","Siège","",
        "DESIGNATION","H South","Décollage","N° Vol","Origine","GO","Num BDC",
        "NOM","ADRESSE","CP","Localité","Tél","Type Nav","PAIEMENT","Caisse"
    ]

    if "Unnamed: 8" in df.columns:
        df = df.rename(columns={"Unnamed: 8": ""})

    for c in cols:
        if c not in df.columns:
            df[c] = ""

    df = df[cols].fillna("")

    data = [cols]
    row_styles = []

    for i, row in df.iterrows():
        data.append([str(row[c]) for c in cols])
        r = len(data) - 1

        # 🎨 Couleurs métier (PDF)
        if int(row.get("IS_INDISPO", 0) or 0) == 1:
            row_styles.append(("BACKGROUND", (0, r), (-1, r), colors.lightgrey))

        if int(row.get("IS_URGENT", 0) or 0) == 1:
            row_styles.append(("BACKGROUND", (0, r), (-1, r), colors.salmon))

        if int(row.get("IS_GROUPAGE", 0) or 0) == 1:
            row_styles.append(("BACKGROUND", (0, r), (-1, r), colors.lightyellow))

        if int(row.get("IS_PARTAGE", 0) or 0) == 1:
            row_styles.append(("BACKGROUND", (0, r), (-1, r), colors.beige))

        paiement = str(row.get("PAIEMENT","")).lower()
        caisse = row.get("Caisse")
        if paiement == "caisse" and caisse not in ("",None,0,"0"):
            row_styles.append(("TEXTCOLOR", (-1, r), (-1, r), colors.red))
            row_styles.append(("FONT", (-1, r), (-1, r), "Helvetica-Bold"))

        if is_navette_confirmed(row):
            row_styles.append(("BACKGROUND", (0, r), (-1, r), colors.lightgreen))

    # 👉 Largeurs calibrées pour tenir sur 1 page A4 paysage
    col_widths = [
        32, 32, 30, 20, 36, 22, 22, 22, 8,
        48, 32, 32, 42, 36, 26, 36,
        44, 90, 28, 44, 44,
        38, 36, 32
    ]

    table = Table(data, colWidths=col_widths, repeatRows=1)

    table.setStyle(TableStyle(
        [
            ("GRID", (0,0), (-1,-1), 0.4, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.lightblue),
            ("FONT", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,-1), 6.5),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING", (0,0), (-1,-1), 2),
            ("RIGHTPADDING", (0,0), (-1,-1), 2),
        ] + row_styles
    ))

    doc.build([table])
    buffer.seek(0)
    return buffer.read()

def style_planning_chauffeur(row):
    styles = [""] * len(row)

    try:
        if int(row.get("IS_INDISPO", 0) or 0) == 1:
            styles = ["background-color:#eeeeee"] * len(row)

        elif int(row.get("IS_URGENT", 0) or 0) == 1:
            styles = ["background-color:#ffcccb"] * len(row)

        elif int(row.get("IS_GROUPAGE", 0) or 0) == 1:
            styles = ["background-color:#fff9c4"] * len(row)

        elif int(row.get("IS_PARTAGE", 0) or 0) == 1:
            styles = ["background-color:#ffe0b2"] * len(row)

        paiement = str(row.get("PAIEMENT", "")).lower()
        caisse = row.get("Caisse")

        if paiement == "caisse" and caisse not in ("", None, 0, "0"):
            idx = list(row.index).index("Caisse")
            styles[idx] = "background-color:#ffebee;font-weight:900;color:#c62828;"

        if is_navette_confirmed(row):
            styles = ["background-color:#e8f5e9"] * len(row)

    except Exception:
        pass

    return styles

def render_tab_chauffeur_driver():
    ch_selected = st.session_state.get("chauffeur_code")
    if not ch_selected:
        st.error("Chauffeur non identifié.")
        return

    df_ch = pd.DataFrame()  # sécurité
    today = date.today()
    # ===================================================
    # 🔄 REFRESH MANUEL (chauffeur)
    # ===================================================
    btn_col, _ = st.columns([1, 6])
    with btn_col:
        if st.button("🔄", key=f"driver_refresh_square_{ch_selected}", help="Mettre à jour le planning depuis Dropbox"):
            _topbar_sync_db_and_refresh(also_send_next3=False)

    # ===================================================
    # 💶 BADGE — CAISSE À REMETTRE
    # ===================================================
    has_caisse_due = False
    total_caisse_due = 0.0
    start_date = date(2026, 1, 1)

    with get_connection() as conn:
        df_badge = pd.read_sql_query(
            """
            SELECT DATE, HEURE, NOM, DESIGNATION, ADRESSE, PAX, PAIEMENT, Caisse
            FROM planning
            WHERE COALESCE(IS_INDISPO,0)=0
              AND COALESCE(IS_SUPERSEDED,0)=0
              AND LOWER(COALESCE(PAIEMENT,''))='caisse'
              AND COALESCE(CAISSE_PAYEE,0)=0
              AND DATE_ISO BETWEEN ? AND ?
              AND UPPER(REPLACE(REPLACE(CH,'*',''),' ','')) LIKE ?
            ORDER BY DATE_ISO, HEURE
            """,
            conn,
            params=(start_date.isoformat(), today.isoformat(), f"{normalize_ch_code(ch_selected)}%"),
        )

    if not df_badge.empty:
        df_badge["Caisse"] = pd.to_numeric(df_badge["Caisse"], errors="coerce").fillna(0.0)
        total_caisse_due = float(df_badge["Caisse"].sum())
        has_caisse_due = total_caisse_due > 0

    if has_caisse_due:
        st.markdown(
            f"""
            <div style="background:#fff3e0;border:1px solid #ff9800;
                        padding:12px;border-radius:10px;margin-bottom:10px;">
                💶 <b>Caisse à remettre :</b>
                <span style="color:#d32f2f;font-weight:900;font-size:18px;">
                    {total_caisse_due:.2f} €
                </span>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if st.toggle("🧾 Voir le détail de la caisse", False):
            st.dataframe(df_badge, use_container_width=True, height=300)
    else:
        st.success("✅ Aucune caisse à remettre pour le moment.")

    # ===================================================
    # 📅 PÉRIODE
    # ===================================================
    scope = st.radio(
        "📅 Quelles navettes veux-tu voir ?",
        ["📍 Aujourd’hui", "➡️ À partir de demain"],
        horizontal=True,
    )

    if scope == "📍 Aujourd’hui":
        from_date, to_date = today, today
        scope_label = "du jour"
    else:
        from_date, to_date = today + timedelta(days=1), None
        scope_label = "à partir de demain"

    df_all = get_planning(
        start_date=from_date,
        end_date=to_date,
        chauffeur=None,
        max_rows=5000,
        source="7j",
    )

    if df_all is None or df_all.empty:
        st.info(f"Aucune navette {scope_label}.")
        return

    # ===================================================
    # 🧭 MODE D’AFFICHAGE
    # ===================================================
    view_mode = st.radio(
        "Vue",
        ["🧾 Mes navettes", "📅 Mon planning"],
        horizontal=True,
        key="chauffeur_view_mode",
    )

    # ===================================================
    # ⚡ FILTRAGE CHAUFFEUR
    # ===================================================
    ch = ch_selected.strip().upper()
    ch_series = df_all["CH"].fillna("").str.upper().str.strip()

    mask = (
        (ch_series == ch)
        | (ch_series == f"{ch}*")
        | (ch_series.str.contains(ch, regex=False))
    ) & (~ch_series.str.match(rf"{ch}\d"))

    df_ch = _sort_df_by_date_heure(df_all[mask].copy())

    if df_ch.empty:
        st.info(f"Aucune navette {scope_label}.")
        return

    cols = df_ch.columns.tolist()

    # ===================================================
    # 📅 VUE PLANNING (TABLEAU)
    # ===================================================
    if view_mode == "📅 Mon planning":

        st.markdown("### 📅 Mon planning (chauffeur)")

        planning_cols_driver = [
            "DATE","HEURE","CH","²²²²","IMMAT","PAX","Reh","Siège",
            "Unnamed: 8","DESIGNATION","H South","Décollage","N° Vol","Origine",
            "GO","Num BDC","NOM","ADRESSE","CP","Localité","Tél",
            "Type Nav","PAIEMENT","Caisse"
        ]

        df_table = df_ch.copy()
        for c in planning_cols_driver:
            if c not in df_table.columns:
                df_table[c] = ""

        df_table = df_table[planning_cols_driver]

        def style_rows(row):
            if str(row.get("PAIEMENT","")).lower()=="caisse" and row.get("Caisse"):
                return ["background-color:#fdecea"] * len(row)
            if int(row.get("IS_GROUPAGE",0) or 0)==1:
                return ["background-color:#fffde7"] * len(row)
            if int(row.get("IS_PARTAGE",0) or 0)==1:
                return ["background-color:#f3e5f5"] * len(row)
            if int(row.get("IS_URGENT",0) or 0)==1:
                return ["background-color:#ffcdd2"] * len(row)
            return [""] * len(row)

        st.dataframe(
            df_table.style.apply(style_planning_chauffeur, axis=1),
            use_container_width=True,
            height=520,
        )

        col_pdf, col_print = st.columns(2)

        with col_pdf:
            if st.button("📄 Télécharger mon planning (PDF)", key="driver_planning_pdf"):
                pdf = export_chauffeur_planning_table_pdf(
                    df_table,
                    ch_selected,
                )

                st.download_button(
                    "⬇️ Télécharger le PDF",
                    data=pdf,
                    file_name=f"planning_{ch_selected}.pdf",
                    mime="application/pdf",
                )

        with col_print:
            if st.button("🖨️ Imprimer mon planning", key="driver_planning_print"):
                html = build_printable_html_planning(df_table, ch_selected)
                print_html_popup(html)

        st.markdown("---")
        st.markdown("### 📚 Export PDF période (historique chauffeur)")
        hist_c1, hist_c2, hist_c3 = st.columns([1, 1, 1])
        today_local = date.today()
        default_start = today_local - timedelta(days=30)

        with hist_c1:
            hist_start = st.date_input(
                "Date début",
                value=default_start,
                key=f"driver_hist_start_{ch_selected}",
            )
        with hist_c2:
            hist_end = st.date_input(
                "Date fin",
                value=today_local,
                key=f"driver_hist_end_{ch_selected}",
            )
        with hist_c3:
            export_hist = st.button(
                "📄 Générer PDF historique",
                key=f"driver_hist_pdf_btn_{ch_selected}",
                use_container_width=True,
            )

        if hist_start > hist_end:
            st.error("La date de début doit être antérieure ou égale à la date de fin.")
        else:
            if export_hist:
                df_hist = get_planning(
                    start_date=hist_start,
                    end_date=hist_end,
                    chauffeur=ch_selected,
                    type_filter=None,
                    search="",
                    max_rows=20000,
                    source="full",
                )

                if df_hist is None or df_hist.empty:
                    st.warning("Aucun transfert trouvé sur cette période.")
                else:
                    if "IS_INDISPO" in df_hist.columns:
                        df_hist = df_hist[df_hist["IS_INDISPO"].fillna(0).astype(int) == 0].copy()
                    if "IS_SUPERSEDED" in df_hist.columns:
                        df_hist = df_hist[df_hist["IS_SUPERSEDED"].fillna(0).astype(int) == 0].copy()

                    planning_cols_driver_hist = [
                        "DATE","HEURE","CH","²²²²","IMMAT","PAX","Reh","Siège",
                        "Unnamed: 8","DESIGNATION","H South","Décollage","N° Vol","Origine",
                        "GO","Num BDC","NOM","ADRESSE","CP","Localité","Tél",
                        "Type Nav","PAIEMENT","Caisse"
                    ]

                    for c in planning_cols_driver_hist:
                        if c not in df_hist.columns:
                            df_hist[c] = ""

                    df_hist = _sort_df_by_date_heure(df_hist.copy())
                    df_hist_pdf = df_hist[planning_cols_driver_hist].copy()
                    pdf_hist = export_chauffeur_planning_table_pdf(df_hist_pdf, ch_selected)
                    st.download_button(
                        "⬇️ Télécharger le PDF historique",
                        data=pdf_hist,
                        file_name=(
                            f"planning_{ch_selected}_{hist_start.strftime('%Y%m%d')}"
                            f"_{hist_end.strftime('%Y%m%d')}.pdf"
                        ),
                        mime="application/pdf",
                        key=f"driver_hist_pdf_dl_{ch_selected}",
                    )

        return  # ⛔ STOP ICI → la vue détaillée n’est PAS affichée




    # ===================================================
    # 🚖 NAVETTES
    # ===================================================
    for _, row in df_ch.iterrows():
        # 🔔 Notification chauffeur si nouvelle réponse admin
        admin_reply = row.get("ADMIN_REPLY")
        admin_reply_read = int(row.get("ADMIN_REPLY_READ") or 0)

        notif_key = f"notif_admin_reply_{row.get('id')}"

        # 🔔 Toast UNE SEULE FOIS
        if admin_reply and admin_reply_read == 0:
            if not st.session_state.get(notif_key):
                st.toast(
                    "💬 Nouveau message du bureau",
                    icon="📨"
                )
                st.session_state[notif_key] = True

        # 👀 AFFICHAGE du message admin
        if admin_reply:
            st.info(admin_reply)

            # ✅ marquer comme lu UNIQUEMENT après affichage
            if admin_reply_read == 0:
                update_planning_row(
                    row.get("id"),
                    {"ADMIN_REPLY_READ": 1}
                )



        nav_id = row.get("id")
        bloc = []

        # ------------------
        # Flags groupage / partage / attente
        # ------------------
        is_groupage = int(row.get("IS_GROUPAGE", 0) or 0) == 1
        is_partage = int(row.get("IS_PARTAGE", 0) or 0) == 1
        is_attente = int(row.get("IS_ATTENTE", 0) or 0) == 1

        prefix = ""
        if is_groupage:
            prefix += "🟡 [GROUPÉE] "
        elif is_partage:
            prefix += "🟡 [PARTAGÉE] "
        if is_attente:
            prefix += "⭐ "
        # 🟢 Statut chauffeur en ligne
        ch_root = normalize_ch_code(row.get("CH", ch_selected))

        # ------------------
        # Chauffeur + statut
        # ------------------
        ch_code = str(row.get("CH", "") or ch_selected).strip()

        if row.get("IS_INDISPO") == 1:
            ch_status = "🚫 Indispo"
        elif is_navette_confirmed(row):
            ch_status = "🟢 Confirmé"
        else:
            ch_status = "🟠 À confirmer"

        bloc.append(f"👨‍✈️ **{ch_code}** — {ch_status}")


        # ------------------
        # Confirmation
        # ------------------
        if is_navette_confirmed(row):
            bloc.append("✅ **Navette confirmée**")
        else:
            bloc.append("🕒 **À confirmer**")

        # ------------------
        # Date / Heure
        # ------------------
        dv = row.get("DATE")
        if isinstance(dv, (datetime, date)):
            date_obj = dv if isinstance(dv, date) else dv.date()
            date_txt = date_obj.strftime("%d/%m/%Y")
        else:
            dtmp = pd.to_datetime(dv, dayfirst=True, errors="coerce")
            date_obj = dtmp.date() if not pd.isna(dtmp) else None
            date_txt = date_obj.strftime("%d/%m/%Y") if date_obj else ""

        heure_txt = normalize_time_string(row.get("HEURE")) or "??:??"
        bloc.append(f"{prefix}📆 {date_txt} | ⏱ {heure_txt}")

        # ------------------
        # Sens / Destination
        # ------------------
        sens_txt = format_sens_ar(row.get("Unnamed: 8"))

        dest_raw = ""
        for cand in ["DESIGNATION", "DESTINATION", "DE/VERS"]:
            if cand in cols and row.get(cand):
                dest_raw = str(row.get(cand)).strip()
                if dest_raw:
                    break

        dest = resolve_client_alias(dest_raw)

        if sens_txt and dest:
            bloc.append(f"➡ {sens_txt} ({dest})")
        elif sens_txt:
            bloc.append(f"➡ {sens_txt}")
        elif dest:
            bloc.append(f"➡ {dest}")

        # ------------------
        # Client
        # ------------------
        nom = str(row.get("NOM", "") or "").strip()
        if nom:
            bloc.append(f"🧑 {nom}")

        # ------------------
        # 👥 PAX
        # ------------------
        pax = row.get("PAX")
        if pax not in ("", None, 0, "0"):
            try:
                pax_i = int(pax)
                if pax_i > 0:
                    bloc.append(f"👥 **{pax_i} pax**")
            except Exception:
                bloc.append(f"👥 **{pax} pax**")

        # ------------------
        # 🚘 Véhicule (SIÈGE BÉBÉ / RÉHAUSSEUR)
        # ------------------
        immat = str(row.get("IMMAT", "") or "").strip()
        if immat:
            bloc.append(f"🚘 Plaque : {immat}")

        siege_bebe = extract_positive_int(row.get("SIEGE", row.get("SIÈGE")))
        if siege_bebe:
            bloc.append(f"🍼 Siège bébé : {siege_bebe}")

        reh_n = extract_positive_int(row.get("REH"))
        if reh_n:
            bloc.append(f"🪑 Rehausseur : {reh_n}")

        # ------------------
        # Adresse / Tel
        # ------------------
        adr = build_full_address_from_row(row)
        nav_adr = build_navigation_address_from_row(row)
        if adr:
            bloc.append(f"📍 {adr}")

        tel = get_client_phone_from_row(row)
        if tel:
            bloc.append(f"📞 {tel}")

        # ------------------
        # Paiement
        # ------------------
        paiement = str(row.get("PAIEMENT", "") or "").lower().strip()
        caisse = row.get("Caisse")

        if paiement == "facture":
            bloc.append("🧾 **FACTURE**")
        elif paiement == "caisse" and caisse:
            bloc.append(
                "<span style='color:#d32f2f;font-weight:800;'>"
                f"💶 {caisse} € (CASH)</span>"
            )
        elif paiement == "bancontact" and caisse:
            bloc.append(
                "<span style='color:#1976d2;font-weight:800;'>"
                f"💳 {caisse} € (BANCONTACT)</span>"
            )

        # ===================================================
        # ✈️ Vol – TOUJOURS AFFICHÉ / STATUT = JOUR J
        # ===================================================
        vol = extract_vol_val(row, cols)
        if vol:
            bloc.append(f"✈️ Vol **{vol}**")

            # 🔎 Vérification statut UNIQUEMENT le jour J
            if date_obj and date_obj == today:
                status, delay_min, *_ = get_flight_status_cached(vol)
                badge = flight_badge(status, delay_min)

                if badge:
                    bloc.append(f"📡 {badge}")

                if delay_min is not None and delay_min >= FLIGHT_ALERT_DELAY_MIN:
                    bloc.append(
                        f"🚨 **ATTENTION : retard {delay_min} min**"
                    )

        # ------------------
        # GO
        # ------------------
        go_val = str(row.get("GO", "") or "").strip()
        if go_val:
            bloc.append(f"🟢 {go_val}")

        # ------------------
        # 🧾 BDC (juste après GO)
        # ------------------
        for cand in ["NUM BDC", "Num BDC", "NUM_BDC", "BDC"]:
            if cand in cols and row.get(cand):
                bloc.append(f"🧾 **BDC : {row.get(cand)}**")
                break

        # ------------------
        # Actions
        # ------------------
        actions = []

        if tel:
            actions.append(f"[📞 Appeler](tel:{clean_phone(tel)})")

        if nav_adr:
            actions.append(f"[🧭 Waze]({build_waze_link(nav_adr)})")
            actions.append(f"[🗺 Google Maps]({build_google_maps_link(nav_adr)})")

        if tel:
            # =========================
            # 📞 GSM CHAUFFEUR(S) – Feuil2
            # =========================
            ch_raw = row.get("CH", "")
            phones = get_chauffeurs_phones(ch_raw)
            tel_chauffeur = " / ".join(phones) if phones else "—"

            msg = build_client_sms_from_driver(
                row,
                ch_selected,
                tel_chauffeur,
            )
            actions.append(
                f"[💬 WhatsApp]({build_whatsapp_link(tel, msg)})"
            )


        if actions:
            bloc.append(" | ".join(actions))

        # ------------------
        # Affichage
        # ------------------
        st.markdown("<br>".join(bloc), unsafe_allow_html=True)

        # ------------------
        # Saisie chauffeur
        # ------------------
        trajet_key = f"trajet_nav_{nav_id}"
        prob_key = f"prob_nav_{nav_id}"

        st.text_input("Trajet compris", key=trajet_key)

        with st.expander("🚨 Signaler un problème"):
            st.text_area("Décrire le problème", key=prob_key)

        st.markdown("---")

    # ===================================================
    # 📅 VUE TABLEAU — PLANNING CHAUFFEUR
    # ===================================================
    if view_mode == "📅 Mon planning":

        st.markdown("### 📅 Mon planning")

        # Colonnes visibles chauffeur (ordre métier)
        planning_cols_driver = [
            "DATE",
            "HEURE",
            "Unnamed: 8",   # SENS
            "DESIGNATION",  # DEST
            "NOM",
            "ADRESSE",
            "CP",
            "Localité",
            "Tél",
            "PAX",
            "PAIEMENT",
            "Caisse",
        ]

        df_table = df_ch.copy()

        # Sécurité colonnes
        for c in planning_cols_driver:
            if c not in df_table.columns:
                df_table[c] = ""

        df_table = df_table[planning_cols_driver]

        # Renommage propre affichage
        df_table = df_table.rename(columns={
            "Unnamed: 8": "SENS",
            "DESIGNATION": "DEST",
            "Localité": "LOCALITÉ",
            "Tél": "TÉL",
        })

        # 🔴 Mise en évidence CAISSE
        def _style_caisse(v):
            try:
                if float(v) > 0:
                    return "background-color:#fdecea;font-weight:800;"
            except Exception:
                pass
            return ""

        st.dataframe(
            df_table.style.applymap(_style_caisse, subset=["Caisse"]),
            use_container_width=True,
            height=420,
        )

        st.markdown("")

        # ===================================================
        # 📄 EXPORT PDF (PAYSAGE)
        # ===================================================
        if st.button("📄 Télécharger mon planning (PDF)", key="driver_planning_pdf"):
            pdf = export_chauffeur_planning_table_pdf(
                df_table,
                ch_selected,
            )

            st.download_button(
                "⬇️ Télécharger le PDF",
                data=pdf,
                file_name=f"planning_{ch_selected}.pdf",
                mime="application/pdf",
            )

    # ===================================================
    # 📤 ENVOI CONFIRMATION (RÉPONSE CHAUFFEUR)
    # ===================================================

    import time

    CONFIRM_MSG_DURATION = 5  # secondes

    confirm_ts = st.session_state.get("confirm_ts")
    confirm_label = st.session_state.get("confirm_label")

    if confirm_ts and confirm_label:
        if (time.time() - confirm_ts) < CONFIRM_MSG_DURATION:
            st.success(confirm_label)
        else:
            st.session_state.pop("confirm_ts", None)
            st.session_state.pop("confirm_label", None)

    if st.button("📤 Envoyer mes informations"):

        recap = []
        now_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        for _, row in df_ch.iterrows():
            nav_id = row.get("id")

            trajet = (
                st.session_state.get(f"trajet_nav_{nav_id}", "")
                .strip()
            )
            probleme = (
                st.session_state.get(f"prob_nav_{nav_id}", "")
                .strip()
            )

            if not trajet and not probleme:
                continue

            # -----------------------------
            # Texte RÉEL de la réponse chauffeur
            # -----------------------------
            ack_parts = []

            if trajet:
                ack_parts.append(f"Trajet compris : {trajet}")

            if probleme:
                ack_parts.append(f"⚠️ Problème signalé : {probleme}")

            ack_text = "\n".join(ack_parts)

            # -----------------------------
            # 🔴 UPDATE DB (CRITIQUE)
            # -----------------------------
            update_planning_row(
                nav_id,
                {
                    "ACK_AT": now_iso,
                    "ACK_TEXT": ack_text,
                },
            )

            recap.append(
                format_navette_ack(
                    row=row,
                    ch_selected=ch_selected,
                    trajet=trajet,
                    probleme=probleme,
                )
            )

        if not recap:
            st.warning("Aucune information encodée.")
            return

        # -----------------------------
        # 📧 Mail admin (inchangé)
        # -----------------------------
        send_mail_admin(
            subject=f"[INFOS CHAUFFEUR] {ch_selected}",
            body="\n\n".join(recap),
        )

        set_chauffeur_last_ack(ch_selected)

        # -----------------------------
        # Feedback UI
        # -----------------------------
        heure_txt = datetime.now().strftime("%H:%M")
        st.session_state["confirm_ts"] = time.time()
        st.session_state["confirm_label"] = f"✅ Informations envoyées à {heure_txt}"
        st.session_state["tab_refresh"]["planning"] = time.time()

        st.rerun()




# ======================================================================
#  ONGLET — Demandes d’indispo côté chauffeur
# ======================================================================

def render_tab_indispo_driver(ch_code: str):
    st.subheader("🚫 Mes indisponibilités")

    today = date.today()

    with st.form("form_indispo"):
        d = st.date_input("Date", value=today)
        col1, col2 = st.columns(2)
        with col1:
            h_debut = st.text_input("Heure début (ex: 08:00)")
        with col2:
            h_fin = st.text_input("Heure fin (ex: 12:00)")
        commentaire = st.text_input("Commentaire (optionnel)")
        submit = st.form_submit_button("📩 Envoyer la demande")

    if submit:
        req_id = create_indispo_request(ch_code, d, h_debut, h_fin, commentaire)

        # mail automatique
        send_mail_admin(
            f"Nouvelle indispo chauffeur {ch_code}",
            f"Chauffeur : {ch_code}\n"
            f"Date : {d.strftime('%d/%m/%Y')}\n"
            f"De {h_debut} à {h_fin}\n"
            f"Commentaire : {commentaire}\n"
            f"ID demande : {req_id}"
        )

        st.success("Demande envoyée à l’admin")
        st.rerun()

    st.markdown("### Mes demandes")
    df = get_indispo_requests(chauffeur=ch_code)

    st.dataframe(df, use_container_width=True, height=300)

# ============================================================
#   ONGLET 👨‍✈️ FEUIL2 / CHAUFFEURS
# ============================================================

def render_tab_chauffeurs():
    st.subheader("👨‍✈️ Chauffeurs (Feuil2)")

    try:
        with get_connection() as conn:
            df = pd.read_sql_query(
                'SELECT * FROM "chauffeurs" ORDER BY INITIALE',
                conn,
            )
    except Exception as e:
        st.error(f"Erreur en lisant la table `chauffeurs` : {e}")
        return

    # 🔒 Sécurité Streamlit : aucune colonne dupliquée
    df = df.loc[:, ~df.columns.duplicated()]

    st.markdown("#### Table chauffeurs (éditable)")
    edited = st.data_editor(
        df,
        use_container_width=True,
        num_rows="dynamic",
        key="chauffeurs_editor",
    )

    if st.button("💾 Enregistrer les modifications (chauffeurs)"):
        try:
            with get_connection() as conn:
                cur = conn.cursor()

                # On repart de zéro pour éviter doublons / lignes fantômes
                cur.execute('DELETE FROM "chauffeurs"')

                cols = [c for c in edited.columns if c != "id"]
                col_list_sql = ",".join(f'"{c}"' for c in cols)
                placeholders = ",".join("?" for _ in cols)

                for _, row in edited.iterrows():
                    values = [
                        row[c] if pd.notna(row[c]) else None
                        for c in cols
                    ]
                    cur.execute(
                        f'INSERT INTO "chauffeurs" ({col_list_sql}) VALUES ({placeholders})',
                        values,
                    )

                conn.commit()

            st.success("Table chauffeurs mise à jour ✅")
            st.rerun()

        except Exception as e:
            st.error(f"Erreur lors de la sauvegarde des chauffeurs : {e}")



# ============================================================
#   ONGLET 📄 FEUIL3 (INFOS DIVERSES)
# ============================================================

def render_tab_feuil3():
    st.subheader("📄 Feuil3 (infos diverses / logins, etc.)")

    try:
        with get_connection() as conn:
            df = pd.read_sql_query(
                "SELECT rowid AS id, * FROM feuil3",
                conn,
            )
    except Exception as e:
        st.warning(f"Table `feuil3` introuvable ou erreur : {e}")
        st.info("Si tu veux l'utiliser, ajoute la feuille Feuil3 dans l'Excel et relance l'import.")
        return

    st.markdown("#### Table Feuil3 (éditable)")
    edited = st.data_editor(
        df,
        use_container_width=True,
        num_rows="dynamic",
        key="feuil3_editor",
    )

    if st.button("💾 Enregistrer les modifications (Feuil3)"):
        try:
            with get_connection() as conn:
                cur = conn.cursor()
                cur.execute("DELETE FROM feuil3")

                cols = [c for c in edited.columns if c != "id"]
                col_list_sql = ",".join(f'"{c}"' for c in cols)
                placeholders = ",".join("?" for _ in cols)

                for _, row in edited.iterrows():
                    values = [row[c] if pd.notna(row[c]) else None for c in cols]
                    cur.execute(
                        f"INSERT INTO feuil3 ({col_list_sql}) VALUES ({placeholders})",
                        values,
                    )
                conn.commit()
            st.success("Table Feuil3 mise à jour ✅")
            st.rerun()
        except Exception as e:
            st.error(f"Erreur lors de la sauvegarde de Feuil3 : {e}")


# ============================================================
#   ONGLET 📂 EXCEL ↔ DB (Dropbox)
# ============================================================

def render_tab_excel_sync():

    from datetime import datetime

    # ===================================================
    # 🔐 SÉCURITÉ — ADMIN UNIQUEMENT
    # ===================================================
    if st.session_state.get("role") != "admin":
        st.warning("🔒 Seuls les administrateurs peuvent synchroniser la base.")
        return

    st.subheader("📂 Synchronisation Excel → Base de données")

    # ===================================================
    # 🟢 DERNIÈRE SYNCHRO
    # ===================================================
    last_sync = st.session_state.get("last_sync_time")
    if last_sync:
        st.success(f"🟢 Dernière mise à jour : {last_sync}")
    else:
        st.info("🔴 Aucune synchronisation effectuée dans cette session")

    st.markdown("---")

    # ===================================================
    # ℹ️ INFO WORKFLOW
    # ===================================================
    st.markdown(
        """
        **Source principale du planning : Dropbox (Excel unique)**

        ---
        🔧 **Workflow normal :**

        1. Ouvre le fichier **Planning 2026.xlsx** dans **Dropbox**
        2. Modifie *Feuil1*, *Feuil2*, *Feuil3*
        3. Enregistre le fichier
        4. Clique sur **FORCER MAJ DROPBOX → DB**
        """
    )

    st.markdown("---")

    # ===================================================
    # 🆘 MODE SECOURS — UPLOAD MANUEL
    # ===================================================
    st.subheader("🆘 Mode secours — Charger un fichier Excel manuellement")

    uploaded_file = st.file_uploader(
        "📤 Charger un fichier Planning Excel (.xlsx)",
        type=["xlsx"],
        accept_multiple_files=False,
    )

    if uploaded_file:
        st.info(
            f"📄 Fichier chargé : {uploaded_file.name}\n\n"
            "⚠️ Les navettes fantômes (non confirmées / non payées) seront nettoyées."
        )

        confirm_upload = st.checkbox(
            "Je confirme vouloir synchroniser la base depuis ce fichier",
            key="confirm_manual_excel_upload",
        )

        if st.button(
            "🆘 SYNCHRONISER DEPUIS LE FICHIER MANUEL",
            type="secondary",
            disabled=not confirm_upload,
        ):
            st.session_state["_do_manual_excel_sync"] = True

    if st.session_state.pop("_do_manual_excel_sync", False):
        with st.spinner("🔄 Synchronisation depuis fichier manuel…"):
            sync_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            inserted = sync_planning_from_uploaded_file(
                uploaded_file,
                excel_sync_ts=sync_ts,
            )
            cleanup_orphan_planning_rows(sync_ts)
            log_event(
                f"Sync fichier manuel + cleanup exécutés ({inserted} lignes)",
                "SYNC",
            )

        st.session_state["last_sync_time"] = datetime.now().strftime("%d/%m/%Y %H:%M")
        st.success(f"✅ DB mise à jour ({inserted} lignes)")

    st.markdown("---")

    # ===================================================
    # 🔄 SYNCHRO MANUELLE DROPBOX
    # ===================================================
    confirm = st.checkbox(
        "Je confirme vouloir forcer la mise à jour de la base depuis Dropbox",
        key="confirm_force_sync_dropbox_v2",
    )

    col1, col2 = st.columns([2, 3])

    with col1:
        if st.button(
            "🔄 FORCER MAJ DROPBOX → DB",
            type="primary",
            disabled=not confirm,
        ):
            st.session_state["_do_dropbox_sync"] = True

    with col2:
        st.caption(
            "⚠️ Les navettes supprimées ou déplacées dans Excel seront nettoyées "
            "si elles ne sont ni confirmées ni payées."
        )

    if st.session_state.pop("_do_dropbox_sync", False):
        with st.spinner("🔄 Synchronisation en cours depuis Dropbox…"):
            sync_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            inserted = sync_planning_from_today(excel_sync_ts=sync_ts)
            cleanup_orphan_planning_rows(sync_ts)
            log_event(
                f"Sync Dropbox + cleanup exécutés ({inserted} lignes)",
                "SYNC",
            )

        st.session_state["last_sync_time"] = datetime.now().strftime("%d/%m/%Y %H:%M")
        st.success(f"✅ DB mise à jour depuis aujourd’hui ({inserted} lignes)")

    st.markdown("---")

    # ===================================================
    # 🔥 RECONSTRUCTION COMPLÈTE DB (DANGER)
    # ===================================================
    st.markdown("### 🔥 Reconstruction complète de la base (DANGER)")

    rebuild_file_1 = st.file_uploader(
        "📂 Sélectionne le PREMIER fichier Excel (ex : Planning 2025)",
        type=["xlsx"],
        key="rebuild_excel_file_1",
    )

    rebuild_file_2 = st.file_uploader(
        "📂 Sélectionne le DEUXIÈME fichier Excel (ex : Planning 2026)",
        type=["xlsx"],
        key="rebuild_excel_file_2",
    )

    confirm_full = st.checkbox(
        "⚠️ Je confirme vouloir reconstruire TOUTE la base",
        key="confirm_full_rebuild",
    )

    if st.button(
        "🔥 RECONSTRUIRE DB COMPLÈTE",
        type="secondary",
        disabled=not (confirm_full and rebuild_file_1 and rebuild_file_2),
    ):
        st.session_state["_do_full_rebuild"] = True

    if st.session_state.pop("_do_full_rebuild", False):
        with st.spinner("🔥 Reconstruction complète de la base en cours…"):
            inserted = rebuild_planning_db_from_two_excel_files(
                rebuild_file_1,
                rebuild_file_2,
            )

        st.session_state["last_sync_time"] = datetime.now().strftime("%d/%m/%Y %H:%M")
        st.success(f"✅ DB reconstruite ({inserted} lignes)")



# ============================================================
#   ONGLET 📦 ADMIN TRANSFERTS (LISTE GLOBALE)
# ============================================================

def render_tab_admin_transferts():
    import pandas as pd
    import streamlit as st
    from datetime import date, datetime, timedelta

    st.subheader("📦 Tous les transferts — vue admin")

    # 🔒 Sécurité anti-UnboundLocalError
    df = pd.DataFrame()


    # ✅ 6 onglets
    tab_transferts, tab_excel, tab_heures, tab_mail, tab_urgences, tab_whatsapp = st.tabs(
        [
            "📋 Transferts / SMS",
            "🟡 À reporter dans Excel",
            "⏱️ Calcul d’heures",
            "📥 Mail → Navette",
            "🚨 Urgences",
            "💬 WhatsApp",
        ]
    )

    # ------------------------------------------------------
    # Helpers dates -> ISO (évite 0 lignes)
    # ------------------------------------------------------
    def _to_iso(d):
        if isinstance(d, (datetime, date)):
            return d.strftime("%Y-%m-%d")
        s = str(d or "").strip()
        if not s:
            return ""
        try:
            dt = pd.to_datetime(s, dayfirst=True, errors="coerce")
            if pd.isna(dt):
                return s
            return dt.strftime("%Y-%m-%d")
        except Exception:
            return s
    # ======================================================
    # 📥 ONGLET MAIL → NAVETTE (ASSISTANT MÉTIER + PDF + APPRENTISSAGE)
    # ======================================================
    st.subheader("📥 Mail → Navette")

    # 📎 Import fichier Word (.docx) — BT Tours / formats variés
    docx_file = st.file_uploader(
        "📎 Importer un fichier Word (.docx) (BT Tours, etc.)",
        type=["docx"],
        key="mail_docx_upload"
    )

    def _extract_text_from_docx(uploaded):
        try:
            from docx import Document
            doc = Document(uploaded)
            chunks = []

            # Paragraphes
            for p in doc.paragraphs:
                if p.text.strip():
                    chunks.append(p.text.strip())

            # Tableaux
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        chunks.append(" | ".join(cells))

            txt = "\n".join(chunks)
            txt = txt.replace("\r", "\n")
            txt = re.sub(r"[ \t]+", " ", txt)
            return txt.strip()
        except Exception:
            return ""

    st.caption("Colle un mail / WhatsApp / texte PDF → propositions → tu corriges → tu valides (apprentissage).")

    consume_soft_refresh("admin_tab_mail")

    # ================= OPTIONS =================
    col_opt1, col_opt2, col_opt3, col_opt4 = st.columns([1, 1, 1, 2])
    with col_opt1:
        urgence_mode = st.checkbox("🚨 Mode urgence", value=False, key="mail_urgence_mode")
    with col_opt2:
        notify_now = st.checkbox("🔔 Notifier chauffeur", value=False, key="mail_notify_now")
    with col_opt3:
        add_retour = st.button("➕ Dupliquer / retour manuel", key="mail_add_retour")
    with col_opt4:
        learn_enabled = st.checkbox("🧠 Apprentissage activé", value=True, key="mail_learn_enabled")

    # ================= PDF =================
    st.markdown("### 📎 Pièce jointe (optionnel)")
    up = st.file_uploader("Dépose un PDF", type=["pdf"], key="mail_pdf_uploader")

    import io
    import re
    from datetime import date, datetime

    def _extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
        if not pdf_bytes:
            return ""
        try:
            import pdfplumber
            out = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for p in pdf.pages[:20]:
                    t = p.extract_text() or ""
                    if t.strip():
                        out.append(t)
            return "\n".join(out).strip()
        except Exception:
            pass
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(pdf_bytes))
            out = []
            for p in reader.pages[:20]:
                t = p.extract_text() or ""
                if t.strip():
                    out.append(t)
            return "\n".join(out).strip()
        except Exception:
            return ""

    if up is not None:
        pdf_bytes = up.read()
        st.download_button(
            "⬇️ Télécharger le PDF",
            data=pdf_bytes,
            file_name=up.name or "piece_jointe.pdf",
            mime="application/pdf",
        )
        if st.button("📄 Extraire texte du PDF → coller", key="mail_pdf_extract_btn"):
            txt = _extract_text_from_pdf_bytes(pdf_bytes)
            if txt:
                st.session_state["mail_raw_input"] = txt
                st.toast("✅ Texte extrait et injecté", icon="📄")
            else:
                st.warning("⚠️ PDF non extractible (scan). Colle le texte manuellement.")

    # ================= TEXTE =================
    raw_mail = st.text_area(
        "📋 Texte à analyser",
        height=260,
        key="mail_raw_input",
    )

    # ======================================================
    # 🧠 UTILITAIRES
    # ======================================================
    def _norm(x):
        return re.sub(r"\s+", " ", str(x or "").upper().strip())

    def _norm_tel(x):
        t = re.sub(r"[^\d+]", "", str(x or ""))
        if t.startswith("00"):
            t = "+" + t[2:]
        return t

    def _normalize_time_txt(s):
        s = str(s or "").upper().replace(":", "H")
        m = re.search(r"(\d{1,2})\s*H\s*(\d{2})?", s)
        if not m:
            return ""
        return f"{int(m.group(1)):02d}:{int(m.group(2) or 0):02d}"

    def _to_date_iso_any(raw):
        """
        Convertit une date vers ISO 'YYYY-MM-DD'.
        Supporte:
          - 12/03/26, 12-03-2026
          - 12 mars 26, 12 mars 2026
        """
        s = str(raw or "").strip()
        if not s:
            return ""

        # 1) formats numériques
        m = re.search(r"(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{2,4})", s)
        if m:
            d = int(m.group(1))
            mo = int(m.group(2))
            y = int(m.group(3))
            if y < 100:
                y += 2000
            try:
                return date(y, mo, d).isoformat()
            except Exception:
                return ""

        # 2) formats "12 mars 26"
        months = {
            "JANVIER": 1, "JANV": 1,
            "FEVRIER": 2, "FÉVRIER": 2, "FEVR": 2, "FÉVR": 2,
            "MARS": 3,
            "AVRIL": 4, "AVR": 4,
            "MAI": 5,
            "JUIN": 6,
            "JUILLET": 7, "JUIL": 7,
            "AOUT": 8, "AOÛT": 8,
            "SEPTEMBRE": 9, "SEPT": 9,
            "OCTOBRE": 10, "OCT": 10,
            "NOVEMBRE": 11, "NOV": 11,
            "DECEMBRE": 12, "DÉCEMBRE": 12, "DEC": 12, "DÉC": 12,
        }

        m2 = re.search(r"(\d{1,2})\s+([A-Za-zÉÈÊËÀÂÄÔÖÛÜÙÇéèêëàâäôöûüùç]+)\s+(\d{2,4})", s, re.IGNORECASE)
        if m2:
            d, mon, y = m2.groups()
            mon_u = mon.strip().upper()
            mon_u = (mon_u
                     .replace("É","E").replace("È","E").replace("Ê","E").replace("Ë","E")
                     .replace("À","A").replace("Â","A").replace("Ä","A")
                     .replace("Ô","O").replace("Ö","O")
                     .replace("Û","U").replace("Ü","U").replace("Ù","U")
                     .replace("Ç","C"))
            mo = months.get(mon_u) or months.get(mon_u[:4]) or months.get(mon_u[:3])
            y = int(y)
            if y < 100:
                y += 2000
            if mo:
                try:
                    return date(y, int(mo), int(d)).isoformat()
                except Exception:
                    return ""

        return ""



    def _next_non_empty(lines, start):
        for k in range(start, len(lines)):
            if str(lines[k]).strip():
                return str(lines[k]).strip()
        return ""

    def _label_value_or_next(lines, i):
        """
        Gère:
        - "Date : 23/02/2026"
        - "Date :" puis valeur à la ligne suivante
        """
        line = str(lines[i]).strip()
        parts = line.split(":", 1)
        if len(parts) == 2 and parts[1].strip():
            return parts[1].strip()
        return _next_non_empty(lines, i + 1)

    def _label_value_or_next_multiline(lines, i, max_lines=6):
        """Valeur sur la même ligne ou concatène les lignes suivantes (utile pour 'Transfert à')."""
        first = _label_value_or_next(lines, i)
        # Si la valeur est déjà sur la même ligne, on peut quand même ajouter les lignes suivantes
        parts = []
        if first:
            parts.append(first)

        # Concatène les lignes suivantes tant qu'on ne retombe pas sur un autre label
        stop_prefixes = (
            "SECTEUR", "S.B.U", "SBU", "DEMANDEUR", "IMPUTATION",
            "DATE", "HEURE", "LIEU", "TRANSFERT", "VOYAGEUR", "N° GSM", "GSM", "IMPUTATION"
        )
        j = i + 1
        taken = 0
        while j < len(lines) and taken < max_lines:
            ln = str(lines[j]).strip()
            if not ln:
                j += 1
                continue
            ul = _norm(ln)
            # stop si c'est un label du genre "XXX :"
            if any(ul.startswith(p) for p in stop_prefixes) and ul.endswith(":"):
                break
            # stop si ligne ressemble à "Label :" (même si pas dans stop_prefixes)
            if re.match(r"^[A-ZÉÈÊËÀÂÄÔÖÛÜÙÇ0-9 ()'\"./-]+\s*:\s*$", ln):
                break

            # éviter de dupliquer first si ln == first
            if (not parts) or (_norm(parts[-1]) != _norm(ln)):
                parts.append(ln)
                taken += 1
            j += 1

        return " ".join([p for p in parts if p]).strip()

    # ======================================================
    # 🧠 MAPPING DESIGNATION (TA RÈGLE)
    # ======================================================
    def _designation_from_place(txt: str) -> str:
        u = _norm(txt)
        # ✅ Spécifique : Gare de Bruxelles Midi
        if "MIDI" in u or "BRUXELLES MIDI" in u or "BRUSSELS MIDI" in u:
            return "MIDI"
        if "ZAVENTEM" in u or "BRU" in u or "BRUSSEL" in u or "BRUX" in u:
            return "BRU"
        if "CHARLEROI" in u or "CRL" in u:
            return "CRL"
        if "LUX" in u or "LUXEMBOURG" in u:
            return "LUX"
        if "DUSS" in u or "DUSSELDORF" in u:
            return "DUSS"
        if "MIDI" in u:
            return "MIDI"
        # John Cockerill sites
        if "JOHN COCKERILL" in u and "SERAING" in u:
            return "JCO"
        if "JOHN COCKERILL" in u and "DEFENSE" in u:
            return "JCD"
        if "JOHN COCKERILL" in u and "CHATEAU" in u:
            return "JCC"
        # fallback: garde brut si pas reconnu
        return str(txt or "").strip()

    def _sens_from_pick_drop(pickup: str, dropoff: str) -> str:
        """
        Ton besoin: sens pas vide.
        - Si pickup est aéroport -> DE
        - Sinon si dropoff est un "site société" (ex John Cockerill Seraing) -> VERS
        - Sinon: si dropoff aéroport -> VERS
        """
        pu = _norm(pickup)
        do = _norm(dropoff)
        aeroport = ["AEROPORT", "AIRPORT", "ZAVENTEM", "BRU", "CRL", "CHARLEROI", "LUX", "MAASTRICHT"]
        if any(k in pu for k in aeroport):
            return "DE"
        if ("JOHN COCKERILL" in do) or ("KNAUF" in do) or ("FN HERSTAL" in do) or ("BT TOUR" in do):
            return "VERS"
        if any(k in do for k in aeroport):
            return "VERS"
        return "VERS"  # par défaut, jamais vide

    def _company_prefix_from_text(txt: str) -> str:
        u = _norm(txt)
        if ("JOHN COCKERILL" in u) or ("JCO" in u) or ("JCD" in u) or ("JCC" in u) or ("JCSA" in u):
            return "JC"
        if ("KNAUF" in u) or ("INSULATION" in u) or ("KI " in u):
            return "KI"
        if ("FN HERSTAL" in u) or ("FNH" in u):
            return "FNH"
        if ("BT TOUR" in u) or ("BT " in u):
            return "BT"
        return ""

    # ======================================================
    # 🧠 PARSE MAIL "CORPORATE" (robuste lignes)
    # ======================================================
    def _parse_mail_corporate(text: str):
        lines = [l.rstrip() for l in (text or "").splitlines()]
        uall = _norm(text)

        nom = ""
        tel = ""
        demandeur = ""
        imputation = ""
        prefix = _company_prefix_from_text(text)

        # champs simples (valeur sur même ligne ou ligne suivante)
        for i in range(len(lines)):
            ul = _norm(lines[i])
            if ul.startswith("VOYAGEUR") or "TRAVELER" in ul:
                nom = _label_value_or_next(lines, i)
            elif ul.startswith("N° GSM") or ul == "GSM" or "N° GSM" in ul:
                tel = _norm_tel(_label_value_or_next(lines, i))
            elif ul.startswith("DEMANDEUR") or "REQUEST SEND BY" in ul:
                demandeur = _label_value_or_next(lines, i)
            elif ul.startswith("IMPUTATION") or "COST CENTER" in ul:
                imputation = _label_value_or_next(lines, i)

        # blocs trajet
        blocks = []
        i = 0
        while i < len(lines):
            ul = _norm(lines[i])
            if ul.startswith("DATE"):
                d_raw = _label_value_or_next(lines, i)
                d_iso = _to_date_iso_any(d_raw)
                heure = ""
                pickup = ""
                dropoff = ""

                j = i + 1
                while j < len(lines):
                    uj = _norm(lines[j])

                    if uj.startswith("DATE"):
                        break

                    if ("HEURE DE PRISE" in uj) or uj.startswith("HEURE"):
                        heure = _normalize_time_txt(_label_value_or_next(lines, j))

                    if ("LIEU DE PRISE" in uj) or ("LIEU DE PICK" in uj):
                        pickup = _label_value_or_next(lines, j)

                    if ("TRANSFERT A" in uj) or ("TRANSFERT À" in uj) or uj.startswith("TRANSFERT"):
                        dropoff = _label_value_or_next_multiline(lines, j, max_lines=8)

                    j += 1

                if d_iso:
                    blocks.append({
                        "DATE": d_iso,
                        "HEURE": heure,
                        "pickup": pickup,
                        "dropoff": dropoff,
                    })

                i = j
                continue

            i += 1

        return {
            "NOM": str(nom or "").strip(),
            "Tél": _norm_tel(tel),
            "DEMANDEUR": str(demandeur or "").strip(),
            "IMPUTATION": str(imputation or "").strip(),
            "PREFIX": prefix,
            "BLOCKS": blocks,
        }

    # ======================================================
    # 🧠 DB: dernier Num BDC du passager (si dispo) + lookup prix trajet
    # ======================================================
    def _db_last_bdc_for_passenger(nom: str, tel: str, prefix: str) -> str:
        n = _norm(nom)
        t = _norm_tel(tel)

        if not n and not t:
            return ""

        where = []
        params = []

        if n:
            toks = [x for x in re.split(r"\s+", n) if len(x) >= 3]
            if toks:
                where.append("(" + " AND ".join(["UPPER(TRIM(NOM)) LIKE ?"] * len(toks)) + ")")
                params.extend([f"%{tok}%" for tok in toks])
            else:
                where.append("UPPER(TRIM(NOM)) LIKE ?")
                params.append(f"%{n}%")

        if t:
            where.append("REPLACE(REPLACE(REPLACE(REPLACE(TRIM(Tél),' ',''),'/',''),'.',''),'-','') LIKE ?")
            params.append(f"%{re.sub(r'[^0-9+]', '', t)}%")

        if prefix:
            where.append("UPPER(TRIM(\"Num BDC\")) LIKE ?")
            params.append(f"{prefix}%")

        sql = f"""
            SELECT TRIM("Num BDC")
            FROM planning
            WHERE {" AND ".join(where)}
            ORDER BY DATE_ISO DESC
            LIMIT 1
        """

        with get_connection() as conn:
            r = conn.execute(sql, params).fetchone()
        return (r[0] if r and r[0] else "")

    def _db_lookup_price_for_route(adresse: str, designation: str, sens: str):
        adr = _norm(adresse)
        des = _norm(designation)
        se = _norm(sens)

        if not adr or not des or not se:
            return None

        with get_connection() as conn:
            row = conn.execute(
                """
                SELECT KM, "H TVA", TTC, "Type Nav", PAIEMENT
                FROM planning
                WHERE
                    UPPER(TRIM(ADRESSE)) = ?
                    AND UPPER(TRIM(DESIGNATION)) = ?
                    AND TRIM("Unnamed: 8") = ?
                ORDER BY DATE_ISO DESC
                LIMIT 1
                """,
                (adr, des, se),
            ).fetchone()
        return row

    # ======================================================
    # 🧠 Construire les lignes “comme tu veux”
    # ======================================================
    def _clean_dest_prefix(s: str) -> str:
        t = str(s or "").strip()
        if not t:
            return ""
        # retire préfixes fréquents
        for pref in ["DOMICILE :", "DOMICILE:", "Domicile :", "Domicile:", "Adresse :", "Adresse:"]:
            if t.startswith(pref):
                t = t[len(pref):].strip()
        return t

    def _split_addr_cp_city(full: str):
        """Retourne (adresse, cp, localite) depuis 'Route ..., 71 4050 Chaudfontaine'."""
        s = _clean_dest_prefix(full)
        s = re.sub(r"\s+", " ", s).strip()
        if not s:
            return "", "", ""
        # cas: '... 4050 Chaudfontaine' ou '... 4050  Chaudfontaine'
        m = re.search(r"\b(\d{4,5})\s+([A-Za-zÀ-ÿ\-\' ]+)$", s)
        if m:
            cp = m.group(1).strip()
            city = m.group(2).strip()
            adr = s[:m.start()].strip(" ,")
            return adr, cp, city
        return s, "", ""

    def _build_rows_business(parsed_mail: dict):
        nom = parsed_mail.get("NOM", "")
        tel = parsed_mail.get("Tél", "")
        demandeur = parsed_mail.get("DEMANDEUR", "")
        imputation = parsed_mail.get("IMPUTATION", "")
        prefix = parsed_mail.get("PREFIX", "")

        last_bdc = _db_last_bdc_for_passenger(nom, tel, prefix) or prefix

        rows = []
        for b in (parsed_mail.get("BLOCKS") or []):
            pickup = b.get("pickup", "")
            dropoff = b.get("dropoff", "")

            sens = _sens_from_pick_drop(pickup, dropoff)

            # ✅ TA RÈGLE:
            # - adresse = pickup (YUST Liège)
            # - designation = destination code (JCO si "John Cockerill Seraing")
            if sens == "VERS":
                adresse = pickup
                designation = _designation_from_place(dropoff)  # -> JCO
            else:
                # DE: adresse = destination client/hôtel, designation = aéroport code
                adresse_full = dropoff
                designation = _designation_from_place(pickup)
                adresse, cp, loc = _split_addr_cp_city(adresse_full)

            # prix depuis DB (ignore CP/localité)
            price = _db_lookup_price_for_route(adresse, designation, sens)
            km = h_tva = ttc = typ = pay = ""
            if price:
                km, h_tva, ttc, typ, pay = price

            cp = locals().get("cp", "")
            loc = locals().get("loc", "")

            rows.append({
                "DATE": b.get("DATE", ""),
                "HEURE": b.get("HEURE", ""),
                "CH": "",
                "²²²²": "",
                "IMMAT": "",
                "PAX": "",
                "Reh": "",
                "Siège": "",
                "Unnamed: 8": sens,
                "DESIGNATION": designation,
                "H South": "",
                "Décollage": "",
                "N° Vol": "",
                "Origine": "",
                "GO": "",
                "Num BDC": last_bdc,
                "NOM": nom,
                "ADRESSE": adresse,
                "CP": cp or "",
                "Localité": loc or "",
                "Tél": tel,
                "Type Nav": typ or "",
                "PAIEMENT": (pay or "Facture"),
                "Caisse": "",
                "KM": km or "",
                "H TVA": h_tva or "",
                "TTC": ttc or "",
                "PARKING": "",
                "ATTENTE": "",
                "PEAGE": "",
                "ADM": "",
                "REMARQUE": "",
                "DEMANDEUR": demandeur,
                "IMPUTATION": imputation,
            })

        if not rows:
            rows = [{"REMARQUE": (raw_mail or "")[:200]}]
        return rows

    # ================= ACTIONS =================
    colA, colB = st.columns(2)

    with colA:
        if st.button("🧠 Analyser", key="mail_analyse_btn"):
            rows = []

            # 1) métier corporate (robuste)
            try:
                parsed_mail = _parse_mail_corporate(raw_mail)
                rows = _build_rows_business(parsed_mail)
            except Exception:
                rows = []

            # 2) IA existante si besoin
            if not rows or (len(rows) == 1 and rows[0].get("REMARQUE")):
                try:
                    parsed = parse_mail_to_navette_v2_cached(raw_mail) or {}
                    rows2 = parsed.get("ROWS") or []
                    if rows2:
                        rows = rows2
                except Exception:
                    pass

            if not rows:
                rows = [{"REMARQUE": raw_mail[:200]}]

            st.session_state.mail_parsed = {"ROWS": rows}

    with colB:
        if st.button("🧹 Réinitialiser", key="mail_reset_btn"):
            st.session_state.pop("mail_parsed", None)

    parsed = st.session_state.get("mail_parsed")

    # ================= ÉDITION =================
    if parsed:
        df = pd.DataFrame(parsed["ROWS"]).fillna("")

        planning_cols = [
            "DATE","HEURE","CH","²²²²","IMMAT","PAX","Reh","Siège",
            "Unnamed: 8",
            "DESIGNATION","H South","Décollage","N° Vol","Origine","GO","Num BDC",
            "NOM","ADRESSE","CP","Localité","Tél",
            "Type Nav","PAIEMENT","Caisse","KM","H TVA","TTC",
            "PARKING","ATTENTE","PEAGE","ADM","REMARQUE","DEMANDEUR","IMPUTATION"
        ]
        for c in planning_cols:
            if c not in df.columns:
                df[c] = ""
        df = df[planning_cols]

        # ===== DUPLICATION / RETOUR =====
        if add_retour and len(df) == 1:
            r2 = df.iloc[0].to_dict()
            r2["DATE"], r2["HEURE"] = "", ""
            s = str(r2.get("Unnamed: 8", "")).upper().strip()
            r2["Unnamed: 8"] = "DE" if s == "VERS" else "VERS"
            r2["REMARQUE"] = (str(r2.get("REMARQUE", "")) + " | RETOUR MANUEL").strip()
            df = pd.concat([df, pd.DataFrame([r2])], ignore_index=True)

        # ===== URGENT =====
        if urgence_mode:
            df["REMARQUE"] = df["REMARQUE"].astype(str).apply(
                lambda x: x if "URGENT" in str(x) else f"{x} | URGENT"
            )

        st.markdown("### ✏️ Prévisualisation")
        st.dataframe(df, use_container_width=True)

        df_edit = st.data_editor(df, use_container_width=True, hide_index=True)

        # ===== COPIE EXCEL =====
        def _clean(v):
            return str(v or "").replace(".0", "")

        st.markdown("### 📋 Lignes Excel prêtes à coller")
        st.code(
            "\n".join(
                "\t".join(_clean(rr.get(c, "")) for c in planning_cols)
                for _, rr in df_edit.iterrows()
            ),
            language="tsv"
        )

        # ===== VALIDATION =====
        if st.button("✅ Valider et envoyer", key="mail_validate_btn"):
            payload = df_edit.fillna("").to_dict(orient="records")
            for r in payload:
                r["_URGENT"] = bool(urgence_mode)
                r["_NOTIFY"] = bool(notify_now)

            from database import insert_planning_rows_from_table
            insert_planning_rows_from_table(payload, ignore_conflict=True)

            request_soft_refresh("planning", clear_cache=True, mute_autosync_sec=10)
            request_soft_refresh("admin_tab_mail")
            st.success("✅ Navette(s) ajoutée(s)")
    # ======================================================
    # 🟡 ONGLET À REPORTER DANS EXCEL
    # ======================================================
    with tab_excel:
        st.subheader("🟡 Modifications à reporter dans Excel (Feuil1)")

        from database import list_pending_actions, mark_actions_done
        from utils import update_excel_rows_by_row_key

        actions = list_pending_actions(limit=300)

        if not actions:
            st.success("✅ Aucune modification en attente. Excel et l’application sont alignés.")
        else:
            rows = []
            for (action_id, row_key, action_type, old_value, new_value, user, created_at) in actions:
                rows.append(
                    {
                        "Type": action_type,
                        "Avant": old_value,
                        "Après": new_value,
                        "Modifié par": user,
                        "Date / heure": created_at,
                        "row_key": row_key,
                        "action_id": action_id,
                    }
                )

            df_actions = pd.DataFrame(rows)
            st.info("Modifs faites dans l’app mais pas encore reportées dans Excel (Feuil1).")
            st.dataframe(df_actions.drop(columns=["row_key", "action_id"], errors="ignore"), width="stretch", hide_index=True)

            st.markdown("### 📤 Envoyer ces modifications vers Excel")

            if st.button("📤 Envoyer vers Excel maintenant", type="primary"):
                try:
                    updates = {}
                    action_ids = []

                    for (action_id, row_key, action_type, old_value, new_value, user, created_at) in actions:
                        if action_type != "CH_CHANGE":
                            continue
                        if not row_key:
                            continue
                        updates.setdefault(row_key, {})
                        updates[row_key]["CH"] = new_value
                        action_ids.append(action_id)

                    if not updates:
                        st.warning("Aucune modification 'CH_CHANGE' à envoyer.")
                    else:
                        updated_count = update_excel_rows_by_row_key(updates)
                        mark_actions_done(action_ids)
                        st.success(f"✅ Excel mis à jour ({updated_count} ligne(s))")
                        st.cache_data.clear()
                        st.rerun()

                except Exception as e:
                    st.error(f"Erreur en envoyant vers Excel : {e}")

    # ======================================================
    # 📋 ONGLET TRANSFERTS / SMS  ✅ FIX (dates ISO + WhatsApp client)
    # ======================================================
    with tab_transferts:
        today = date.today()
        start_60j = today - timedelta(days=60)

        col1, col2 = st.columns(2)
        with col1:
            start_date = st.date_input("Date de début", value=start_60j, key="admin_start_date")
        with col2:
            end_date = st.date_input("Date de fin", value=today, key="admin_end_date")

        start_iso = _to_iso(start_date)
        end_iso = _to_iso(end_date)

        # ======================================================
        # 📱 WHATSAPP CLIENT — J+1 / J+3 (AUTO, GROUPÉ)
        # ======================================================
        # ======================================================
        # 📋 VUE TRANSFERTS
        # ======================================================

        # ✅ charge FULL (puis fallback si ton get_planning ne supporte pas source="full")
        df = get_planning(
            start_date=start_iso,
            end_date=end_iso,
            chauffeur=None,
            type_filter=None,
            search="",
            max_rows=5000,
            source="full",
        )

        # fallback “safe”
        if df is None or df.empty:
            df = get_planning(
                start_date=start_iso,
                end_date=end_iso,
                chauffeur=None,
                type_filter=None,
                search="",
                max_rows=5000,
                source="7j",
            )

        if df is None:
            df = pd.DataFrame()

        # 🔒 SÉCURITÉ DATE_ISO (évite incohérences UI / WhatsApp)
        if not df.empty and "DATE_ISO" in df.columns:
            try:
                df["DATE_ISO"] = pd.to_datetime(df["DATE_ISO"], errors="coerce").dt.date
            except Exception:
                pass

        st.caption(f"DEBUG admin transferts — lignes chargées : {len(df)}")

        try:
            df = apply_actions_overrides(df)
        except Exception:
            pass

        if df.empty:
            st.warning("Aucun transfert pour cette période.")
        else:
            # 🔽 filtres
            col3, col4, col5 = st.columns(3)
            with col3:
                bdc_prefix = st.text_input("Filtrer par Num BDC", "", key="admin_bdc_prefix")
            with col4:
                paiement_filter = st.text_input("Filtrer par paiement", "", key="admin_paiement_filter")
            with col5:
                ch_filter = st.text_input("Filtrer par chauffeur", "", key="admin_ch_filter")

            if bdc_prefix.strip() and "Num BDC" in df.columns:
                df = df[df["Num BDC"].astype(str).str.upper().str.startswith(bdc_prefix.upper())]

            if paiement_filter.strip() and "PAIEMENT" in df.columns:
                df = df[df["PAIEMENT"].astype(str).str.upper().str.contains(paiement_filter.upper(), na=False)]

            if ch_filter.strip() and "CH" in df.columns:
                df = df[df["CH"].astype(str).str.upper() == ch_filter.upper()]

            if df.empty:
                st.warning("Aucun transfert après filtres.")
            else:
                sort_mode = st.radio("Tri", ["DATE + HEURE", "CH + DATE + HEURE"], horizontal=True)

                sort_cols = []
                if sort_mode == "CH + DATE + HEURE" and "CH" in df.columns:
                    sort_cols.append("CH")
                for c in ["DATE", "HEURE"]:
                    if c in df.columns:
                        sort_cols.append(c)

                if sort_cols:
                    try:
                        df = df.sort_values(sort_cols)
                    except Exception:
                        pass

                # ✅ Badges
                if "Badges" not in df.columns:
                    try:
                        df["Badges"] = df.apply(navette_badges, axis=1)
                    except Exception:
                        df["Badges"] = ""

                # ======================================================
                # 📋 AFFICHAGE (sans WhatsApp ici)
                # ======================================================
                st.dataframe(df, width="stretch", height=520)

    # ======================================================
    # 🚨 ONGLET URGENCES (✅ déplacé ici)
    # ======================================================
    with tab_urgences:
        st.subheader("🚨 Urgences — actions rapides")

        # Charge DB (safe)
        with get_connection() as conn:
            urgence_df = pd.read_sql_query(
                """
                SELECT *
                FROM planning
                ORDER BY DATE_ISO, HEURE
                """,
                conn,
            )

        # Filtrage urgences
        if urgence_df is not None and not urgence_df.empty:
            adm = urgence_df.get("ADM", pd.Series([""])).fillna("").astype(str).str.upper()
            rem = urgence_df.get("REMARQUE", pd.Series([""])).fillna("").astype(str).str.upper()

            urgence_df = urgence_df[(adm.eq("URGENT")) | (rem.str.contains("URGENCE", na=False))]

        if urgence_df is None or urgence_df.empty:
            st.success("✅ Aucune urgence en cours")
        else:
            for _, rep in urgence_df.iterrows():
                rid = rep.get("id")
                st.markdown(
                    f"**Urgence #{rid}** — {rep.get('DATE', '')} {rep.get('HEURE', '')} — {rep.get('CH', '')}"
                )

                gsm = rep.get("Tél") or rep.get("TEL") or ""
                if gsm:
                    wa_link = f"https://wa.me/{str(gsm).replace(' ', '').replace('/', '').replace('+','')}"
                    st.markdown(f"[💬 WhatsApp chauffeur]({wa_link})")

                pdf_path = rep.get("PDF_PATH")
                if pdf_path:
                    try:
                        with open(pdf_path, "rb") as f:
                            st.download_button(
                                f"📄 PDF mission (#{rid})",
                                data=f,
                                file_name=f"MISSION_URGENTE_{rid}.pdf",
                                mime="application/pdf",
                                key=f"dl_urg_pdf_{rid}",
                            )
                    except Exception:
                        st.warning("⚠️ PDF introuvable")
    # ======================================================
    # ⏱️ ONGLET CALCUL D’HEURES
    # ======================================================
    with tab_heures:
        render_tab_calcul_heures()
  

    # ======================================================
    # 💬 ONGLET WHATSAPP (CLIENTS) — J+1 / J+3
    # ======================================================
    with tab_whatsapp:
        st.subheader("💬 WhatsApp — confirmations clients")
        st.caption("Envoi manuel • 1 message par client et par jour • basé sur planning 7 jours (aujourd’hui → J+7).")

        # --------------------------------------------------
        # Helpers
        # --------------------------------------------------
        def _norm_gsm(x: str) -> str:
            s = str(x or "").strip()
            s = s.replace(" ", "").replace("/", "").replace(".", "").replace("-", "")
            s = s.replace("+", "")
            if s.startswith("00"):
                s = s[2:]
            return s

        def _split_chauffeurs_codes(ch_value: str) -> list:
            raw = str(ch_value or "").upper().replace("*", "").replace(" ", "")
            codes, buf = [], ""
            for c in raw:
                buf += c
                if len(buf) == 2:
                    codes.append(buf)
                    buf = ""
            return [c for c in codes if c]




        def _resolve_dest(x: str) -> str:
            try:
                from database import resolve_location_alias
                return str(resolve_location_alias(x))
            except Exception:
                return str(x or "")

        def _build_client_message(client_name: str, date_txt: str, rows: list) -> str:
            lines = [
                "*🚐 Airports Lines*",
                "",
                f"Bonjour {client_name or ''},",
                "",
                "Voici votre(vos) transfert(s) prévu(s) :",
                "",
                f"📅 {date_txt}",
                "",
            ]

            for r in rows:
                sens = str(r.get("Unnamed: 8", "") or "").strip()
                dest = str(r.get("DESIGNATION", "") or "").strip()
                adr = str(r.get("ADRESSE", "") or "").strip()
                cp = str(r.get("CP", "") or "").strip()
                loc = str(r.get("Localité", "") or "").strip()
                h = str(r.get("HEURE", "") or "").strip()

                ch = str(r.get("CH", "") or "").replace("*", "").strip()
                ch_gsms = r.get("_CH_GSMS") or []

                lines.append(f"⏰ {h}")
                lines.append(f"➡️ {sens} {dest}")

                if adr or cp or loc:
                    lines.append(f"📍 {adr}, {cp} {loc}".replace(".0 ", " ").strip())

                # Chauffeur + téléphone(s)
                if ch:
                    if ch_gsms:
                        lines.append(
                            f"👨‍✈️ Chauffeur : {ch} — 📞 " + " / ".join(ch_gsms)
                        )
                    else:
                        lines.append(f"👨‍✈️ Chauffeur : {ch}")

                lines.append("")

            lines.append("Merci de votre confiance.")
            return "\n".join(lines)


        def _get_client_gsm_col(df_in: pd.DataFrame):
            for c in ["Tél", "TEL", "GSM", "Tel", "Téléphone"]:
                if c in df_in.columns:
                    return c
            return None

        def _load_whatsapp_source_df():
            # Source fixée: aujourd’hui -> J+7, planning 7j
            d0 = date.today()
            d7 = date.today() + timedelta(days=7)
            dfw = get_planning(
                start_date=d0.isoformat(),
                end_date=d7.isoformat(),
                chauffeur=None,
                type_filter=None,
                search="",
                max_rows=5000,
                source="7j",
            )
            if dfw is None:
                dfw = pd.DataFrame()
            return dfw

        def _render_send_links(df_src: pd.DataFrame, target_from: date, target_to: date):
            if df_src is None or df_src.empty:
                st.info("Aucun transfert sur la période source.")
                return

            df_src = df_src.copy()

            # Filtre dates cible (par DATE_ISO en priorité)
            t_from = target_from.isoformat()
            t_to = target_to.isoformat()

            if "DATE_ISO" in df_src.columns:
                df_src["_DATE_ISO"] = df_src["DATE_ISO"].astype(str)
                df_src = df_src[(df_src["_DATE_ISO"] >= t_from) & (df_src["_DATE_ISO"] <= t_to)]
            else:
                # fallback sur DATE (moins fiable)
                df_src["_DATE_TXT"] = df_src.get("DATE", "").astype(str)
                df_src = df_src[df_src["_DATE_TXT"].astype(str).str.contains(str(target_from.year), na=False)]

            if df_src.empty:
                st.info("Aucun transfert pour cette période.")
                return

            # Destination alias + GSM chauffeurs
            if "DESIGNATION" in df_src.columns:
                df_src["DESIGNATION"] = df_src["DESIGNATION"].apply(_resolve_dest)

            def _get_chauffeur_gsms_from_db(ch_value: str) -> list:
                if not ch_value:
                    return []

                gsms = []
                for ch in split_chauffeurs(ch_value):
                    try:
                        tel, _mail = get_chauffeur_contact(ch)
                        if tel:
                            gsms.append(clean_phone(tel))
                    except Exception:
                        pass
                return gsms

            df_src["_CH_GSMS"] = df_src["CH"].apply(_get_chauffeur_gsms_from_db)

            # GSM client
            gsm_col = _get_client_gsm_col(df_src)
            if not gsm_col:
                st.warning("Colonne téléphone client introuvable (Tél/TEL/GSM).")
                return
            df_src["_CLIENT_GSM"] = df_src[gsm_col].apply(_norm_gsm)

            # Groupement 1 message par client et par jour
            if "DATE" in df_src.columns:
                date_group = "DATE"
            elif "DATE_ISO" in df_src.columns:
                date_group = "DATE_ISO"
            else:
                date_group = gsm_col  # fallback (moche mais safe)

            groups = df_src.groupby(["_CLIENT_GSM", date_group], dropna=True)

            sent = 0
            for (gsm, dval), g in groups:
                if not gsm:
                    continue
                client_name = str(g.iloc[0].get("NOM", "") or "").strip()
                date_txt = str(dval or "").strip()
                msg = _build_client_message(client_name, date_txt, g.to_dict(orient="records"))
                try:
                    link = build_whatsapp_link(gsm, msg)
                except Exception:
                    link = f"https://wa.me/{gsm}"
                st.markdown(f"[💬 Envoyer WhatsApp — {client_name or gsm} ({date_txt})]({link})")
                sent += 1

            st.success(f"✅ Messages prêts : {sent}")

        # --------------------------------------------------
        # UI
        # --------------------------------------------------
        df_source = _load_whatsapp_source_df()

        st.caption(
            f"Source : planning 7j • Aujourd’hui : {date.today().strftime('%d/%m/%Y')} • "
            f"Fenêtre source : {date.today().strftime('%d/%m/%Y')} → {(date.today()+timedelta(days=7)).strftime('%d/%m/%Y')}"
        )

        c1, c2 = st.columns(2)
        with c1:
            if st.button("📅 Préparer WhatsApp J+1 (demain)", key="wa_btn_j1"):
                d1 = date.today() + timedelta(days=1)
                _render_send_links(df_source, d1, d1)

        with c2:
            if st.button("📅 Préparer WhatsApp J+3 (3 prochains jours)", key="wa_btn_j3"):
                d1 = date.today() + timedelta(days=1)
                d3 = date.today() + timedelta(days=3)
                _render_send_links(df_source, d1, d3)

        with st.expander("🧪 Debug WhatsApp (source)", expanded=False):
            st.caption(f"Lignes source chargées : {len(df_source)}")
            try:
                st.dataframe(df_source.head(50), use_container_width=True, hide_index=True)
            except Exception:
                pass




# ============================================================
# ⏱️ HELPERS RÈGLES HEURES (OBLIGATOIRES)
# ============================================================

def _coerce_minutes(val) -> int:
    """
    Accepte: 150 | "150" | "2h30" | "2:30" | "2.5"
    Retourne des minutes (int)
    """
    if val is None:
        return 0

    if isinstance(val, (int, float)):
        return int(val * 60) if val < 24 else int(val)

    s = str(val).strip().lower()
    if not s:
        return 0

    # 2h30
    if "h" in s:
        try:
            h, m = s.split("h", 1)
            return int(h) * 60 + int(m or 0)
        except Exception:
            return 0

    # 2:30
    if ":" in s:
        try:
            h, m = s.split(":", 1)
            return int(h) * 60 + int(m)
        except Exception:
            return 0

    # 2.5
    try:
        f = float(s.replace(",", "."))
        return int(f * 60) if f < 24 else int(f)
    except Exception:
        return 0


def _rules_prepare(df_rules: pd.DataFrame) -> pd.DataFrame:
    """
    Normalise les règles pour calcul heures
    Colonnes attendues :
    - ch_base
    - is_star (0/1)
    - sens
    - dest_contains
    - minutes
    """
    if df_rules is None or df_rules.empty:
        return pd.DataFrame()

    df = df_rules.copy()

    for col in ["ch_base", "sens", "dest_contains"]:
        if col not in df.columns:
            df[col] = ""
        df[col] = (
            df[col]
            .fillna("")
            .astype(str)
            .str.upper()
            .str.strip()
        )

    if "is_star" not in df.columns:
        df["is_star"] = 0

    df["is_star"] = df["is_star"].fillna(0).astype(int)

    if "minutes" not in df.columns:
        df["minutes"] = 0

    df["minutes_norm"] = df["minutes"].apply(_coerce_minutes)

    # garder uniquement règles valides
    df = df[df["minutes_norm"] > 0]

    return df
def render_tab_confirmation_chauffeur():
    st.subheader("✅ Confirmation chauffeur")
    st.caption(
        "Validation définitive des navettes après réponse chauffeur. "
        "Les confirmations sont traçables dans l’historique."
    )

    # ======================================================
    # 🧪 TEST DB (ADMIN)
    # ======================================================
    with st.expander("🧪 TEST DB (ADMIN)"):

        from database import get_connection
        import pandas as pd

        with get_connection() as conn:

            total = conn.execute(
                "SELECT COUNT(*) FROM planning"
            ).fetchone()[0]
            st.write("Total lignes planning :", total)

            ack_count = conn.execute(
                "SELECT COUNT(*) FROM planning WHERE ACK_AT IS NOT NULL AND ACK_AT != ''"
            ).fetchone()[0]
            st.write("Lignes avec ACK_AT :", ack_count)

            non_conf = conn.execute(
                "SELECT COUNT(*) FROM planning WHERE COALESCE(CONFIRMED,0) != 1"
            ).fetchone()[0]
            st.write("Non confirmées :", non_conf)

            try:
                count_7j = conn.execute(
                    "SELECT COUNT(*) FROM planning_7j"
                ).fetchone()[0]
                st.write("Lignes planning_7j :", count_7j)
            except Exception as e:
                st.write("Erreur planning_7j :", e)

            df_debug = pd.read_sql(
                """
                SELECT DATE, CH, NOM, ACK_AT, CONFIRMED, DATE_ISO
                FROM planning
                ORDER BY DATE DESC
                LIMIT 20
                """,
                conn,
            )

            st.dataframe(df_debug, use_container_width=True)

            df_ack = pd.read_sql(
                """
                SELECT DATE, DATE_ISO, ACK_AT
                FROM planning
                WHERE ACK_AT IS NOT NULL AND ACK_AT != ''
                ORDER BY DATE DESC
                """,
                conn,
            )

            st.write("🔎 Vérif DATE_ISO des ACK :")
            st.dataframe(df_ack, use_container_width=True)

    # ======================================================
    # SOUS-ONGLETS
    # ======================================================
    tab_confirm, tab_history, tab_messages = st.tabs(
        ["🟢 À confirmer", "🧾 Historique", "📩 Messages chauffeurs"]
    )

    # ======================================================
    # 🟢 À CONFIRMER
    # ======================================================
    with tab_confirm:

        if "confirm_periode" not in st.session_state:
            st.session_state.confirm_periode = "Aujourd’hui"

        periode = st.radio(
            "📅 Navettes à confirmer",
            ["Aujourd’hui", "À partir de demain"],
            horizontal=True,
            key="confirm_periode",
        )

        from database import get_connection
        import pandas as pd

        with get_connection() as conn:
            df = pd.read_sql("SELECT * FROM planning", conn)

        if df is None or df.empty:
            st.info("Aucune navette à afficher.")
            return

        # --------------------------------------------------
        # Nettoyage fort DATE_ISO
        # --------------------------------------------------
        df["DATE_ISO"] = (
            df["DATE_ISO"]
            .astype(str)
            .str.strip()
        )

        df["DATE_ISO_DT"] = pd.to_datetime(
            df["DATE_ISO"],
            errors="coerce"
        )

        today_dt = pd.to_datetime(date.today())

        if periode == "Aujourd’hui":
            df = df[df["DATE_ISO_DT"] >= today_dt]
        else:
            df = df[df["DATE_ISO_DT"] > today_dt]

        # --------------------------------------------------
        # Filtres métier
        # --------------------------------------------------
        df["CH"] = df["CH"].fillna("NON_ATTRIBUE")
        df = df[df.get("IS_INDISPO", 0) == 0]
        df = df[df.get("CONFIRMED", 0) != 1]

        if df.empty:
            st.info("Aucune navette à confirmer pour cette période.")
            return

        # --------------------------------------------------
        # Normalisation chauffeur
        # --------------------------------------------------
        df["CH_ROOT"] = df["CH"].apply(normalize_ch_code)

        chauffeurs = (
            df.groupby("CH_ROOT")
            .size()
            .sort_index()
            .index
            .tolist()
        )

        for ch_root in chauffeurs:

            df_ch = df[df["CH_ROOT"] == ch_root].copy()
            has_reply = df_ch["ACK_AT"].notna().any()

            badge = "🟠🆕" if has_reply else "🟠"
            title = f"{badge} Chauffeur {ch_root} — {len(df_ch)} navette(s)"

            with st.expander(title, expanded=has_reply):

                df_ch = df_ch.sort_values(
                    by=["DATE_ISO_DT", "HEURE"],
                    ascending=[True, True],
                )

                for _, row in df_ch.iterrows():

                    date_txt = row.get("DATE", "—")
                    heure_txt = row.get("HEURE", "—")
                    client = row.get("NOM", "—")
                    chauffeur = row.get("CH", "—")

                    sens = str(row.get("Unnamed: 8", "") or "").strip()
                    trajet = resolve_client_alias(
                        row.get("DESIGNATION", row.get("DESTINATION", "—"))
                    )

                    adresse_complete = build_full_address_from_row(row)

                    st.markdown(
                        f"""
                        ### 📅 {date_txt} ⏰ {heure_txt}
                        👤 **Client :** {client}  
                        👨‍✈️ **Chauffeur :** {chauffeur}  
                        ➡️ **Sens :** {sens}  
                        📍 **Adresse :** {adresse_complete or "—"}  
                        🧭 **Destination :** {trajet}
                        """
                    )

                    # -----------------------------
                    # ✉️ Réponse chauffeur
                    # -----------------------------
                    ack_at = row.get("ACK_AT")
                    ack_txt = row.get("ACK_TEXT")

                    if ack_at:
                        st.markdown(f"📝 **Réponse chauffeur** ({ack_at})")
                        st.info(ack_txt or "— réponse vide —")
                    else:
                        st.warning("⏳ Aucune réponse chauffeur")

                    # -----------------------------
                    # 💬 Réponse admin au chauffeur
                    # -----------------------------
                    admin_reply_key = f"admin_reply_{row['id']}"
                    admin_reply = st.text_area(
                        "💬 Message pour le chauffeur",
                        key=admin_reply_key,
                        height=90,
                        placeholder="Ex : OK pour cette navette 👍 / ou indique le problème à corriger…",
                    )

                    col_ok, col_ko = st.columns(2)

                    # ✅ OK → Confirme + message chauffeur
                    with col_ok:
                        if st.button(
                            "✅ C’est OK",
                            key=f"ok_{row['id']}",
                            use_container_width=True,
                        ):
                            now_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                            update_planning_row(
                                row["id"],
                                {
                                    "CONFIRMED": 1,
                                    "CONFIRMED_AT": now_iso,
                                    "ADMIN_REPLY": admin_reply.strip() if admin_reply and admin_reply.strip() else "OK",
                                    "ADMIN_REPLY_AT": now_iso,
                                    "ADMIN_REPLY_READ": 0,
                                },
                            )

                            with get_connection() as conn:
                                conn.execute(
                                    """
                                    INSERT INTO planning_audit
                                    (ts, user, action, row_key, details)
                                    VALUES (?, ?, ?, ?, ?)
                                    """,
                                    (
                                        now_iso,
                                        st.session_state.get("username"),
                                        "CONFIRM_OK",
                                        row.get("row_key"),
                                        f"CH={chauffeur}",
                                    ),
                                )
                                conn.commit()

                            st.toast("🟢 Validé — le chauffeur est informé", icon="✅")
                            st.rerun()

                    # ❌ Pas OK → Message chauffeur, pas confirmé
                    with col_ko:
                        if st.button(
                            "❌ Pas OK / Problème",
                            key=f"ko_{row['id']}",
                            use_container_width=True,
                        ):
                            if not admin_reply or not admin_reply.strip():
                                st.warning("Merci d’indiquer le problème dans le message.")
                            else:
                                now_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                                update_planning_row(
                                    row["id"],
                                    {
                                        "ADMIN_REPLY": admin_reply.strip(),
                                        "ADMIN_REPLY_AT": now_iso,
                                        "ADMIN_REPLY_READ": 0,
                                    },
                                )

                                with get_connection() as conn:
                                    conn.execute(
                                        """
                                        INSERT INTO planning_audit
                                        (ts, user, action, row_key, details)
                                        VALUES (?, ?, ?, ?, ?)
                                        """,
                                        (
                                            now_iso,
                                            st.session_state.get("username"),
                                            "CONFIRM_KO",
                                            row.get("row_key"),
                                            admin_reply.strip()[:200],
                                        ),
                                    )
                                    conn.commit()

                                st.toast("⚠️ Problème envoyé au chauffeur", icon="📩")
                                st.rerun()

                    st.divider()

    # ======================================================
    # HISTORIQUE
    # ======================================================
    with tab_history:

        st.markdown("### 🧾 Historique des confirmations")

        from database import get_connection
        import pandas as pd

        with get_connection() as conn:
            df_hist = pd.read_sql(
                """
                SELECT ts AS "Date",
                       user AS "Admin",
                       action AS "Action",
                       row_key AS "Navette",
                       details AS "Détails"
                FROM planning_audit
                ORDER BY ts DESC
                LIMIT 500
                """,
                conn,
            )

        if df_hist.empty:
            st.info("Aucune confirmation enregistrée.")
            return

        st.dataframe(df_hist, use_container_width=True, hide_index=True)

    # ======================================================
    # MESSAGES
    # ======================================================
    with tab_messages:
        st.info("📩 Vue messages chauffeurs — à venir")

def _match_rule_minutes(rules_norm, ch, sens, dest):
    """
    Retourne le nombre de minutes selon les règles définies par l'utilisateur.

    Règles :
    - ch : 'NP', 'NP*', '*', 'ALL'
    - sens : 'VERS', 'DE', '*'
    - dest : texte contenu dans la destination (BRU, ZAVENTEM, CDG), ou '*'
    - la première règle la plus spécifique qui matche gagne
    """

    if rules_norm is None or rules_norm.empty:
        return 0

    ch = str(ch or "").strip().upper()
    sens = str(sens or "").strip().upper()
    dest = str(dest or "").strip().upper()

    has_star = "*" in ch
    ch_base = ch.replace("*", "").strip()

    for _, rule in rules_norm.iterrows():

        # -----------------------------
        # Chauffeur
        # -----------------------------
        rule_ch = str(rule.get("ch_base", "")).strip().upper()
        rule_star = int(rule.get("is_star", 0))

        if rule_ch not in ("", "ALL", "*"):
            if rule_ch != ch_base:
                continue

        if rule_star == 1 and not has_star:
            continue

        # -----------------------------
        # Sens
        # -----------------------------
        rule_sens = str(rule.get("sens", "")).strip().upper()
        if rule_sens not in ("", "*"):
            if rule_sens != sens:
                continue

        # -----------------------------
        # Destination
        # -----------------------------
        rule_dest = str(rule.get("dest_contains", "")).strip().upper()
        if rule_dest not in ("", "*"):
            if rule_dest not in dest:
                continue

        # -----------------------------
        # ✅ MATCH OK → minutes NORMALISÉES
        # -----------------------------
        minutes = int(rule.get("minutes_norm", 0) or 0)
        return minutes

    return 0


def render_tab_calcul_heures():
    import pandas as pd
    import streamlit as st
    from datetime import date, timedelta, datetime
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas

    from database import (
        get_time_rules_df,
        save_time_rules_df,
        _detect_sens_dest_from_row,
        _minutes_to_hhmm,
        split_chauffeurs,
        get_last_caisse_paid_dates,
        init_time_adjustments_table,
        get_time_adjustments_df,
        insert_time_adjustment,
        get_connection,
    )

    st.subheader("⏱️ Calcul d’heures")

    tab_calc, tab_rules, tab_caisse = st.tabs([
        "📊 Heures (60 jours)",
        "⚙️ Règles (éditables)",
        "💶 Caisse non rentrée (60j)",
    ])

    # ======================================================
    # 📊 HEURES
    # ======================================================
    with tab_calc:
        today = date.today()

        mode = st.radio(
            "📅 Période",
            ["Mois complet", "Période personnalisée"],
            horizontal=True,
        )

        if mode == "Mois complet":
            mois = st.selectbox("Mois", list(range(1, 13)), index=today.month - 1)
            annee = st.selectbox("Année", list(range(2026, today.year + 1)), index=len(list(range(2026, today.year + 1))) - 1)

            d1 = date(annee, mois, 1)
            d2 = (
                date(annee + 1, 1, 1) - timedelta(days=1)
                if mois == 12
                else date(annee, mois + 1, 1) - timedelta(days=1)
            )
        else:
            c1, c2 = st.columns(2)
            with c1:
                d1 = st.date_input("Du", date(2026, 1, 1))
            with c2:
                d2 = st.date_input("Au", today)

        if d1 > d2:
            st.error("La date de début est après la date de fin.")
            return

        # 🔒 LECTURE STRICTE TABLE planning (ANTI EXPLOSION)
        with get_connection() as conn:
            df_hours = pd.read_sql_query(
                """
                SELECT *
                FROM planning
                WHERE
                    COALESCE(IS_INDISPO,0) = 0
                    AND COALESCE(IS_SUPERSEDED,0) = 0
                    AND DATE_ISO >= ?
                    AND DATE_ISO <= ?
                ORDER BY DATE_ISO, HEURE
                """,
                conn,
                params=(d1.isoformat(), d2.isoformat()),
            )

        if df_hours.empty:
            st.info("Aucune navette sur cette période.")
            return

        # Anti-doublons
        if "row_key" in df_hours.columns:
            df_hours = df_hours.drop_duplicates(subset=["row_key"])
        elif "id" in df_hours.columns:
            df_hours = df_hours.drop_duplicates(subset=["id"])

        # Chauffeurs
        df_hours["CH_LIST"] = (
            df_hours["CH"]
            .fillna("")
            .astype(str)
            .str.upper()
            .apply(split_chauffeurs)
        )

        # Sens / destination
        df_hours[["SENS", "DEST"]] = df_hours.apply(
            lambda r: pd.Series(_detect_sens_dest_from_row(r.to_dict())),
            axis=1,
        )

        rules_norm = _rules_prepare(get_time_rules_df())

        totals = {}
        rows_not_matched = []
        debug_rows = []

        for _, r in df_hours.iterrows():

            ch_list = r.get("CH_LIST") or []
            if not ch_list:
                continue

            sens = r.get("SENS")
            dest = r.get("DEST")

            # Normalisation destination via alias si dispo
            try:
                from database import resolve_location_alias
                dest = resolve_location_alias(dest)
            except Exception:
                pass

            # Attente éventuelle (minutes)
            attente_raw = (
                r.get("ATTENTE")
                or r.get("Attente")
                or r.get("ATTENTE_MIN")
                or r.get("MIN_ATTENTE")
                or 0
            )
            attente_min = _coerce_minutes(attente_raw)

            for ch in ch_list:
                ch_u = str(ch or "").strip().upper()
                if not ch_u:
                    continue

                minutes_base = _match_rule_minutes(
                    rules_norm,
                    ch_u,
                    sens,
                    dest,
                )

                total_minutes = int(minutes_base or 0) + int(attente_min or 0)

                debug_rows.append({
                    "DATE": r.get("DATE_ISO"),
                    "CH": ch_u,
                    "SENS": sens,
                    "DEST": dest,
                    "BASE_MIN": int(minutes_base or 0),
                    "ATTENTE_MIN": int(attente_min or 0),
                    "MINUTES": int(total_minutes or 0),
                })

                if minutes_base <= 0 and attente_min <= 0:
                    rows_not_matched.append({
                        "DATE": r.get("DATE_ISO"),
                        "CH": ch_u,
                        "SENS": sens,
                        "DEST": dest,
                    })
                    continue

                if total_minutes > 0:
                    totals[ch_u] = totals.get(ch_u, 0) + total_minutes

        if totals:
            df_tot = pd.DataFrame([
                {
                    "Chauffeur": ch,
                    "Heures": _minutes_to_hhmm(mins),
                }
                for ch, mins in sorted(totals.items())
            ])

            st.markdown("#### ✅ Heures calculées")
            st.dataframe(df_tot, use_container_width=True, hide_index=True)

        if rows_not_matched:
            st.markdown("#### ⚠️ Navettes sans règle")
            st.dataframe(pd.DataFrame(rows_not_matched), use_container_width=True, hide_index=True)

        # ================= DEBUG =================
        with st.expander("🧪 Debug calcul heures", expanded=False):
            df_dbg = pd.DataFrame(debug_rows)

            st.caption(f"Navettes analysées : {len(df_dbg)}")
            st.caption(f"Minutes totales : {df_dbg['MINUTES'].sum()}")

            st.markdown("🔴 Top 10 navettes les plus lourdes")
            st.dataframe(
                df_dbg.sort_values("MINUTES", ascending=False).head(10),
                use_container_width=True,
                hide_index=True,
            )

            st.markdown("⚠️ Navettes à 0 minute")
            st.dataframe(
                df_dbg[df_dbg["MINUTES"] == 0].head(20),
                use_container_width=True,
                hide_index=True,
            )


            # 📤 Export paie (CSV)
            try:
                csv_tot = df_tot.to_csv(index=False, sep=";", encoding="utf-8")
                st.download_button(
                    "📤 Télécharger heures chauffeurs (CSV)",
                    data=csv_tot,
                    file_name=f"heures_chauffeurs_{d1.strftime('%Y%m%d')}_{d2.strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                )
            except Exception:
                pass

        if rows_not_matched:
            st.markdown("#### ⚠️ Navettes non calculées (ajouter des règles)")
            st.dataframe(
                pd.DataFrame(rows_not_matched),
                use_container_width=True,
                hide_index=True,
            )

        # ==================================================
        # ➕ AJUSTEMENTS MANUELS PAR CHAUFFEUR (attente, détour…)
        # ==================================================
        init_time_adjustments_table()

        d1_iso = d1.strftime("%Y-%m-%d")
        d2_iso = d2.strftime("%Y-%m-%d")

        st.markdown("#### ➕ Ajustements manuels (par chauffeur)")
        colx1, colx2, colx3, colx4 = st.columns([1, 1, 1, 3])

        with colx1:
            adj_ch = st.selectbox("Chauffeur", sorted(totals.keys()) if totals else [])
        with colx2:
            adj_hours = st.text_input("Heures (+/-)", "0.0", help="Ex: 0.5, -0.5, 1.0, 2h30")
        with colx3:
            adj_scope = st.selectbox("Portée", ["Sur la période", "Toujours"], index=0)
        with colx4:
            adj_reason = st.text_input("Raison", "", placeholder="Attente, détour, accident, etc.")

        if st.button("✅ Ajouter l’ajustement"):
            mins = _coerce_minutes(adj_hours)
            if mins == 0:
                st.warning("Ajustement invalide (0).")
            elif not adj_ch:
                st.warning("Choisis un chauffeur.")
            else:
                if adj_scope == "Toujours":
                    df_from = None
                    df_to = None
                else:
                    df_from = d1_iso
                    df_to = d2_iso

                insert_time_adjustment(
                    chauffeur=adj_ch,
                    date_from_iso=df_from,
                    date_to_iso=df_to,
                    minutes=mins,
                    reason=adj_reason,
                )
                st.success("Ajustement ajouté ✅")
                st.rerun()

        df_adj = get_time_adjustments_df(d1_iso, d2_iso)

        # Somme ajustements par chauffeur
        adj_map = {}
        if df_adj is not None and not df_adj.empty:
            df_adj["chauffeur"] = df_adj["chauffeur"].astype(str).str.upper().str.strip()
            df_adj["minutes"] = pd.to_numeric(df_adj["minutes"], errors="coerce").fillna(0).astype(int)
            for ch_u, g in df_adj.groupby("chauffeur"):
                adj_map[ch_u] = int(g["minutes"].sum())

            st.caption("Ajustements enregistrés sur la période :")
            st.dataframe(df_adj[["chauffeur", "date_from", "date_to", "minutes", "reason", "created_at"]], use_container_width=True, hide_index=True)

        # ==================================================
        # ✅ TABLE FINALE (calcul + ajustements)
        # ==================================================
        if totals:
            df_final = pd.DataFrame(
                [
                    {
                        "Chauffeur": ch,
                        "Heures calculées": _minutes_to_hhmm(totals.get(ch, 0)),
                        "Ajustement": _minutes_to_hhmm(adj_map.get(ch, 0)) if adj_map.get(ch, 0) else "—",
                        "Heures finales": _minutes_to_hhmm(totals.get(ch, 0) + adj_map.get(ch, 0)),
                    }
                    for ch in sorted(totals.keys())
                ]
            )

            st.markdown("#### ✅ Total final (avec ajustements)")
            st.dataframe(df_final, use_container_width=True, hide_index=True)

            # 📄 Export PDF
            def _hours_pdf_bytes(df_pdf: pd.DataFrame, d1_local: date, d2_local: date) -> bytes:
                buf = BytesIO()
                c = canvas.Canvas(buf, pagesize=A4)
                w, h = A4
                x = 2 * cm
                y = h - 2 * cm
                c.setFont("Helvetica-Bold", 14)
                c.drawString(x, y, "Heures chauffeurs")
                y -= 0.8 * cm
                c.setFont("Helvetica", 10)
                c.drawString(x, y, f"Période : {d1_local.strftime('%d/%m/%Y')} → {d2_local.strftime('%d/%m/%Y')}")
                y -= 1.0 * cm

                c.setFont("Helvetica-Bold", 10)
                cols = ["Chauffeur", "Heures calculées", "Ajustement", "Heures finales"]
                col_w = [4*cm, 4*cm, 4*cm, 4*cm]
                for i, col in enumerate(cols):
                    c.drawString(x + sum(col_w[:i]), y, col)
                y -= 0.5 * cm
                c.setFont("Helvetica", 10)

                for _, rr in df_pdf.iterrows():
                    if y < 2.2 * cm:
                        c.showPage()
                        y = h - 2 * cm
                        c.setFont("Helvetica-Bold", 10)
                        for i, col in enumerate(cols):
                            c.drawString(x + sum(col_w[:i]), y, col)
                        y -= 0.5 * cm
                        c.setFont("Helvetica", 10)

                    c.drawString(x + 0, y, str(rr.get("Chauffeur", "")))
                    c.drawString(x + col_w[0], y, str(rr.get("Heures calculées", "")))
                    c.drawString(x + col_w[0] + col_w[1], y, str(rr.get("Ajustement", "")))
                    c.drawString(x + col_w[0] + col_w[1] + col_w[2], y, str(rr.get("Heures finales", "")))
                    y -= 0.45 * cm

                c.save()
                buf.seek(0)
                return buf.read()

            pdf_bytes = _hours_pdf_bytes(df_final, d1, d2)
            st.download_button(
                "📄 Export PDF heures chauffeurs",
                data=pdf_bytes,
                file_name=f"heures_chauffeurs_{d1.strftime('%Y%m%d')}_{d2.strftime('%Y%m%d')}.pdf",
                mime="application/pdf",
            )
    # ======================================================
    # ⚙️ RÈGLES (ÉDITABLES)
    # ======================================================
    with tab_rules:
        st.markdown("### ⚙️ Règles de calcul des heures")
        st.caption(
            "Les règles sont lues de haut en bas. "
            "La première règle qui correspond est utilisée."
        )

        df_rules = get_time_rules_df()

        # Adapter DB → UI
        if df_rules is None or df_rules.empty:
            df_ui = pd.DataFrame(
                [
                    {
                        "ch": "*",
                        "sens": "*",
                        "dest": "*",
                        "heures": "1",
                    }
                ]
            )
        else:
            # 🔁 DB => UI
            # DB: ch_base + is_star (0/1)  => UI: ch = "*" ou "FA" ou "FA*" si tu veux distinguer
            df_ui = pd.DataFrame()

            if "ch_base" in df_rules.columns:
                ch_base = df_rules["ch_base"].fillna("").astype(str).str.upper().str.strip()
            else:
                ch_base = pd.Series("", index=df_rules.index)

            is_star = df_rules.get("is_star", 0)
            try:
                is_star = is_star.fillna(0).astype(int)
            except Exception:
                is_star = 0

            # "*" = règle pour tous
            df_ui["ch"] = [
                "*" if (cb in ("", "*")) else (cb + "*" if int(star or 0) == 1 else cb)
                for cb, star in zip(ch_base.tolist(), list(is_star))
            ]

            df_ui["sens"] = df_rules.get("sens", "*")
            df_ui["dest"] = df_rules.get("dest_contains", "*")
            df_ui["heures"] = (df_rules.get("minutes", 0).fillna(0).astype(float) / 60).astype(str)




        df_edit = st.data_editor(
            df_ui,
            use_container_width=True,
            hide_index=True,
            num_rows="dynamic",
            column_config={
                "ch": st.column_config.TextColumn(
                    "Chauffeur",
                    help="*, NP, NP* …",
                ),
                "sens": st.column_config.TextColumn(
                    "Sens",
                    help="VERS, DE ou *",
                ),
                "dest": st.column_config.TextColumn(
                    "Destination",
                    help="BRU, ZAVENTEM, CDG, *",
                ),
                "heures": st.column_config.TextColumn(
                    "Heures",
                    help="2.5, 2h30, 150…",
                ),
            },
        )

        from database import is_time_rules_locked, set_time_rules_locked, get_time_rules_audit_df

        locked = is_time_rules_locked()
        if locked:
            st.warning("🔒 Règles verrouillées (modification désactivée)")
        
        # Boutons lock/unlock (admin)
        if st.session_state.get('role') == 'admin':
            c1, c2, c3 = st.columns([1,1,2])
            with c1:
                if st.button("🔒 Verrouiller", disabled=locked):
                    set_time_rules_locked(True, user=st.session_state.get('username',''), details='UI')
                    st.toast('Règles verrouillées', icon='🔒')
                    st.rerun()
            with c2:
                if st.button("🔓 Déverrouiller", disabled=not locked):
                    set_time_rules_locked(False, user=st.session_state.get('username',''), details='UI')
                    st.toast('Règles déverrouillées', icon='🔓')
                    st.rerun()
            with c3:
                st.caption("Le verrou empêche toute modification des règles.")

        if st.button("💾 Sauvegarder les règles", disabled=locked):
            try:
                save_time_rules_df(df_edit, user=st.session_state.get('username',''))
                st.success("Règles sauvegardées ✅")
                st.rerun()
            except PermissionError:
                st.error("🔒 Règles verrouillées")

        with st.expander("🧾 Historique (audit)", expanded=False):
            try:
                df_a = get_time_rules_audit_df(limit=30)
                st.dataframe(df_a, use_container_width=True, hide_index=True)
            except Exception:
                st.info("Aucun audit disponible.")




    # ======================================================
    # 💶 CAISSE NON RENTRÉE — GESTION BUREAU
    # ======================================================
            # ======================================================
        # 🗺️ ALIAS LIEUX / CLIENTS (JCO/JCC/GUIL/BRU/...)
        # ======================================================
        st.markdown("---")
        st.markdown("### 🗺️ Alias lieux / clients (normalisation)")

        try:
            from database import get_location_aliases_df, save_location_aliases_df
            df_alias = get_location_aliases_df()
            df_alias_edit = st.data_editor(
                df_alias,
                use_container_width=True,
                hide_index=True,
                num_rows="dynamic",
                column_config={
                    "code": st.column_config.TextColumn("Code (ex: BRU)"),
                    "label": st.column_config.TextColumn("Libellé (ex: Zaventem)"),
                },
                key="alias_editor",
            )
            if st.button("💾 Sauvegarder les alias", key="save_aliases"):
                save_location_aliases_df(df_alias_edit)
                st.success("✅ Alias sauvegardés")
                st.cache_data.clear()
                st.rerun()
        except Exception as e:
            st.error(f"Erreur alias: {e}")

        # ======================================================
        # 🧠 MÉMOIRE PRIX & DEMANDEUR (mail → navette)
        # ======================================================
        st.markdown("---")
        st.markdown("### 🧠 Mémoire (prix / demandeur)")
        colM1, colM2 = st.columns(2)

        with colM1:
                st.markdown("#### 💶 Prix mémorisés")
                try:
                        from database import init_price_memory_table
                        init_price_memory_table()
                        with get_connection() as conn:
                                df_pm = pd.read_sql(
                                        """
                                        SELECT dest_code, sens, nom_key, paiement,
                                               prix_ttc, prix_htva, updated_at
                                        FROM price_memory
                                        ORDER BY updated_at DESC
                                        LIMIT 200
                                        """,
                                        conn,
                                )
                        st.dataframe(df_pm, use_container_width=True, hide_index=True)
                except Exception:
                        st.info("Aucun prix mémorisé ou erreur lecture.")

        with colM2:
                st.markdown("#### 👤 Demandeurs mémorisés")
                try:
                        from database import init_requester_memory_table
                        init_requester_memory_table()
                        with get_connection() as conn:
                                df_rm = pd.read_sql(
                                        """
                                        SELECT demandeur, societe, tva, bdc,
                                               imputation, updated_at
                                        FROM requester_memory
                                        ORDER BY updated_at DESC
                                        LIMIT 200
                                        """,
                                        conn,
                                )
                        st.dataframe(df_rm, use_container_width=True, hide_index=True)
                except Exception:
                        st.info("Aucun demandeur mémorisé ou erreur lecture.")

        # ======================================================
        # 💶 ONGLET CAISSE
        # ======================================================
        with tab_caisse:
                st.markdown("### 💶 Caisse non rentrée (60 jours)")
                consume_soft_refresh("caisse")

                render_excel_modified_indicator()

                if st.button("🔄 Rafraîchir la caisse depuis Excel"):
                        request_soft_refresh("caisse")

                # ----------------- Période -----------------
                today = date.today()
                d1 = today - timedelta(days=60)
                if d1 < date(2026, 1, 1):
                        d1 = date(2026, 1, 1)

                # ----------------- Chauffeur -----------------
                chs = get_chauffeurs_for_ui()
                ch_filter = st.selectbox(
                        "👨‍✈️ Chauffeur",
                        ["(Tous)"] + chs,
                )
                if ch_filter == "(Tous)":
                        ch_filter = None

                # ==================================================
                # 🔒 LECTURE DB DIRECTE (PAS get_planning)
                # ==================================================
                with get_connection() as conn:
                        df_cash = pd.read_sql_query(
                                """
                                SELECT *
                                FROM planning
                                WHERE
                                        COALESCE(IS_INDISPO,0) = 0
                                        AND COALESCE(IS_SUPERSEDED,0) = 0
                                        AND LOWER(COALESCE(PAIEMENT,'')) = 'caisse'
                                        AND DATE_ISO >= ?
                                        AND DATE_ISO <= ?
                                ORDER BY DATE_ISO, HEURE
                                """,
                                conn,
                                params=(d1.isoformat(), today.isoformat()),
                        )


                # ==================================================
                # 📊 RÉCAP PAR CHAUFFEUR (comme la vue chauffeur)
                # ==================================================
                if not df_cash.empty:
                        df_cash2 = df_cash.copy()
                        df_cash2["Caisse"] = pd.to_numeric(df_cash2.get("Caisse", 0), errors="coerce").fillna(0.0)
                        if "CAISSE_PAYEE" in df_cash2.columns:
                                df_cash2 = df_cash2[df_cash2["CAISSE_PAYEE"].fillna(0).astype(int).eq(0)]
                        recap = (
                                df_cash2.groupby(df_cash2["CH"].fillna("").astype(str).str.strip().str.upper())["Caisse"]
                                .sum()
                                .reset_index()
                                .rename(columns={"CH": "Chauffeur", "Caisse": "Caisse due (€)"})
                                .sort_values("Caisse due (€)", ascending=False)
                        )
                        recap = recap[recap["Chauffeur"] != ""]
                        if not recap.empty:
                                st.markdown("#### 💶 Caisse due par chauffeur (60 jours)")
                                st.dataframe(recap, use_container_width=True, height=220)

                # DEBUG
                if not df_cash.empty:
                        st.caption(
                                f"DEBUG caisse — lignes chargées : {len(df_cash)} | "
                                f"date min = {df_cash['DATE_ISO'].min()} | "
                                f"date max = {df_cash['DATE_ISO'].max()}"
                        )

                if df_cash is None or df_cash.empty:
                        st.success("✅ Aucune caisse à rentrer")
                        st.stop()

                df_cash = df_cash.copy()

                # ----------------- Filtre chauffeur -----------------
                if ch_filter and "CH" in df_cash.columns:
                        ch_norm = normalize_ch_code(ch_filter)
                        df_cash = df_cash[
                                df_cash["CH"]
                                .fillna("")
                                .astype(str)
                                .str.upper()
                                .str.replace("*", "", regex=False)
                                .str.replace(" ", "", regex=False)
                                .str.startswith(ch_norm)
                        ]

                if df_cash.empty:
                        st.success("✅ Aucune caisse à rentrer")
                        st.stop()

                # ----------------- Montant > 0 -----------------
                df_cash["Caisse"] = (
                        df_cash.get("Caisse", pd.Series(0, index=df_cash.index))
                        .pipe(pd.to_numeric, errors="coerce")
                        .fillna(0)
                )
                df_cash = df_cash[df_cash["Caisse"] > 0]

                if df_cash.empty:
                        st.success("✅ Aucune caisse à rentrer")
                        st.stop()

                # ----------------- Dernière caisse payée -----------------
                try:
                        last_paid = get_last_caisse_paid_dates(ch_filter)
                except Exception:
                        last_paid = {}

                if last_paid:
                        df_cash["_date_iso"] = pd.to_datetime(df_cash["DATE_ISO"], errors="coerce")

                        def _after_last_paid(row):
                                ch_norm = normalize_ch_code(str(row.get("CH", "")))
                                lp = last_paid.get(ch_norm)
                                if not lp:
                                        return True
                                try:
                                        return row["_date_iso"].date() > datetime.fromisoformat(lp).date()
                                except Exception:
                                        return True

                        df_cash = df_cash[df_cash.apply(_after_last_paid, axis=1)]
                        df_cash = df_cash.drop(columns=["_date_iso"], errors="ignore")

                # ----------------- Non payées uniquement -----------------
                if "CAISSE_PAYEE" in df_cash.columns:
                        df_cash = df_cash[df_cash["CAISSE_PAYEE"].fillna(0).astype(int) == 0]

                if df_cash.empty:
                        st.success("✅ Aucune caisse à rentrer")
                        st.stop()

                # ==================================================
                # 📋 TABLE ÉDITABLE
                # ==================================================
                df_out = df_cash[["id", "DATE", "CH", "NOM", "Caisse"]].copy()

                df_out.rename(
                        columns={
                                "NOM": "Client",
                                "Caisse": "Montant €",
                        },
                        inplace=True,
                )

                df_out["Valider"] = False

                edited = st.data_editor(
                        df_out,
                        use_container_width=True,
                        hide_index=True,
                        column_config={
                                "Valider": st.column_config.CheckboxColumn("Payé"),
                        },
                )

                total_due = float(edited["Montant €"].sum())
                st.metric("💶 Total à rentrer", f"{total_due:.2f} €")

                # ==================================================
                # 📝 COMMENTAIRE
                # ==================================================
                comment = st.text_input(
                        "📝 Commentaire (ex : finalement paiement bancontact)",
                        "",
                )

                # ==================================================
                # ✅ VALIDATION
                # ==================================================
                colv1, colv2 = st.columns(2)

                with colv1:
                        if st.button("✅ Valider la sélection"):
                                ids = edited[edited["Valider"] == True]["id"].tolist()

                                if not ids:
                                        st.warning("Aucune ligne sélectionnée.")
                                else:
                                        now_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                        for rid in ids:
                                                apply_row_update(
                                                        int(rid),
                                                        {
                                                                "CAISSE_PAYEE": 1,
                                                                "CAISSE_PAYEE_AT": now_iso,
                                                                "CAISSE_COMMENT": comment or "Validé manuellement",
                                                        },
                                                        lock_row=True,
                                                        touch_is_new=True,
                                                )

                                        # 📤 Export DB -> Excel (sans conflit) : uniquement la sélection
                                        try:
                                                export_db_changes_to_excel_dropbox(row_ids=[int(x) for x in ids])
                                        except Exception:
                                                pass

                                        st.success("Caisse validée ✅")
                                        request_soft_refresh("caisse")


                with colv2:
                        if ch_filter and st.button("✅ Tout valider pour ce chauffeur"):
                                ch_norm = normalize_ch_code(ch_filter)
                                now_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                                # IDs concernés (caisse non payée)
                                with get_connection() as conn:
                                        rows = conn.execute(
                                                """
                                                SELECT id
                                                FROM planning
                                                WHERE UPPER(REPLACE(REPLACE(CH,'*',''),' ','')) LIKE ?
                                                  AND LOWER(COALESCE(PAIEMENT,'')) = 'caisse'
                                                  AND COALESCE(CAISSE_PAYEE,0) = 0
                                                  AND DATE_ISO >= ?
                                                """,
                                                (f"{ch_norm}%", d1.isoformat()),
                                        ).fetchall()

                                ids2 = [int(r[0]) for r in rows] if rows else []

                                if not ids2:
                                        st.info("Aucune ligne à valider pour ce chauffeur.")
                                else:
                                        for rid in ids2:
                                                apply_row_update(
                                                        rid,
                                                        {
                                                                "CAISSE_PAYEE": 1,
                                                                "CAISSE_PAYEE_AT": now_iso,
                                                                "CAISSE_COMMENT": comment or "Validé globalement",
                                                        },
                                                        lock_row=True,
                                                        touch_is_new=True,
                                                )

                                        try:
                                                export_db_changes_to_excel_dropbox(row_ids=ids2)
                                        except Exception:
                                                pass

                                        st.success(f"Toute la caisse de {ch_filter} est validée ✅")
                                        request_soft_refresh("caisse")





# ==========================================================================
#  ONGLET Admin — Validation des indispos
# ==========================================================================

def render_tab_indispo_admin():
    st.subheader("🚫 Indisponibilités chauffeurs")

    # Toutes les demandes
    df = get_indispo_requests()

    if df.empty:
        st.info("Aucune demande d'indisponibilité.")
        return

    st.markdown("### 🔍 Toutes les demandes")
    st.dataframe(df, use_container_width=True, height=250)

    # Demandes en attente
    if "STATUT" not in df.columns:
        st.error("Colonne STATUT manquante dans la table chauffeur_indispo.")
        return

    df_pending = df[df["STATUT"] == "EN_ATTENTE"].copy()

    if df_pending.empty:
        st.info("Aucune demande en attente.")
        return

    st.warning(f"🔔 {len(df_pending)} demande(s) en attente")
    st.markdown("### 📝 Traiter une demande")

    # Sélecteur avec un joli label
    def _format_option(row):
        ch = str(row.get("CH", "") or "")
        d = str(row.get("DATE", "") or "")
        h1 = str(row.get("HEURE_DEBUT", "") or "")
        h2 = str(row.get("HEURE_FIN", "") or "")
        com = str(row.get("COMMENTAIRE", "") or "")

        label = f"#{row['id']} — {ch} {d} {h1}→{h2}"
        if com:
            label += f" — {com[:40]}"
        return label

    options = [int(v) for v in df_pending["id"].tolist()]
    labels_map = {int(row["id"]): _format_option(row) for _, row in df_pending.iterrows()}

    selected_id = st.selectbox(
        "Sélectionne une demande",
        options=options,
        format_func=lambda x: labels_map.get(int(x), f"#{x}"),
    )

    row = df_pending[df_pending["id"] == selected_id].iloc[0]

    colA, colB = st.columns(2)

    with colA:
        if st.button("✅ Accepter"):
            # Création d'une ligne INDISPO dans le planning
            data_planning = {
                "DATE": row.get("DATE", ""),
                "HEURE": row.get("HEURE_DEBUT", ""),
                "²²²²": row.get("HEURE_FIN", ""),
                "CH": row.get("CH", ""),
                "REMARQUE": f"INDISPO {row.get('CH','')} - {row.get('COMMENTAIRE','')}",
            }
            planning_id = insert_planning_row(data_planning)

            # MAJ statut + lien vers la ligne planning
            set_indispo_status(int(row["id"]), "ACCEPTEE", planning_id=planning_id)

            st.success("Indisponibilité acceptée et ajoutée au planning.")
            st.rerun()

    with colB:
        if st.button("❌ Refuser"):
            set_indispo_status(int(row["id"]), "REFUSEE")
            st.error("La demande a été refusée.")
            st.rerun()

# ============================================================
#   MAIN — ROUTAGE PAR RÔLE (admin / restricted / driver)
# ============================================================

def main():
    print("🔄 MAIN RUN", datetime.now())

    # =====================================================
    # 🔧 DEBUG_STEP
    # 0 = mode normal (tous les onglets)
    # 1 = admin sans onglets
    # 2 = admin : Planning seul
    # 3 = admin : Confirmation seule
    # 4 = admin : Vue jour (mobile) seule
    # 5 = admin : Tableau / Édition seul
    # 6 = admin : Clients / Historique seul
    # 7 = admin : Vue Chauffeur seule
    # 8 = admin : Feuil2 / Chauffeurs seul
    # 9 = admin : Feuil3 seul
    # 10 = admin : Admin transferts seul
    # 11 = admin : Excel ↔ DB seul
    # 12 = admin : Indispos chauffeurs seul
    # =====================================================
    DEBUG_STEP = 0   # 👈 change UNIQUEMENT ce chiffre

    # ======================================
    # 1️⃣ INITIALISATION SESSION + DB
    # ======================================
    init_session_state()
    bootstrap_login_persistence()

    # 🔁 1 rerun max pour laisser le JS écrire al_cid / al_session en cookie/localStorage
    # (très important dans l'app Streamlit)
    if "CID_BOOT_ONCE" not in st.session_state:
        st.session_state.CID_BOOT_ONCE = True
        st.rerun()

    init_db_once()
    init_all_db_once()
    ensure_persistent_sessions_table()

    # 1 rerun max pour laisser le JS remonter cid/session
    if "BOOTSTRAP_TRY" not in st.session_state:
        st.session_state.BOOTSTRAP_TRY = 0

    restored = restore_login_from_cookie()

    # ✅ Si connecté + remember_me, on attache la session à CE device (client_id) une seule fois
    # (évite FA -> MA et autres mélanges)
    if st.session_state.get("logged_in") and st.session_state.get("remember_me"):
        if not st.session_state.get("_persist_state_saved", False):
            login = str(st.session_state.get("username") or "").strip().lower()
            token = str(st.session_state.get("session_token") or "").strip()
            cid = str(get_client_id() or "").strip()

            if login and token and cid:
                ok = save_persistent_session(login, token, True)
                if ok:
                    st.session_state["_persist_state_saved"] = True

    # Si pas connecté, on autorise 1 rerun si cid/session pas encore visibles
    if not st.session_state.get("logged_in"):
        cid = str(get_client_id() or "").strip()
        sess = get_login_cookie()

        if st.session_state.BOOTSTRAP_TRY < 1 and (not cid or not sess):
            st.session_state.BOOTSTRAP_TRY += 1
            st.rerun()
    else:
        st.session_state.BOOTSTRAP_TRY = 0
    # ======================================
    # 2️⃣ LOGIN
    # ======================================
    if not st.session_state.logged_in:
        login_screen()
        st.stop()
    # ======================================
    # 3️⃣ UI MINIMALE
    # ======================================
    render_top_bar()
    role = st.session_state.role

    # 🔄 Synchro silencieuse (uniquement si Excel modifié)
    # auto_sync_planning_if_needed()  # ⛔ DISABLED DEBUG

    # ====================== ADMIN ===========================
    if role == "admin":

        # ---------------- DEBUG MODES ----------------
        if DEBUG_STEP == 1:
            st.success("✅ TEST 1 : Admin sans onglets")
            st.stop()

        if DEBUG_STEP == 2:
            st.success("✅ TEST 2 : Planning seul")
            (tab1,) = st.tabs(["📅 Planning"])
            with tab1:
                render_tab_planning()
            st.stop()

        if DEBUG_STEP == 3:
            st.success("✅ TEST 3 : Confirmation chauffeur seule")
            (tab_confirm,) = st.tabs(["✅ Confirmation chauffeur"])
            with tab_confirm:
                render_tab_confirmation_chauffeur()
            st.stop()

        if DEBUG_STEP == 4:
            st.success("✅ TEST 4 : Vue jour (mobile) seule")
            (tab2,) = st.tabs(["⚡ Vue jour (mobile)"])
            with tab2:
                render_tab_quick_day_mobile()
            st.stop()

        if DEBUG_STEP == 5:
            st.success("✅ TEST 5 : Tableau / Édition seul")
            (tab3,) = st.tabs(["📊 Tableau / Édition"])
            with tab3:
                render_tab_table()
            st.stop()

        if DEBUG_STEP == 6:
            st.success("✅ TEST 6 : Clients / Historique seul")
            (tab4,) = st.tabs(["🔍 Clients / Historique"])
            with tab4:
                render_tab_clients()
            st.stop()

        if DEBUG_STEP == 7:
            st.success("✅ TEST 7 : Vue Chauffeur seule")
            (tab5,) = st.tabs(["🚖 Vue Chauffeur"])
            with tab5:
                render_tab_vue_chauffeur()
            st.stop()

        if DEBUG_STEP == 8:
            st.success("✅ TEST 8 : Feuil2 / Chauffeurs seul")
            (tab6,) = st.tabs(["👨‍✈️ Feuil2 / Chauffeurs"])
            with tab6:
                render_tab_chauffeurs()
            st.stop()

        if DEBUG_STEP == 9:
            st.success("✅ TEST 9 : Feuil3 seul")
            (tab7,) = st.tabs(["📄 Feuil3"])
            with tab7:
                render_tab_feuil3()
            st.stop()

        if DEBUG_STEP == 10:
            st.success("✅ TEST 10 : Admin transferts seul")
            (tab8,) = st.tabs(["📦 Admin transferts"])
            with tab8:
                render_tab_admin_transferts()
            st.stop()

        if DEBUG_STEP == 11:
            st.success("✅ TEST 11 : Excel ↔ DB seul")
            (tab9,) = st.tabs(["📂 Excel ↔ DB"])
            with tab9:
                render_tab_excel_sync()
            st.stop()

        if DEBUG_STEP == 12:
            st.success("✅ TEST 12 : Indispos chauffeurs seul")
            (tab10,) = st.tabs(["🚫 Indispos chauffeurs"])
            with tab10:
                render_tab_indispo_admin()
            st.stop()

        # ---------------- MODE NORMAL ----------------
        pending = count_pending_confirmations()
        confirm_label = (
            f"✅ Confirmation chauffeur ({pending})"
            if pending > 0
            else "✅ Confirmation chauffeur"
        )

        (
            tab1,
            tab_confirm,
            tab2,
            tab3,
            tab4,
            tab5,
            tab6,
            tab7,
            tab8,
            tab9,
            tab10,
        ) = st.tabs(
            [
                "📅 Planning",
                confirm_label,
                "⚡ Vue jour (mobile)",
                "📊 Tableau / Édition",
                "🔍 Clients / Historique",
                "🚖 Vue Chauffeur",
                "👨‍✈️ Feuil2 / Chauffeurs",
                "📄 Feuil3",
                "📦 Admin transferts",
                "📂 Excel ↔ DB",
                "🚫 Indispos chauffeurs",
            ]
        )

        with tab1:
            render_tab_planning()
        with tab_confirm:
            render_tab_confirmation_chauffeur()
        with tab2:
            render_tab_quick_day_mobile()
        with tab3:
            render_tab_table()
        with tab4:
            render_tab_clients()
        with tab5:
            render_tab_vue_chauffeur()
        with tab6:
            render_tab_chauffeurs()
        with tab7:
            render_tab_feuil3()
        with tab8:
            render_tab_admin_transferts()
        with tab9:
            render_tab_excel_sync()
        with tab10:
            render_tab_indispo_admin()

    # ==================== RESTRICTED ========================
    elif role == "restricted":
        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(
            [
                "📅 Planning",
                "📊 Tableau / Édition",
                "🔍 Clients / Historique",
                "🚖 Vue Chauffeur",
                "👨‍✈️ Feuil2 / Chauffeurs",
                "📄 Feuil3",
            ]
        )

        with tab1:
            render_tab_planning()
        with tab2:
            render_tab_table()
        with tab3:
            render_tab_clients()
        with tab4:
            render_tab_vue_chauffeur()
        with tab5:
            render_tab_chauffeurs()
        with tab6:
            render_tab_feuil3()

    # ==================== DRIVER ============================
    elif role == "driver":
        ch_code = st.session_state.get("chauffeur_code")
        if not ch_code:
            st.error("Aucun code chauffeur configuré.")
            return

        tab1, tab2 = st.tabs(["🚖 Mon planning", "🚫 Mes indispos"])
        with tab1:
            render_tab_chauffeur_driver()
        with tab2:
            render_tab_indispo_driver(ch_code)

    # ==================== ERREUR ============================
    else:
        st.error(f"Rôle inconnu : {role}")


if __name__ == "__main__":
    main()

import base64

def build_printable_html(df):
    html_table = df.to_html(index=False)
    return f"""
    <html>
    <head>
        <style>
            @page {{
                size: A4 landscape;
                margin: 10mm;
            }}
            body {{
                font-family: Arial, sans-serif;
                zoom: 0.85;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                font-size: 10px;
            }}
            th, td {{
                border: 1px solid #ccc;
                padding: 4px;
                text-align: left;
            }}
            th {{
                background-color: #f0f0f0;
            }}
        </style>
    </head>
    <body onload="window.print();">
        <h2>Planning chauffeur</h2>
        {html_table}
    </body>
    </html>
    """


